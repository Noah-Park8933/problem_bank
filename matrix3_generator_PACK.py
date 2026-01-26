# matrix3_generator_PACK.py
# ------------------------------------------------------------
# Matrix3 출제기 (2연관 1독립 ONLY, 연관쌍 AB/AD/BD 랜덤, 문제에서 연관쌍 숨김)
# - 재조합률 r은 문제에 명시(필수로 줘야 유일정답이 잘 나옴)
# - 주어진 정보:
#   1) 자손의 "가능한 유전자형 가짓수" (distinct genotype count)
#   2) 임의의 자손 유전자형 1개와 그 확률
# - 요구: P1, P2 유전자형 추론 + 자손 유전자형 분포(확률) 계산
# - unique solver: 위 단서로 가능한 (연관쌍, P1, P2, 상(phase)) 후보가 유일해야 통과
# - 출력: DOCX(2단, 문제/정답분리), PACK JSON 저장
#
# 실행:
#   pip install python-docx
#   python matrix3_generator_PACK.py
# ------------------------------------------------------------

import os
import json
import time
import uuid
import random
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional, Any
from fractions import Fraction

from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

# -----------------------------
# CONFIG
# -----------------------------
N_PROBLEMS = 30
FONT_NAME = "바탕"
FONT_SIZE_PT = 9

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")

MODULE_CODE = "MAT3"
ID_PREFIX = "MAT3_"

# 재조합률 후보(0.5는 연관 의미가 약해져 솔버가 흔들릴 수 있어 제외)
R_CANDIDATES = [Fraction(0), Fraction(1, 10), Fraction(1, 5), Fraction(3, 10), Fraction(2, 5)]  # 0,0.1,0.2,0.3,0.4

# 생성 시도
MAX_WORLD_TRIES = 20000

# 문제 단서 난이도 조절(확률 너무 작거나 너무 크면 제외)
MIN_TARGET_PROB = Fraction(1, 32)
MAX_TARGET_PROB = Fraction(31, 32)

# -----------------------------
# Utility
# -----------------------------
def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)

def frac_str(f: Fraction) -> str:
    if f.denominator == 1:
        return str(f.numerator)
    return f"{f.numerator}/{f.denominator}"

def geno_to_str3(g: Dict[str, str]) -> str:
    # {"A":"Aa","B":"bb","D":"Dd"} -> "Aa bb Dd" -> 사용자 요구: 붙여쓰기 "AaBbdd" 형태
    return f"{g['A']}{g['B']}{g['D']}"

def childgeno_to_str3(cg: Dict[str, str]) -> str:
    return f"{cg['A']}{cg['B']}{cg['D']}"

def simplify_prob_map(prob_map: Dict[str, Fraction]) -> Dict[str, Fraction]:
    # 이미 Fraction이라 OK. 0 제거만.
    return {k: v for k, v in prob_map.items() if v != 0}

def weighted_choice(items: List[Tuple[Any, Fraction]]) -> Any:
    total = sum([w for _, w in items], Fraction(0))
    r = Fraction(random.randint(0, 10**9), 10**9) * total
    s = Fraction(0)
    for val, w in items:
        s += w
        if r <= s:
            return val
    return items[-1][0]

# -----------------------------
# Genetics core
# -----------------------------
ALLELES = {
    "A": ("A", "a"),
    "B": ("B", "b"),
    "D": ("D", "d"),
}

GENO_STATES = {
    "A": ["AA", "Aa", "aa"],
    "B": ["BB", "Bb", "bb"],
    "D": ["DD", "Dd", "dd"],
}

PAIR_CHOICES = [("A", "B"), ("A", "D"), ("B", "D")]  # linked pair candidates


@dataclass(frozen=True)
class Parent:
    geno: Dict[str, str]        # locus -> genotype string (e.g., "Aa")
    phase: Optional[Tuple[str, str]]  # haplotype pair for linked loci if needed (e.g., ("AB","ab")) else None


def locus_gamete_probs(geno: str, locus: str) -> Dict[str, Fraction]:
    # returns allele -> prob
    maj, mino = ALLELES[locus]
    if geno == maj + maj:
        return {maj: Fraction(1)}
    if geno == mino + mino:
        return {mino: Fraction(1)}
    return {maj: Fraction(1, 2), mino: Fraction(1, 2)}  # hetero


def possible_haplotype_pairs_for_linked(l1: str, l2: str, g1: str, g2: str) -> List[Tuple[str, str]]:
    # Determine possible phased haplotype pairs consistent with diploid genotypes at l1,l2
    # Example: A: Aa, B: Bb -> possible phases: AB/ab or Ab/aB
    a1, a2 = g1[0], g1[1]
    b1, b2 = g2[0], g2[1]

    # if either locus homozygous, phase is essentially fixed (only one type at that locus)
    # But we still can represent as haplotype pair.
    # Build all haplotype pairs by pairing one allele from locus1 with one from locus2 on each chromatid.
    # There are two chromatids: hap1, hap2.
    # A alleles are a1,a2; B alleles are b1,b2.
    # Two ways to pair:
    #   (a1 with b1) and (a2 with b2)
    #   (a1 with b2) and (a2 with b1)
    h1 = a1 + b1
    h2 = a2 + b2
    k1 = a1 + b2
    k2 = a2 + b1

    # normalize ordering within pair to avoid duplicates
    pairs = []
    pairs.append(tuple(sorted([h1, h2])))
    pairs.append(tuple(sorted([k1, k2])))
    # remove duplicates
    uniq = []
    for p in pairs:
        if p not in uniq:
            uniq.append(p)
    return uniq


def linked_gamete_probs(hap_pair: Tuple[str, str], r: Fraction) -> Dict[str, Fraction]:
    # hap_pair: two haplotypes like ("AB","ab") or ("Ab","aB") etc.
    h1, h2 = hap_pair
    if h1 == h2:
        return {h1: Fraction(1)}
    # If haplotypes differ at both loci: recombination matters.
    # If differ at only one locus: crossover doesn't create new types; still 1/2 each.
    diff = sum(1 for i in range(2) if h1[i] != h2[i])
    if diff <= 1:
        return {h1: Fraction(1, 2), h2: Fraction(1, 2)}
    # diff == 2
    # parental types: (1-r)/2 each
    # recombinant types: r/2 each
    recomb1 = h1[0] + h2[1]
    recomb2 = h2[0] + h1[1]
    return {
        h1: (Fraction(1) - r) / 2,
        h2: (Fraction(1) - r) / 2,
        recomb1: r / 2,
        recomb2: r / 2
    }


def combine_alleles_to_geno(a: str, b: str, locus: str) -> str:
    # make genotype like "Aa" not "aA"
    maj, mino = ALLELES[locus]
    if a == b:
        return a + b
    # hetero:
    return maj + mino


def offspring_distribution(p1: Parent, p2: Parent, linked_pair: Tuple[str, str], r: Fraction) -> Dict[str, Fraction]:
    l1, l2 = linked_pair
    # independent locus:
    indep = ({"A", "B", "D"} - set(linked_pair)).pop()

    # P gamete distributions:
    # linked haplotype probs:
    assert p1.phase is not None and p2.phase is not None
    g1_link = linked_gamete_probs(p1.phase, r)
    g2_link = linked_gamete_probs(p2.phase, r)

    g1_ind = locus_gamete_probs(p1.geno[indep], indep)
    g2_ind = locus_gamete_probs(p2.geno[indep], indep)

    # combine to full gametes (hap + allele)
    # gamete representation: (hap, allele_indep)
    gam1 = []
    for hap, ph in g1_link.items():
        for al, pa in g1_ind.items():
            gam1.append(((hap, al), ph * pa))
    gam2 = []
    for hap, ph in g2_link.items():
        for al, pa in g2_ind.items():
            gam2.append(((hap, al), ph * pa))

    # offspring genotype distribution over 3 loci, canonical string key
    dist: Dict[str, Fraction] = {}
    for (hap1, al1), p_g1 in gam1:
        for (hap2, al2), p_g2 in gam2:
            # linked loci from hap
            a_l1 = hap1[0]
            b_l2 = hap1[1]
            a2_l1 = hap2[0]
            b2_l2 = hap2[1]

            child = {"A": "", "B": "", "D": ""}
            child[l1] = combine_alleles_to_geno(a_l1, a2_l1, l1)
            child[l2] = combine_alleles_to_geno(b_l2, b2_l2, l2)
            child[indep] = combine_alleles_to_geno(al1, al2, indep)

            key = f"{child['A']}{child['B']}{child['D']}"
            dist[key] = dist.get(key, Fraction(0)) + (p_g1 * p_g2)
    return simplify_prob_map(dist)


def distinct_genotype_count(dist: Dict[str, Fraction]) -> int:
    return sum(1 for _k, v in dist.items() if v != 0)


# -----------------------------
# World generation
# -----------------------------
def random_parent_genotype() -> Dict[str, str]:
    return {L: random.choice(GENO_STATES[L]) for L in ["A", "B", "D"]}

def build_parent_with_phase(geno: Dict[str, str], linked_pair: Tuple[str, str]) -> List[Parent]:
    l1, l2 = linked_pair
    phase_pairs = possible_haplotype_pairs_for_linked(l1, l2, geno[l1], geno[l2])
    # return all possible phase variants (for solver enumeration)
    out = []
    for ph in phase_pairs:
        out.append(Parent(geno=geno, phase=ph))
    return out


def sample_world() -> Tuple[Tuple[str, str], Fraction, Parent, Parent, Dict[str, Fraction]]:
    # choose linked pair randomly
    linked_pair = random.choice(PAIR_CHOICES)
    r = random.choice(R_CANDIDATES)

    # random genotypes
    g1 = random_parent_genotype()
    g2 = random_parent_genotype()

    # ensure not identical too often (조금 다양하게)
    if g1 == g2 and random.random() < 0.7:
        g2 = random_parent_genotype()

    # pick a random phase among possible
    p1_candidates = build_parent_with_phase(g1, linked_pair)
    p2_candidates = build_parent_with_phase(g2, linked_pair)
    p1 = random.choice(p1_candidates)
    p2 = random.choice(p2_candidates)

    dist = offspring_distribution(p1, p2, linked_pair, r)

    return linked_pair, r, p1, p2, dist


# -----------------------------
# Clue making + unique solver
# -----------------------------
@dataclass
class Clue:
    genotype_count: int
    target_child: str
    target_prob: Fraction


def pick_target_child(dist: Dict[str, Fraction]) -> Optional[Tuple[str, Fraction]]:
    # pick a child genotype with mid probability (not too tiny)
    items = [(k, v) for k, v in dist.items() if MIN_TARGET_PROB <= v <= MAX_TARGET_PROB]
    if not items:
        return None
    # favor moderate probabilities
    weights = []
    for k, v in items:
        # weight peaks near 1/8~1/4
        w = v if v <= Fraction(1, 2) else (Fraction(1) - v)
        weights.append((k, w))
    return weighted_choice(weights), dict(weights)[weighted_choice(weights)]  # not stable

def pick_target_child_stable(dist: Dict[str, Fraction]) -> Optional[Tuple[str, Fraction]]:
    items = [(k, v) for k, v in dist.items() if MIN_TARGET_PROB <= v <= MAX_TARGET_PROB]
    if not items:
        return None
    # build weight list once
    witems = []
    for k, v in items:
        w = v if v <= Fraction(1, 2) else (Fraction(1) - v)
        if w <= 0:
            w = Fraction(1, 1000)
        witems.append((k, w))
    ksel = weighted_choice(witems)
    return ksel, dist[ksel]


def make_clue(dist: Dict[str, Fraction]) -> Optional[Clue]:
    cnt = distinct_genotype_count(dist)
    pick = pick_target_child_stable(dist)
    if pick is None:
        return None
    tg, tp = pick
    return Clue(genotype_count=cnt, target_child=tg, target_prob=tp)


def enumerate_all_candidates(clue: Clue, r: Fraction) -> List[Tuple[Tuple[str, str], Parent, Parent]]:
    # problem hides which pair is linked
    # solver tries all linked pairs + all genotypes + phases consistent and checks clues
    sols = []

    # enumerate parental genotypes (3^3 each = 27; total 729 pairs) × phase variants × 3 linked pairs
    # manageable.
    all_g = []
    for a in GENO_STATES["A"]:
        for b in GENO_STATES["B"]:
            for d in GENO_STATES["D"]:
                all_g.append({"A": a, "B": b, "D": d})

    for linked_pair in PAIR_CHOICES:
        for g1 in all_g:
            for g2 in all_g:
                p1s = build_parent_with_phase(g1, linked_pair)
                p2s = build_parent_with_phase(g2, linked_pair)
                for p1 in p1s:
                    for p2 in p2s:
                        dist = offspring_distribution(p1, p2, linked_pair, r)
                        if distinct_genotype_count(dist) != clue.genotype_count:
                            continue
                        if dist.get(clue.target_child, Fraction(0)) != clue.target_prob:
                            continue
                        sols.append((linked_pair, p1, p2))
    return sols


def generate_unique_problem() -> Dict[str, Any]:
    for t in range(1, MAX_WORLD_TRIES + 1):
        linked_pair, r, p1, p2, dist = sample_world()
        clue = make_clue(dist)
        if clue is None:
            continue

        # unique solver
        sols = enumerate_all_candidates(clue, r)
        if len(sols) != 1:
            continue

        # Passed
        return {
            "linked_pair": linked_pair,
            "r": r,
            "P1": p1,
            "P2": p2,
            "dist": dist,
            "clue": clue,
            "solver_solutions": sols,
            "tries": t,
        }

    raise RuntimeError("유일정답 문제 생성 실패: MAX_WORLD_TRIES 증가 또는 단서 설계 변경 필요")


# -----------------------------
# Markdown builders for PACK
# -----------------------------
def dist_to_md_table(dist: Dict[str, Fraction]) -> str:
    # sorted by probability desc
    rows = sorted(dist.items(), key=lambda kv: (-kv[1], kv[0]))
    md = "| 자손 유전자형 | 확률 |\n|---|---|\n"
    for g, p in rows:
        md += f"| {g} | {frac_str(p)} |\n"
    return md

def build_problem_text_md(pnum: int, r: Fraction, clue: Clue) -> str:
    # Note: linked info hidden intentionally
    lines = []
    lines.append(f"**[문제 {pnum}] Matrix3 (2연관 1독립)**")
    lines.append("")
    lines.append("A/a, B/b, D/d 3쌍의 유전자에 의해 결정되는 단일인자 일반유전을 가정한다(모두 상염색체).")
    lines.append("세 유전자쌍 중 **두 유전자쌍은 연관**되어 있고, 나머지 한 유전자쌍은 **독립**이다.")
    lines.append(f"연관된 두 유전자쌍의 **재조합률 r = {frac_str(r)}** 이다.")
    lines.append("")
    lines.append(f"부모 P1과 P2를 교배시킬 때 나오는 자손의 **가능한 유전자형 가짓수는 {clue.genotype_count}가지**이다.")
    lines.append(f"또한 자손의 유전자형이 **{clue.target_child}** 일 확률은 **{frac_str(clue.target_prob)}** 이다.")
    return "\n".join(lines)

def build_ask_line_md() -> str:
    return "P1, P2의 유전자형을 구하고, 자손의 유전자형 종류와 각 확률을 구하시오."

def build_answer_md(world: Dict[str, Any]) -> str:
    p1: Parent = world["P1"]
    p2: Parent = world["P2"]
    linked_pair = world["linked_pair"]
    r = world["r"]
    dist = world["dist"]

    lines = []
    lines.append(f"- 연관 유전자쌍: ({linked_pair[0]}/{linked_pair[1]})  (재조합률 r={frac_str(r)})")
    lines.append(f"- P1 유전자형: **{geno_to_str3(p1.geno)}**  (phase={p1.phase[0]}/{p1.phase[1]})")
    lines.append(f"- P2 유전자형: **{geno_to_str3(p2.geno)}**  (phase={p2.phase[0]}/{p2.phase[1]})")
    lines.append("")
    lines.append("**[자손 유전자형 분포]**")
    lines.append(dist_to_md_table(dist))
    return "\n".join(lines)

# -----------------------------
# PACK writer (standalone, no external problem_pack needed)
# -----------------------------
def new_problem_id() -> str:
    return ID_PREFIX + uuid.uuid4().hex[:10]

def save_pack_json(items: List[Dict[str, Any]], out_dir: str) -> str:
    ensure_dir(out_dir)
    path = os.path.join(out_dir, f"PACK_{MODULE_CODE}_{time.strftime('%Y%m%d-%H%M%S')}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(
            {"module": MODULE_CODE, "created": time.time(), "items": items},
            f,
            ensure_ascii=False,
            indent=2
        )
    return path

# -----------------------------
# DOCX helpers (2-column per page, 2 problems/page)
# -----------------------------
def set_normal_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    # for Korean font name compatibility
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    except Exception:
        pass
    style.font.size = Pt(FONT_SIZE_PT)

def add_problem_block(cell, problem_text: str, ask_line: str):
    # problem_text and ask_line are plain text with newlines
    for line in problem_text.split("\n"):
        p = cell.add_paragraph(line)
        if line.startswith("**[문제"):
            # crude bold marker
            p.runs[0].bold = True
    cell.add_paragraph("")
    cell.add_paragraph(ask_line)

def add_answer_block(doc: Document, idx: int, pid: str, answer_md: str):
    title = doc.add_paragraph(f"[정답/해설 {idx}]  ID: {pid}")
    title.runs[0].bold = True
    for line in answer_md.split("\n"):
        doc.add_paragraph(line)

def make_docx_and_pack():
    ensure_dir(OUT_DIR)
    doc = Document()
    set_normal_style(doc)

    pack_items: List[Dict[str, Any]] = []

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        w = generate_unique_problem()
        pid = new_problem_id()
        clue: Clue = w["clue"]
        problem_text_md = build_problem_text_md(i, w["r"], clue)
        ask_line_md = build_ask_line_md()
        answer_md = build_answer_md(w)

        # store
        problems.append((pid, w, problem_text_md, ask_line_md, answer_md))

        payload = {
            "r": float(w["r"]),  # streamlit 편의용
            "genotype_count": clue.genotype_count,
            "target_child": clue.target_child,
            "target_prob": frac_str(clue.target_prob),
            "solution": {
                "linked_pair": list(w["linked_pair"]),
                "P1": geno_to_str3(w["P1"].geno),
                "P2": geno_to_str3(w["P2"].geno),
                "phase_P1": list(w["P1"].phase),
                "phase_P2": list(w["P2"].phase),
            },
            "dist": {k: frac_str(v) for k, v in w["dist"].items()},
            "tries": w["tries"],
        }

        pack_items.append({
            "id": pid,
            "module": MODULE_CODE,
            "id_prefix": ID_PREFIX,
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "answer_md": answer_md,
            "payload": payload,
        })

    # ---- DOCX: 2-column, 2 problems per page (table with 1 row, 2 cols)
    idx = 0
    pnum = 1
    while idx < len(problems):
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        left = t.rows[0].cells[0]
        right = t.rows[0].cells[1]

        pid, _w, ptxt, ask, _ans = problems[idx]
        add_problem_block(left, ptxt.replace("**", ""), ask)
        idx += 1
        pnum += 1

        if idx < len(problems):
            pid, _w, ptxt, ask, _ans = problems[idx]
            add_problem_block(right, ptxt.replace("**", ""), ask)
            idx += 1
            pnum += 1

        if idx < len(problems):
            doc.add_page_break()

    # ---- Answers at end
    doc.add_page_break()
    h = doc.add_paragraph("[정답/해설]")
    h.runs[0].bold = True
    for i, (pid, _w, _ptxt, _ask, ans) in enumerate(problems, start=1):
        add_answer_block(doc, i, pid, ans)
        doc.add_paragraph("")

    # save files
    docx_path = os.path.join(OUT_DIR, f"Matrix3_{time.strftime('%Y%m%d-%H%M%S')}.docx")
    doc.save(docx_path)

    pack_path = save_pack_json(pack_items, OUT_DIR)

    print("✅ 생성 완료")
    print("DOCX:", docx_path)
    print("PACK:", pack_path)


if __name__ == "__main__":
    make_docx_and_pack()
