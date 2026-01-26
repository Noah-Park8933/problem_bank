# gene_detecting_ox_v4_randomX_randomLinkage_NO_SHUFFLE.py
# ------------------------------------------------------------
# ✅ Gene Detecting O/X 자동 출제기 (X연관 개수 0~2 랜덤 + 상염색체 연관/비연관 명시)
#
# 표 방향:
#   - 가로(열) = 세포: A B C D E F
#   - 세로(행) = 유전자: 가 나 다 라 마 바
#
# 핵심:
# 1) A,B,C=1번 사람 / D,E,F=2번 사람 (고정, 섞지 않음)
# 2) 성염색체(X) 유전자좌 개수: 0~2 랜덤
#    - X에 2개면 "연관되어 함께 이동"으로 처리(같은 X에서 같이 선택)
# 3) 상염색체 유전자좌도:
#    - 서로 다른 상염색체(연관 없음) 또는
#    - 2개가 같은 상염색체에 있어 연관(함께 이동) 랜덤
# 4) 문제 지문에 "X연관 개수 + 상염색체 연관 여부 + 연관이면 함께 이동"을 매문항 명시
# 5) 정답 타겟팅(①~⑤) 유지
#    - 선지: ①ㄱ ②ㄴ ③ㄱ,ㄴ ④ㄱ,ㄷ ⑤ㄴ,ㄷ
#
# 실행:
#   pip install python-docx
#   python gene_detecting_ox_v4_randomX_randomLinkage_NO_SHUFFLE.py
# ------------------------------------------------------------
import time
import os
import random
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)



# ============================================================
# CONFIG
# ============================================================
N_PROBLEMS = 30

ANSWER_COUNTS = {1: 6, 2: 6, 3: 7, 4: 6, 5: 5}

CELL_COLS = ["A", "B", "C", "D", "E", "F"]
GENE_ROWS = ["가", "나", "다", "라", "마", "바"]

MAX_TOTAL_MASK = 12
MAX_MASK_PER_GENE = 2
MIN_KNOWN_PER_CELL_COL = 2

MAX_TRIES_PER_PROBLEM = 900
RELAX_AFTER_1 = 350
RELAX_AFTER_2 = 650

CHOICE_TO_TRUESET = {
    1: {"ㄱ"},
    2: {"ㄴ"},
    3: {"ㄱ", "ㄴ"},
    4: {"ㄱ", "ㄷ"},
    5: {"ㄴ", "ㄷ"},
}


# ============================================================
# 핵상 판정(출제 안정화)
# - (O/O) 쌍이 하나라도 있으면 2n
# - 그 외에는 O 개수 기준: O>=4 ->2n, O<=3 -> n (O=3은 n 고정)
# ============================================================
def infer_ploidy(row_gene_presence: Dict[str, str], allele_pairs: List[Tuple[str, str]]) -> str:
    for a, b in allele_pairs:
        if row_gene_presence[a] == "O" and row_gene_presence[b] == "O":
            return "2n"
    o_cnt = sum(1 for v in row_gene_presence.values() if v == "O")
    return "2n" if o_cnt >= 4 else "n"


def describe_ploidy_reason(row_gene_presence: Dict[str, str], allele_pairs: List[Tuple[str, str]]) -> str:
    for a, b in allele_pairs:
        if row_gene_presence[a] == "O" and row_gene_presence[b] == "O":
            return f"같은 유전자좌({a}/{b})가 (O/O)이므로 핵상은 2n이다."
    o_cnt = sum(1 for v in row_gene_presence.values() if v == "O")
    if o_cnt >= 4:
        return f"O의 개수가 {o_cnt}개(>50%)이므로 핵상은 2n이다."
    return f"O의 개수가 {o_cnt}개(≤50% 포함)이므로 핵상은 n이다(O=3은 n 처리)."


# ============================================================
# 유전자형/연관 모델
# - 3개의 유전자좌(locus 0,1,2) = (대립유전자쌍) 3개
# - X연관 locus 개수: 0~2 랜덤
# - 상염색체 locus는 "연관 그룹"을 만들 수 있음(2개가 같이 이동)
# ============================================================
@dataclass
class Person:
    sex: str  # "male" or "female"

    # X 관련: male=1 hap, female=2 hap
    x_haps: List[Dict[int, str]]  # hap -> {locus_index: "A"/"a"}

    # autosome 관련: 두 상동염색체(2 hap)
    auto_haps: List[Dict[int, str]]  # two hap: each {locus_index: "A"/"a"}


@dataclass
class Cell:
    owner: int             # 0 or 1
    ploidy: str            # "2n" or "n"
    gamete: Optional[str]  # male n: "X" or "Y"; else None


def rand_allele() -> str:
    return random.choice(["A", "a"])


def alleles_to_geno(a1: str, a2: str) -> str:
    if a1 == "A" and a2 == "A":
        return "AA"
    if a1 == "a" and a2 == "a":
        return "aa"
    return "Aa"


def diploid_presence(geno: str) -> Tuple[str, str]:
    # AA -> (O,X), Aa -> (O,O), aa -> (X,O)
    if geno == "AA":
        return ("O", "X")
    if geno == "Aa":
        return ("O", "O")
    return ("X", "O")


def haploid_presence_from_diploid(geno: str) -> Tuple[str, str]:
    # n세포에서 같은 좌의 두 대립유전자가 동시에 O 금지
    if geno == "AA":
        return ("O", "X")
    if geno == "aa":
        return ("X", "O")
    return random.choice([("O", "X"), ("X", "O")])


def make_linkage_plan() -> Tuple[List[int], List[List[int]]]:
    """
    returns:
      x_loci: list of locus indices on X (size 0~2)
      auto_groups: autosome linkage groups among remaining loci
         - each group is list of locus indices that move together in n cells
    """
    loci = [0, 1, 2]
    x_count = random.choice([0, 1, 2])
    x_loci = sorted(random.sample(loci, x_count))

    autos = [li for li in loci if li not in x_loci]

    auto_groups: List[List[int]] = []
    if len(autos) <= 1:
        auto_groups = [autos] if autos else []
    elif len(autos) == 2:
        # 두 개가 상염색체: 연관 or 비연관(서로 다른 염색체)
        if random.random() < 0.5:
            auto_groups = [autos]           # 연관(같이 이동)
        else:
            auto_groups = [[autos[0]], [autos[1]]]  # 비연관
    else:
        # autos=3인 경우(=x_count=0): (2개 연관+1개 단독) 또는 전부 비연관
        if random.random() < 0.5:
            pair = random.sample(autos, 2)
            solo = [li for li in autos if li not in pair][0]
            auto_groups = [sorted(pair), [solo]]
        else:
            auto_groups = [[autos[0]], [autos[1]], [autos[2]]]

    # 빈 그룹 제거
    auto_groups = [g for g in auto_groups if g]
    return x_loci, auto_groups


def make_person(sex: str, x_loci: List[int], auto_loci: List[int]) -> Person:
    # X haplotypes
    if sex == "male":
        x_haps = [{li: rand_allele() for li in x_loci}]
    else:
        x_haps = [
            {li: rand_allele() for li in x_loci},
            {li: rand_allele() for li in x_loci},
        ]
        # 둘이 완전 동일하면 약간 재미 없어서 한 번만 비틀기 시도
        for _ in range(20):
            if x_haps[0] != x_haps[1]:
                break
            x_haps[1] = {li: rand_allele() for li in x_loci}

    # autosome haplotypes (2 homologs)
    auto_haps = [
        {li: rand_allele() for li in auto_loci},
        {li: rand_allele() for li in auto_loci},
    ]
    return Person(sex=sex, x_haps=x_haps, auto_haps=auto_haps)


def build_cell_gene_map(person: Person,
                        cell: Cell,
                        allele_pairs: List[Tuple[str, str]],
                        x_loci: List[int],
                        auto_groups: List[List[int]]) -> Dict[str, str]:
    """
    returns presence map for one cell: {gene_symbol: "O"/"X"}
    linkage rule:
      - X에 2개 이상이면 n세포에서 같은 X hap 선택(함께 이동) → 한 번만 pick
      - autosome은 group 단위로 한 homolog 선택(연관 loci 함께 이동)
    """
    presence = {g: "X" for g in GENE_ROWS}

    # n세포에서 X hap pick(여성만 의미 있음), X에 여러 locus면 같이 움직이게 "한 번만" pick
    if cell.ploidy == "n" and person.sex == "female" and len(x_loci) >= 1:
        x_pick = random.choice([0, 1])
    else:
        x_pick = None

    # n세포에서 autosome group pick(연관 반영)
    auto_pick: Dict[int, int] = {}  # locus -> homolog index 0/1
    if cell.ploidy == "n":
        for group in auto_groups:
            pick = random.choice([0, 1])
            for li in group:
                auto_pick[li] = pick

    for li, (g1, g2) in enumerate(allele_pairs):
        is_x = (li in x_loci)

        if is_x:
            # X locus 처리
            if cell.ploidy == "2n":
                if person.sex == "male":
                    # 남성 2n: X는 1개(hemizygous)로 가정 → 한쪽만 O
                    allele = person.x_haps[0][li]
                    pres = ("O", "X") if allele == "A" else ("X", "O")
                else:
                    a1 = person.x_haps[0][li]
                    a2 = person.x_haps[1][li]
                    geno = alleles_to_geno(a1, a2)
                    pres = diploid_presence(geno)
                presence[g1], presence[g2] = pres

            else:
                # n
                if person.sex == "male":
                    # 남성 n: Y 보유면 X 유전자 없음
                    if cell.gamete == "Y":
                        presence[g1], presence[g2] = ("X", "X")
                    else:
                        allele = person.x_haps[0][li]
                        pres = ("O", "X") if allele == "A" else ("X", "O")
                        presence[g1], presence[g2] = pres
                else:
                    # 여성 n: 선택된 X hap에서 allele 결정(연관이면 같이 pick)
                    allele = person.x_haps[x_pick][li]
                    pres = ("O", "X") if allele == "A" else ("X", "O")
                    presence[g1], presence[g2] = pres

        else:
            # autosome locus 처리
            if cell.ploidy == "2n":
                a1 = person.auto_haps[0][li]
                a2 = person.auto_haps[1][li]
                geno = alleles_to_geno(a1, a2)
                pres = diploid_presence(geno)
                presence[g1], presence[g2] = pres
            else:
                # n: linkage group 단위로 같은 homolog 선택(연관 loci는 같이 이동)
                pick = auto_pick[li]
                allele = person.auto_haps[pick][li]
                pres = ("O", "X") if allele == "A" else ("X", "O")
                presence[g1], presence[g2] = pres

    return presence


def make_problem_world():
    # 성별: 한 명 male, 한 명 female
    sexes = ["male", "female"]
    random.shuffle(sexes)

    # linkage plan
    x_loci, auto_groups = make_linkage_plan()
    auto_loci = [li for li in [0, 1, 2] if li not in x_loci]

    p0 = make_person(sexes[0], x_loci, auto_loci)
    p1 = make_person(sexes[1], x_loci, auto_loci)

    # 두 사람 유전형(=haplotypes)이 완전히 같으면 재생성
    for _ in range(120):
        if (p0.sex != p1.sex) and (p0.x_haps != p1.x_haps or p0.auto_haps != p1.auto_haps):
            break
        p1 = make_person(sexes[1], x_loci, auto_loci)

    # 각 사람: 2n 1개 + n 2개
    cells0: List[Cell] = []
    cells1: List[Cell] = []

    if p0.sex == "male":
        cells0 += [Cell(0, "2n", None), Cell(0, "n", "X"), Cell(0, "n", "Y")]
    else:
        cells0 += [Cell(0, "2n", None), Cell(0, "n", None), Cell(0, "n", None)]

    if p1.sex == "male":
        cells1 += [Cell(1, "2n", None), Cell(1, "n", "X"), Cell(1, "n", "Y")]
    else:
        cells1 += [Cell(1, "2n", None), Cell(1, "n", None), Cell(1, "n", None)]

    # ✅ 사람 내부만 셔플 (ABC는 1번, DEF는 2번 유지)
    random.shuffle(cells0)
    random.shuffle(cells1)
    cells = cells0 + cells1

    # 유전자 6개를 3쌍(대립유전자쌍)으로 랜덤 페어링
    genes = GENE_ROWS[:]
    random.shuffle(genes)
    allele_pairs = [(genes[0], genes[1]), (genes[2], genes[3]), (genes[4], genes[5])]

    # 완성표
    full_table: Dict[str, Dict[str, str]] = {}
    for cell_col, cell in zip(CELL_COLS, cells):
        person = p0 if cell.owner == 0 else p1
        full_table[cell_col] = build_cell_gene_map(person, cell, allele_pairs, x_loci, auto_groups)

    # 각 사람의 n세포 2개가 서로 달라야 함(완전 동일이면 재생성)
    def owner_n_cols(owner: int, col_names: List[str]) -> List[Dict[str, str]]:
        out = []
        for col, cobj in zip(col_names, cells):
            if cobj.owner == owner and cobj.ploidy == "n":
                out.append(full_table[col])
        return out

    for owner in [0, 1]:
        ncols = owner_n_cols(owner, CELL_COLS)
        if len(ncols) == 2 and ncols[0] == ncols[1]:
            return None

    x_genes = set()
    for li in x_loci:
        x_genes.add(allele_pairs[li][0])
        x_genes.add(allele_pairs[li][1])

    auto_linked = any(len(g) >= 2 for g in auto_groups)

    meta = {
        "persons": [p0, p1],
        "cells": cells,
        "allele_pairs": allele_pairs,
        "x_loci": x_loci,
        "auto_groups": auto_groups,
        "x_genes": x_genes,
        "auto_linked": auto_linked,
    }
    return full_table, meta


# ============================================================
# Masking
# ============================================================
def mask_table(full_table: Dict[str, Dict[str, str]], total_mask: int) -> Dict[str, Dict[str, str]]:
    total_mask = max(0, min(MAX_TOTAL_MASK, total_mask))
    masked = {c: dict(full_table[c]) for c in CELL_COLS}

    gene_mask_count = {g: 0 for g in GENE_ROWS}
    cell_mask_count = {c: 0 for c in CELL_COLS}

    candidates = [(c, g) for c in CELL_COLS for g in GENE_ROWS]
    random.shuffle(candidates)

    done = 0
    for c, g in candidates:
        if done >= total_mask:
            break
        if gene_mask_count[g] >= MAX_MASK_PER_GENE:
            continue

        known_now = len(GENE_ROWS) - cell_mask_count[c]
        if known_now <= MIN_KNOWN_PER_CELL_COL:
            continue

        masked[c][g] = "?"
        gene_mask_count[g] += 1
        cell_mask_count[c] += 1
        done += 1

    return masked


# ============================================================
# 지문(연관 정보) 생성
# ============================================================
def linkage_info_text(meta) -> str:
    x_cnt = len(meta["x_loci"])
    auto_groups = meta["auto_groups"]

    # X 설명
    if x_cnt == 0:
        x_txt = "성염색체(X)에 위치한 유전자좌는 0개이다."
    elif x_cnt == 1:
        x_txt = "성염색체(X)에 위치한 유전자좌는 1개이다."
    else:
        x_txt = "성염색체(X)에 위치한 유전자좌는 2개이며, 두 유전자좌는 연관되어 함께 이동한다."

    # autosome 설명
    auto_linked = any(len(g) >= 2 for g in auto_groups)
    if not auto_groups:
        a_txt = "상염색체에 위치한 유전자좌는 없다."
    elif not auto_linked:
        a_txt = "상염색체에 위치한 유전자좌들은 서로 다른 상염색체에 위치하며 연관되지 않는다."
    else:
        # 어떤 그룹이 연관인지 말은 해 주되, '몇 개가' 정도로만
        linked_sizes = [len(g) for g in auto_groups if len(g) >= 2]
        k = max(linked_sizes) if linked_sizes else 0
        a_txt = f"상염색체에 위치한 유전자좌 중 {k}개는 같은 상염색체에 있어 연관되어 함께 이동한다."

    return x_txt + " " + a_txt


# ============================================================
# 진술 풀(pool) 확장 (x_count=0/1/2에서도 정답타겟팅 안정)
# ============================================================
def make_statement_pool(full_table, meta):
    allele_pairs = meta["allele_pairs"]
    x_genes = meta["x_genes"]
    auto_linked = meta["auto_linked"]

    pool = []

    # (1) 핵상 2n 여부(임의 세포)
    cell = random.choice(CELL_COLS)
    pl = infer_ploidy(full_table[cell], allele_pairs)
    truth = "O" if pl == "2n" else "X"
    pool.append(("PLOIDY", f"{cell} 세포의 핵상은 2n이다.", truth, {"cell": cell}))

    # (2) 특정 유전자가 X에 있다
    g = random.choice(GENE_ROWS)
    truth2 = "O" if g in x_genes else "X"
    pool.append(("XLOC", f"({g}) 유전자는 성염색체 위에 있다.", truth2, {"gene": g}))

    # (3) X연관 유전자좌가 0개이다
    truth3 = "O" if len(meta["x_loci"]) == 0 else "X"
    pool.append(("XCOUNT0", "성염색체에 위치한 유전자좌는 없다.", truth3, {}))

    # (4) X연관 유전자좌가 2개이며 연관이다
    truth4 = "O" if len(meta["x_loci"]) == 2 else "X"
    pool.append(("XCOUNT2", "성염색체에 위치한 유전자좌는 2개이며 서로 연관되어 함께 이동한다.", truth4, {}))

    # (5) 상염색체에서 연관이 존재한다
    truth5 = "O" if auto_linked else "X"
    pool.append(("ALINK", "상염색체에 위치한 유전자좌 중 일부는 연관되어 함께 이동한다.", truth5, {}))

    # (6) 대립유전자 여부(임의 두 유전자)
    g1, g2 = random.sample(GENE_ROWS, 2)
    is_allelic = any((g1, g2) == p or (g2, g1) == p for p in allele_pairs)
    truth6 = "O" if is_allelic else "X"
    pool.append(("ALLELE", f"({g1})와 ({g2})는 대립유전자이다.", truth6, {"g1": g1, "g2": g2}))

    return pool


def assign_labels_to_hit_answer(pool, answer_num: int):
    target_true = CHOICE_TO_TRUESET[answer_num]
    o_items = [s for s in pool if s[2] == "O"]
    x_items = [s for s in pool if s[2] == "X"]

    need_o = len(target_true)
    if len(o_items) < need_o or len(x_items) < (3 - need_o):
        return None

    random.shuffle(o_items)
    random.shuffle(x_items)

    labeled = []
    for lab in ["ㄱ", "ㄴ", "ㄷ"]:
        item = o_items.pop() if lab in target_true else x_items.pop()
        _id, text, truth, info = item
        labeled.append((lab, _id, text, truth, info))
    return labeled


def ox_string(stmts_labeled) -> str:
    d = {lab: truth for (lab, _id, _text, truth, _info) in stmts_labeled}
    return d["ㄱ"] + d["ㄴ"] + d["ㄷ"]


def build_answer_list() -> List[int]:
    if sum(ANSWER_COUNTS.values()) != N_PROBLEMS:
        raise ValueError("ANSWER_COUNTS 합이 30이어야 함")
    out = []
    for k, v in ANSWER_COUNTS.items():
        out.extend([k] * v)
    random.shuffle(out)
    return out


# ============================================================
# 해설
# ============================================================
def explain_statement(stmt_id: str, truth: str, info: dict, full_table, meta) -> str:
    allele_pairs = meta["allele_pairs"]
    x_genes = meta["x_genes"]

    verdict = "옳다(O)" if truth == "O" else "옳지 않다(X)"

    if stmt_id == "PLOIDY":
        cell = info["cell"]
        reason = describe_ploidy_reason(full_table[cell], allele_pairs)
        return f"{verdict} — {cell}열에서 {reason}"

    if stmt_id == "XLOC":
        gene = info["gene"]
        if gene in x_genes:
            return f"{verdict} — ({gene})는 X염색체에 위치한 유전자좌에 속한다."
        return f"{verdict} — ({gene})는 X염색체에 위치한 유전자좌에 속하지 않는다."

    if stmt_id == "XCOUNT0":
        return f"{verdict} — X연관 유전자좌 개수는 {len(meta['x_loci'])}개이다."

    if stmt_id == "XCOUNT2":
        return f"{verdict} — X연관 유전자좌 개수는 {len(meta['x_loci'])}개이며, 2개일 때 연관 이동으로 처리한다."

    if stmt_id == "ALINK":
        return f"{verdict} — 상염색체 연관 여부는 문제 조건(지문)에 따라 결정된다."

    if stmt_id == "ALLELE":
        g1, g2 = info["g1"], info["g2"]
        pairs_txt = ", ".join([f"{a}/{b}" for a, b in allele_pairs])
        is_allelic = any((g1, g2) == p or (g2, g1) == p for p in allele_pairs)
        if is_allelic:
            return f"{verdict} — 대립유전자쌍({pairs_txt}) 중 ({g1})/({g2})가 같은 쌍이다."
        return f"{verdict} — 대립유전자쌍은 {pairs_txt}이며, ({g1})와 ({g2})는 같은 쌍이 아니다."

    return verdict


# ============================================================
# 문제 1개 생성
# ============================================================
def generate_one_problem(answer_num: int):
    mask_now = MAX_TOTAL_MASK

    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        if attempt == RELAX_AFTER_1:
            mask_now = max(8, mask_now - 2)
        if attempt == RELAX_AFTER_2:
            mask_now = max(6, mask_now - 2)

        world = make_problem_world()
        if world is None:
            continue

        full_table, meta = world

        pool = make_statement_pool(full_table, meta)
        stmts_labeled = assign_labels_to_hit_answer(pool, answer_num)
        if stmts_labeled is None:
            continue

        masked = mask_table(full_table, mask_now)

        explanations = []
        for lab, stmt_id, text, truth, info in stmts_labeled:
            explanations.append((lab, text, truth, explain_statement(stmt_id, truth, info, full_table, meta)))

        return {
            "full_table": full_table,
            "masked_table": masked,
            "meta": meta,
            "stmts": stmts_labeled,
            "ox": ox_string(stmts_labeled),
            "answer_num": answer_num,
            "mask_used": mask_now,
            "explanations": explanations,
        }

    raise RuntimeError(f"정답번호({answer_num}) 생성 실패")


# ============================================================
# DOCX 출력
# ============================================================
def add_ox_table(container, table):
    t = container.add_table(rows=1 + len(GENE_ROWS), cols=1 + len(CELL_COLS))
    t.style = "Table Grid"

    t.rows[0].cells[0].text = ""
    for j, cell in enumerate(CELL_COLS):
        t.rows[0].cells[j + 1].text = cell

    for i, gene in enumerate(GENE_ROWS):
        t.rows[i + 1].cells[0].text = gene
        for j, cell in enumerate(CELL_COLS):
            t.rows[i + 1].cells[j + 1].text = table[cell][gene]


def render_choices_line() -> str:
    return "① ㄱ    ② ㄴ    ③ ㄱ, ㄴ    ④ ㄱ, ㄷ    ⑤ ㄴ, ㄷ"


def add_problem_to_cell(doc_cell, pnum: int, pdata):
    h = doc_cell.add_paragraph(f"[문제 {pnum}]")
    h.runs[0].bold = True

    doc_cell.add_paragraph("다음 표는 감수분열 중인 세포에서 유전자의 존재 여부(O/X)를 나타낸 것이다.")
    doc_cell.add_paragraph("※ A, B, C는 1번 사람의 세포이고 D, E, F는 2번 사람의 세포이다.")
    doc_cell.add_paragraph("※ " + linkage_info_text(pdata["meta"]))
    doc_cell.add_paragraph("")
    add_ox_table(doc_cell, pdata["masked_table"])

    doc_cell.add_paragraph("")
    doc_cell.add_paragraph("다음 진술의 O/X를 판단하여, 옳은 것만을 고른 번호를 선택하시오.")
    for lab, stmt_id, text, truth, info in pdata["stmts"]:
        doc_cell.add_paragraph(f"{lab}. {text}   (O / X)")
    doc_cell.add_paragraph("")
    doc_cell.add_paragraph(render_choices_line())



def gene_table_to_md(table: Dict[str, Dict[str, str]]) -> str:
    # table[cell][gene]
    header = [""] + CELL_COLS
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for g in GENE_ROWS:
        row = [g] + [str(table[c][g]) for c in CELL_COLS]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

def make_docx():
    # output 폴더 생성
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="GENEDETX", out_dir=out_dir, id_prefix="GENEDETX_")

    os.makedirs(out_dir, exist_ok=True)
    # 저장 파일명 (시간 붙여서 중복 방지)
    filename = os.path.join(out_dir, f"Gene_Detecting_{int(time.time())}.docx")
    doc = Document()
    style = doc.styles["Normal"]
    set_page_margins(doc)
    style.font.name = "바탕"
    style.font.size = Pt(9)

    answer_nums = build_answer_list()

    problems = []
    for i in range(N_PROBLEMS):
        pdata = generate_one_problem(answer_nums[i])
        problems.append(pdata)
        # PACK record
        # PACK record (표시용 텍스트/표/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "다음 표는 감수분열 중인 세포에서 유전자의 존재 여부(O/X)를 나타낸 것이다.",
            "가로(열)=세포 A~F, 세로(행)=유전자 가~바",
        ])
        ask_line_md = "다음 진술의 O/X를 판단하여, 옳은 것만을 고른 번호를 선택하시오."

        table_md = gene_table_to_md(pdata["masked_table"])
        full_table_md = gene_table_to_md(pdata["full_table"])

        answer_md = f"정답 {pdata['answer_num']} (ㄱㄴㄷ={pdata['ox']})"
        explanation_md = "\n".join([f"- {lab}: {exp}" for (lab, _t, _ox, exp) in pdata.get("explanations", [])])

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": table_md,
            "full_table_md": full_table_md,
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "mask_used": pdata.get("mask_used"),
            "data": pdata,
        }))

    # 문제(2단, 2문제/페이지)
    idx = 0
    pnum = 1
    while idx < N_PROBLEMS:
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        add_problem_to_cell(left, pnum, problems[idx])
        idx += 1
        pnum += 1

        if idx < N_PROBLEMS:
            add_problem_to_cell(right, pnum, problems[idx])
            idx += 1
            pnum += 1

        if idx < N_PROBLEMS:
            doc.add_page_break()

    # 정답
    doc.add_page_break()
    h = doc.add_paragraph("[정답]")
    h.runs[0].bold = True
    for i, pdata in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번: 정답 {pdata['answer_num']}   (ㄱㄴㄷ={pdata['ox']}, ?={pdata['mask_used']})")

    # 해설 + 완성표
    doc.add_page_break()
    h2 = doc.add_paragraph("[해설(완성표 포함)]")
    h2.runs[0].bold = True

    for i, pdata in enumerate(problems, start=1):
        title = doc.add_paragraph(f"{i}번 해설")
        title.runs[0].bold = True

        doc.add_paragraph("조건(연관 정보): " + linkage_info_text(pdata["meta"]))
        doc.add_paragraph("")

        doc.add_paragraph("① 완성표(원본, ? 없음)")
        add_ox_table(doc, pdata["full_table"])
        doc.add_paragraph("")

        doc.add_paragraph("② 제시표(물음표 포함)")
        add_ox_table(doc, pdata["masked_table"])
        doc.add_paragraph("")

        doc.add_paragraph("③ 진술 판단")
        for lab, text, truth, exp in pdata["explanations"]:
            doc.add_paragraph(f"{lab}. {text}")
            doc.add_paragraph(f"   → {exp}")

        doc.add_paragraph("")

    doc.save(filename)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {filename}")


if __name__ == "__main__":
    make_docx()