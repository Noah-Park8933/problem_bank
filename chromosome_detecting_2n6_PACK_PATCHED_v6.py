# chromosome_detecting_2n6.py
# ------------------------------------------------------------
# Chromosome Detecting (2n=6, 상염색체만) 자동 문제 생성기
# 표 방향(선택 2):
#   - 세로(행) = 세포 (A, B, C, D, ...)
#   - 가로(열) = 염색체 (가, 나, 다, 라, 마, 바)
#
# 규칙 반영(사용자 정의):
# 1-1: 한 세포에서 O가 50%보다 많으면(>3) 핵상 2n이고 모든 염색체가 O
# 1-2: 한 세포에서 X가 1개라도 있으면 핵상 n이고 전체 염색체의 50%가 O (즉 O=3, X=3)
# 2-1: n세포에는 1~n번 염색체가 1개씩(=각 상동쌍에서 하나씩) 들어있다
# 2-2: n세포에서 O인 염색체는 서로 상동염색체가 아니다 (즉 같은 쌍에서 둘 다 O 불가)
#
# 출력:
# - 문제 파트: 물음표 포함 제시표
# - 해설 파트: 상동염색체 쌍 + 완성표(물음표 없는 원본)
#
# 실행:
#   pip install python-docx
#   python chromosome_detecting_2n6.py
# ------------------------------------------------------------
import time
import os
import random
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, Cm
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30

# 표 구성
CELLS = ["A", "B", "C", "D"]              # 세포(행) 수: 필요하면 늘려도 됨(예: ["A","B","C","D","E"])
CHRS  = ["가", "나", "다", "라", "마", "바"]  # 염색체(열) = 2n=6

# 물음표(마스킹) 설정
MASK_TOTAL_RANGE = (8, 12)   # 한 문제에서 ? 개수 범위(원하면 조절)
MAX_MASK_PER_COL = 2         # 같은 염색체(열)에서 최대 ? 개수 (너가 예전에 쓰던 제한 감각 유지)
MIN_KNOWN_PER_ROW = 2        # 한 세포(행)에서 최소 공개 칸

# =========================
# World 생성(완성표)
# =========================
def random_homolog_pairs(chrs: List[str]) -> List[Tuple[str, str]]:
    xs = chrs[:]
    random.shuffle(xs)
    return [(xs[0], xs[1]), (xs[2], xs[3]), (xs[4], xs[5])]  # 3쌍

def make_full_table() -> Tuple[Dict[str, Dict[str, str]], List[Tuple[str, str]]]:
    """
    반환:
      full[cell][chr] = 'O'/'X'
      pairs = [(..., ...), ...]  # 상동쌍 3개
    구성:
      - 2n 세포 1개: 전부 O
      - n 세포 나머지: 각 상동쌍에서 하나씩만 O → O=3, X=3 자동 보장
    """
    pairs = random_homolog_pairs(CHRS)

    # 어떤 세포가 2n인지 선택
    two_n_cell = random.choice(CELLS)

    full: Dict[str, Dict[str, str]] = {}
    for cell in CELLS:
        full[cell] = {c: "X" for c in CHRS}

        if cell == two_n_cell:
            # 2n: 전부 O
            for c in CHRS:
                full[cell][c] = "O"
        else:
            # n: 각 상동쌍에서 1개씩만 O
            for a, b in pairs:
                pick = random.choice([a, b])
                full[cell][pick] = "O"
                # 나머지는 X 그대로
            # 이 시점에 O=3, X=3 확정

    return full, pairs

# =========================
# 마스킹(제시표)
# =========================
def mask_table(full: Dict[str, Dict[str, str]]) -> Dict[str, Dict[str, str]]:
    masked = {r: dict(full[r]) for r in CELLS}

    total_mask = random.randint(MASK_TOTAL_RANGE[0], MASK_TOTAL_RANGE[1])
    col_mask = {c: 0 for c in CHRS}
    row_mask = {r: 0 for r in CELLS}

    candidates = [(r, c) for r in CELLS for c in CHRS]
    random.shuffle(candidates)

    done = 0
    for r, c in candidates:
        if done >= total_mask:
            break
        if col_mask[c] >= MAX_MASK_PER_COL:
            continue

        known_now = len(CHRS) - row_mask[r]
        if known_now <= MIN_KNOWN_PER_ROW:
            continue

        masked[r][c] = "?"
        col_mask[c] += 1
        row_mask[r] += 1
        done += 1

    return masked

# =========================
# DOCX 출력 유틸
# =========================
def add_table(doc_or_cell, table: Dict[str, Dict[str, str]]):
    """
    표 방향(선택2):
      - 행(세로) = 세포 A,B,C,D
      - 열(가로) = 염색체 가,나,다,라,마,바
    """
    t = doc_or_cell.add_table(rows=1 + len(CELLS), cols=1 + len(CHRS))
    t.style = "Table Grid"

    # header row
    t.rows[0].cells[0].text = ""
    for j, chr_ in enumerate(CHRS):
        t.rows[0].cells[j + 1].text = chr_

    # body
    for i, cell in enumerate(CELLS):
        t.rows[i + 1].cells[0].text = cell
        for j, chr_ in enumerate(CHRS):
            t.rows[i + 1].cells[j + 1].text = table[cell][chr_]

def pairs_to_text(pairs: List[Tuple[str, str]]) -> str:
    # 보기 좋게 정렬해서 출력(쌍 내부/쌍들 순서)
    norm = [tuple(sorted(p)) for p in pairs]
    norm.sort()
    return ", ".join([f"({a}, {b})" for a, b in norm])

# =========================
# 문제 1개 생성
# =========================
def generate_problem():
    full, pairs = make_full_table()
    masked = mask_table(full)
    return {
        "full": full,
        "masked": masked,
        "pairs": pairs,
    }


def grid_to_md(grid: Dict[str, Dict[str, str]], row_labels: List[str], col_labels: List[str]) -> str:
    header = [""] + [str(c) for c in col_labels]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for r in row_labels:
        row = [str(r)] + [str(grid[r][c]) for c in col_labels]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

def count_masks_grid(grid: Dict[str, Dict[str, str]]) -> int:
    return sum(1 for r in grid for c in grid[r] if grid[r][c] == "?")

# =========================
# DOCX 생성(30문제)
# =========================
def make_docx():
    # output 폴더 생성
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="CHROM26", out_dir=out_dir, id_prefix="CHROM26_")

    os.makedirs(out_dir, exist_ok=True)
    # 저장 파일명 (시간 붙여서 중복 방지)
    filename = os.path.join(out_dir, f"Chromosome_Detecting_{int(time.time())}.docx")
    doc = Document()
    style = doc.styles["Normal"]
    set_page_margins(doc)
    style.font.name = "바탕"
    style.font.size = Pt(9)

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        p = generate_problem()
        problems.append(p)

        problem_text_md = "\n".join([
            "다음 표는 감수분열 중인 임의의 세포에서 염색체의 존재 여부(O/X)를 나타낸 것이다.",
            "단, 2n = 6이다.",
        ])
        ask_line_md = "물음표(?)를 채우고, 각 세포의 핵상(2n/n)과 상동염색체 쌍을 구하시오."

        table_md = grid_to_md(p["masked"], CELLS, CHRS)
        full_table_md = grid_to_md(p["full"], CELLS, CHRS)

        answer_md = f"상동염색체 쌍: {pairs_to_text(p['pairs'])}"
        explanation_md = answer_md

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": table_md,
            "full_table_md": full_table_md,
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "mask_used": count_masks_grid(p["masked"]),
            "data": p,
        }))

    # 문제 파트
    title = doc.add_paragraph("[문제]")
    title.runs[0].bold = True
    doc.add_paragraph("다음 표는 감수분열 중인 임의의 세포에서 염색체의 존재 여부(O/X)를 나타낸 것이다.")
    doc.add_paragraph("단, 2n = 6이다.")
    doc.add_paragraph("물음표(?)를 채우고, 각 세포의 핵상(2n/n)과 상동염색체 쌍을 구하시오.")
    doc.add_paragraph("")

    for i, p in enumerate(problems, start=1):
        h = doc.add_paragraph(f"[{i}번]")
        h.runs[0].bold = True
        add_table(doc, p["masked"])
        doc.add_paragraph("")

    # 해설 파트(요구: 상동쌍 + 원래표만)
    doc.add_page_break()
    title2 = doc.add_paragraph("[해설]")
    title2.runs[0].bold = True
    doc.add_paragraph("※ 각 문항의 해설에는 상동염색체 쌍과 완성표(원본)만 제시한다.")
    doc.add_paragraph("")

    for i, p in enumerate(problems, start=1):
        h = doc.add_paragraph(f"[{i}번 해설]")
        h.runs[0].bold = True
        doc.add_paragraph(f"상동염색체 쌍: {pairs_to_text(p['pairs'])}")
        doc.add_paragraph("완성표(원본):")
        add_table(doc, p["full"])
        doc.add_paragraph("")

    doc.save(filename)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {filename}")

if __name__ == "__main__":
    make_docx()