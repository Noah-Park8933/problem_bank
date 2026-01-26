# deduction_1_2_30gen_unique.py
# ------------------------------------------------------------
# ✅ Deduction 1-2 자동 출제기 (표현형→유전자형 버전)
# - 4유전자 A,B,D,E / 2개는 7번(연관), 2개는 9번(연관)
# - 문제에 "(A/a)는 7번 염색체에 있다"를 추가해 7/9 뒤집힘 모호성 제거
# - 조건:
#   (1) P1: 주어진 생식세포 "AbDE" 생성 가능
#   (2) 자손의 "유전자형"이 P2와 같을 확률 = 0
# - 솔버: 가능한 모든 연관 분할(AB/DE, AD/BE, AE/BD) + (각 염색체) 결합/반발 위상 전부 탐색
#         조건 만족 해가 "유일"한 문제만 통과
# - 출력: DOCX(2단, 2문제/페이지) + 정답/해설
#
# 설치/실행:
#   pip install python-docx
#   python deduction_1_2_30gen_unique.py
# ------------------------------------------------------------

import os, time, random
from dataclasses import dataclass
from typing import Dict, Tuple, List, Optional

from docx import Document
from docx.shared import Pt
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
MAX_TRIES_PER_PROBLEM = 200000000
PROGRESS_EVERY = 2000

OUT_DIR = "output"
OUT_FILE_PREFIX = "Deduction_1_2_unique"

FONT_NAME = "바탕"
FONT_SIZE = 9

LOCI = ["A", "B", "D", "E"]

# 완전연관(재조합 0) 가정: Deduction류에서 가장 안정적으로 유일정답 만들기 쉬움
R = 0.0

# 선호: 너무 단순한 문제(거의 고정) 제외
MIN_HET_SUM = 2  # P1+P2 헤테로 수 합 최소

# =========================
# Helpers
# =========================
def norm_pair(x: str, y: str) -> str:
    # 'a','A' -> 'Aa'
    if x.isupper() and y.islower():
        return x + y
    if x.islower() and y.isupper():
        return y + x
    return x + y

def alleles_of(geno: str) -> List[str]:
    # "Aa" -> ["A","a"], "AA" -> ["A"]
    if geno[0] == geno[1]:
        return [geno[0]]
    return [geno[0], geno[1]]

def geno_random(locus: str) -> str:
    U = locus
    L = locus.lower()
    return random.choice([U+U, U+L, L+L])

def hetero_count(G: Dict[str,str]) -> int:
    return sum(1 for L in LOCI if G[L][0] != G[L][1])

def pheno_label(geno: Dict[str,str]) -> str:
    # "A- B- D- E-" 형태
    def dom(locus: str) -> str:
        g = geno[locus]
        if g[0].isupper() or g[1].isupper():
            return f"{locus}-"
        return f"{locus.lower()}{locus.lower()}"
    return f"{dom('A')} {dom('B')} {dom('D')} {dom('E')}"

def child_from_gametes(g1: Dict[str,str], g2: Dict[str,str]) -> Dict[str,str]:
    return {L: norm_pair(g1[L], g2[L]) for L in LOCI}

# =========================
# Linkage gamete model (R=0)
# phase: "coupling" or "repulsion"
# for linked pair (X,Y):
#  - coupling: XY and xy only (if both hetero)
#  - repulsion: Xy and xY only (if both hetero)
#  - if any locus homozygous: phase irrelevant (collapse)
# =========================
def linked_gametes(pair: Tuple[str,str], G: Dict[str,str], phase: str) -> Dict[Tuple[str,str], float]:
    X, Y = pair
    ax = alleles_of(G[X])
    ay = alleles_of(G[Y])

    if len(ax) == 1 and len(ay) == 1:
        return {(ax[0], ay[0]): 1.0}
    if len(ax) == 1 and len(ay) == 2:
        return {(ax[0], ay[0]): 0.5, (ax[0], ay[1]): 0.5}
    if len(ax) == 2 and len(ay) == 1:
        return {(ax[0], ay[0]): 0.5, (ax[1], ay[0]): 0.5}

    # both hetero
    Ux, lx = X, X.lower()
    Uy, ly = Y, Y.lower()
    if phase == "coupling":
        return {(Ux, Uy): 0.5, (lx, ly): 0.5}
    else:
        return {(Ux, ly): 0.5, (lx, Uy): 0.5}

def indep_gametes(locus: str, G: Dict[str,str]) -> Dict[str,float]:
    g = G[locus]
    if g[0] == g[1]:
        return {g[0]: 1.0}
    return {g[0]: 0.5, g[1]: 0.5}

# =========================
# Hypothesis space:
# Partition into 2 linked pairs:
#  1) (A,B) & (D,E)
#  2) (A,D) & (B,E)
#  3) (A,E) & (B,D)
#
# Additionally, we FIX: "(A/a)는 7번 염색체"
# → the linked pair containing A is on chr7, the other pair on chr9 (unique)
# We still search phases for:
#   - P1 chr7-pair phase
#   - P1 chr9-pair phase
#   - P2 chr7-pair phase
#   - P2 chr9-pair phase
# (If parent not hetero at both loci in the pair, phase doesn't matter but we still allow both)
# =========================
PARTITIONS = [
    (("A","B"), ("D","E")),
    (("A","D"), ("B","E")),
    (("A","E"), ("B","D")),
]
PHASES = ["coupling", "repulsion"]

@dataclass
class Solution:
    chr7: Tuple[str,str]  # contains A
    chr9: Tuple[str,str]
    p1_phase_7: str
    p1_phase_9: str
    p2_phase_7: str
    p2_phase_9: str
    p_pheno_P1: float
    p_pheno_P2: float

# =========================
# Condition checks
# =========================
def gamete_possible_P1_AbDE(P1: Dict[str,str], chr7: Tuple[str,str], chr9: Tuple[str,str],
                           phase7: str, phase9: str) -> bool:
    # need gamete with A, b, D, E (AbDE)
    g7 = linked_gametes(chr7, P1, phase7)
    g9 = linked_gametes(chr9, P1, phase9)

    for (x7,y7), p7 in g7.items():
        if p7 == 0: continue
        for (x9,y9), p9 in g9.items():
            if p9 == 0: continue
            gam = {chr7[0]: x7, chr7[1]: y7, chr9[0]: x9, chr9[1]: y9}
            # normalize dictionary for all loci
            # chr7/chr9 cover all four loci
            if gam["A"] == "A" and gam["B"] == "b" and gam["D"] == "D" and gam["E"] == "E":
                return True
    return False

def prob_child_genotype_equals_P2(P1: Dict[str,str], P2: Dict[str,str],
                                 chr7: Tuple[str,str], chr9: Tuple[str,str],
                                 p1ph7: str, p1ph9: str, p2ph7: str, p2ph9: str) -> float:
    # iterate gametes
    g1_7 = linked_gametes(chr7, P1, p1ph7)
    g1_9 = linked_gametes(chr9, P1, p1ph9)
    g2_7 = linked_gametes(chr7, P2, p2ph7)
    g2_9 = linked_gametes(chr9, P2, p2ph9)

    target = P2  # exact genotype equality

    prob = 0.0
    for (a7,b7), p7 in g1_7.items():
        for (a9,b9), p9 in g1_9.items():
            pr1 = p7*p9
            if pr1 == 0: continue
            gam1 = {chr7[0]: a7, chr7[1]: b7, chr9[0]: a9, chr9[1]: b9}
            for (c7,d7), q7 in g2_7.items():
                for (c9,d9), q9 in g2_9.items():
                    pr2 = q7*q9
                    if pr2 == 0: continue
                    gam2 = {chr7[0]: c7, chr7[1]: d7, chr9[0]: c9, chr9[1]: d9}
                    child = child_from_gametes(gam1, gam2)
                    if child == target:
                        prob += pr1*pr2
    return prob

def prob_child_phenotypes(P1: Dict[str,str], P2: Dict[str,str],
                          chr7: Tuple[str,str], chr9: Tuple[str,str],
                          p1ph7: str, p1ph9: str, p2ph7: str, p2ph9: str) -> Tuple[float,float]:
    # prob phenotype equals P1 phenotype / equals P2 phenotype
    g1_7 = linked_gametes(chr7, P1, p1ph7)
    g1_9 = linked_gametes(chr9, P1, p1ph9)
    g2_7 = linked_gametes(chr7, P2, p2ph7)
    g2_9 = linked_gametes(chr9, P2, p2ph9)

    ph1 = pheno_label(P1)
    ph2 = pheno_label(P2)

    p_same_1 = 0.0
    p_same_2 = 0.0

    for (a7,b7), p7 in g1_7.items():
        for (a9,b9), p9 in g1_9.items():
            pr1 = p7*p9
            if pr1 == 0: continue
            gam1 = {chr7[0]: a7, chr7[1]: b7, chr9[0]: a9, chr9[1]: b9}
            for (c7,d7), q7 in g2_7.items():
                for (c9,d9), q9 in g2_9.items():
                    pr2 = q7*q9
                    if pr2 == 0: continue
                    gam2 = {chr7[0]: c7, chr7[1]: d7, chr9[0]: c9, chr9[1]: d9}
                    child = child_from_gametes(gam1, gam2)
                    ph = pheno_label(child)
                    pr = pr1*pr2
                    if ph == ph1:
                        p_same_1 += pr
                    if ph == ph2:
                        p_same_2 += pr

    return p_same_1, p_same_2

# =========================
# Solver: enumerate all hypotheses and enforce "unique"
# =========================
def solve_unique(P1: Dict[str,str], P2: Dict[str,str], eps: float = 1e-12) -> Tuple[Optional[Solution], List[Solution]]:
    sols: List[Solution] = []

    for (pairA, pairB) in PARTITIONS:
        # Determine chr7 pair (must include A)
        if "A" in pairA:
            chr7 = pairA
            chr9 = pairB
        else:
            chr7 = pairB
            chr9 = pairA

        # search phases
        for p1ph7 in PHASES:
            for p1ph9 in PHASES:
                # condition: P1 can make AbDE gamete
                if not gamete_possible_P1_AbDE(P1, chr7, chr9, p1ph7, p1ph9):
                    continue

                for p2ph7 in PHASES:
                    for p2ph9 in PHASES:
                        # condition: P(child genotype == P2) == 0
                        p_eq = prob_child_genotype_equals_P2(P1, P2, chr7, chr9, p1ph7, p1ph9, p2ph7, p2ph9)
                        if abs(p_eq - 0.0) > eps:
                            continue

                        pP1, pP2 = prob_child_phenotypes(P1, P2, chr7, chr9, p1ph7, p1ph9, p2ph7, p2ph9)
                        sols.append(Solution(chr7, chr9, p1ph7, p1ph9, p2ph7, p2ph9, pP1, pP2))

    # unique by (chr7-pair, chr9-pair) and P1/P2 phases (full uniqueness)
    if not sols:
        return None, sols

    keyset = set((tuple(sorted(s.chr7)), tuple(sorted(s.chr9)),
                  s.p1_phase_7, s.p1_phase_9, s.p2_phase_7, s.p2_phase_9) for s in sols)
    if len(keyset) != 1:
        return None, sols

    return sols[0], sols

# =========================
# Random parent generation (varied)
# =========================
def random_parents() -> Tuple[Dict[str,str], Dict[str,str]]:
    while True:
        P1 = {L: geno_random(L) for L in LOCI}
        P2 = {L: geno_random(L) for L in LOCI}

        # too easy filter
        if hetero_count(P1) + hetero_count(P2) < MIN_HET_SUM:
            continue

        # avoid identical genotype (boring)
        if all(P1[L] == P2[L] for L in LOCI):
            continue

        return P1, P2

# =========================
# Problem object
# =========================
@dataclass
class Problem:
    P1: Dict[str,str]
    P2: Dict[str,str]
    sol: Solution

# =========================
# Build one problem (unique solution)
# =========================
def make_one_problem(idx: int) -> Problem:
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        P1, P2 = random_parents()

        sol, all_s = solve_unique(P1, P2)
        if sol is not None:
            return Problem(P1, P2, sol)

        if attempt % PROGRESS_EVERY == 0:
            print(f"[P{idx:02d}] attempt={attempt} searching... (still no unique hit)")

    raise RuntimeError(f"[P{idx:02d}] 생성 실패: MAX_TRIES_PER_PROBLEM 증가 또는 필터 완화 필요")

# =========================
# Explanation (결정적 단서 스타일)
# =========================
def phase_kor(ph: str) -> str:
    return "결합형" if ph == "coupling" else "반발형"

def pair_text(pair: Tuple[str,str]) -> str:
    x,y = pair
    return f"({x}/{x.lower()})-({y}/{y.lower()})"

def build_explanation(pr: Problem) -> str:
    s = pr.sol
    # 결정 단서 1: AbDE 가능 → 두 유전자쌍 분해
    # (완전 자동으로 '반드시'라는 증명까지 쓰면 길어져서, 평가원식 핵심만)
    lines = []
    lines.append(f"AbDE 생식세포가 가능하므로 A와 함께 이동하는 7번 염색체 쌍과, 나머지 9번 염색체 쌍이 어떻게 짝지어지는지 후보(AB/DE, AD/BE, AE/BD)를 모두 점검한다.")
    lines.append(f"유일하게 성립하는 연관쌍은 7번: {pair_text(s.chr7)}, 9번: {pair_text(s.chr9)}이다.")
    lines.append(f"또한 배열은 P1에서 7번 {phase_kor(s.p1_phase_7)}, 9번 {phase_kor(s.p1_phase_9)}; P2에서 7번 {phase_kor(s.p2_phase_7)}, 9번 {phase_kor(s.p2_phase_9)}로 정해진다.")
    lines.append(f"이때 자손 유전자형이 P2와 정확히 같아지려면(P2와 동일 4좌 모두 일치) 특정 배우자 조합이 필요하지만, 위 연관·배열에서는 그 조합이 생성되지 않아 확률이 0이 된다.")
    lines.append(f"따라서 P(자손 표현형=P1)={s.p_pheno_P1:.6f}, P(자손 표현형=P2)={s.p_pheno_P2:.6f}.")
    return " ".join(lines)

# =========================
# DOCX writer
# =========================
def set_style(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = Pt(FONT_SIZE)

def add_problem_to_cell(cell, idx: int, pr: Problem):
    p = cell.add_paragraph(f"[문제 {idx}]")
    if p.runs:
        p.runs[0].bold = True

    cell.add_paragraph("형질 (가)는 A/a에 의해 결정되고 형질 (나)는 B/b에 의해, 형질 (다)는 D/d에 의해, 형질 (라)는 E/e에 의해 결정된다.")
    cell.add_paragraph("유전자의 우열관계는 A>a이고 B>b이며 D>d이고 E>e이다.")
    cell.add_paragraph("(A/a), (B/b), (D/d), (E/e) 중 2개는 7번 염색체에 있고 나머지 2개는 9번 염색체에 있다.")
    cell.add_paragraph("또한 (A/a)는 7번 염색체에 존재한다.")  # ✅ 7/9 모호성 제거

    cell.add_paragraph(f"유전자형이 {pr.P1['A']}{pr.P1['B']}{pr.P1['D']}{pr.P1['E']}인 부모 P1과 "
                       f"{pr.P2['A']}{pr.P2['B']}{pr.P2['D']}{pr.P2['E']}인 부모 P2 사이에서 태어나는 자손의 "
                       "유전자형이 P2와 같을 확률은 0이다.")
    cell.add_paragraph("P1에게서 유전자형이 AbDE인 생식세포가 생성될 수 있다.")
    cell.add_paragraph("P1과 P2에서 유전자 사이의 관계를 구하고 자손의 표현형이 P1과 같을 확률과 자손의 표현형이 P2와 같을 확률을 구하시오.")

def add_two_col_pages(doc: Document, problems: List[Problem]):
    i = 0
    pnum = 1
    while i < len(problems):
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        add_problem_to_cell(left, pnum, problems[i])
        i += 1
        pnum += 1

        if i < len(problems):
            add_problem_to_cell(right, pnum, problems[i])
            i += 1
            pnum += 1

        if i < len(problems):
            doc.add_page_break()

def add_answers(doc: Document, problems: List[Problem]):
    doc.add_page_break()
    h = doc.add_paragraph("[정답·해설]")
    if h.runs:
        h.runs[0].bold = True

    for idx, pr in enumerate(problems, start=1):
        s = pr.sol
        doc.add_paragraph(f"{idx}번 정답")
        doc.add_paragraph(f" - 7번 염색체: {pair_text(s.chr7)} / 9번 염색체: {pair_text(s.chr9)}")
        doc.add_paragraph(f" - 배열: P1(7번 {phase_kor(s.p1_phase_7)}, 9번 {phase_kor(s.p1_phase_9)}), "
                          f"P2(7번 {phase_kor(s.p2_phase_7)}, 9번 {phase_kor(s.p2_phase_9)})")
        doc.add_paragraph(f" - P(자손 표현형=P1) = {s.p_pheno_P1:.6f}")
        doc.add_paragraph(f" - P(자손 표현형=P2) = {s.p_pheno_P2:.6f}")
        doc.add_paragraph(f" - 해설: {build_explanation(pr)}")
        doc.add_paragraph("")

def make_docx():
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(out_dir, exist_ok=True)
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="DED12", out_dir=out_dir, id_prefix="DED12_")

    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, f"{OUT_FILE_PREFIX}_{int(time.time())}.docx")

    problems: List[Problem] = []
    for i in range(1, N_PROBLEMS + 1):
        pr = make_one_problem(i)
        problems.append(pr)
        # PACK record
        # PACK record (표시용 텍스트/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "형질 (가)는 A/a에 의해 결정되고 형질 (나)는 B/b에 의해, 형질 (다)는 D/d에 의해, 형질 (라)는 E/e에 의해 결정된다.",
            "유전자의 우열관계는 A>a, B>b, D>d, E>e이다.",
            "(A/a), (B/b), (D/d), (E/e) 중 2개는 7번 염색체에 있고 나머지 2개는 9번 염색체에 있다.",
            f"유전자형이 {pr.P1['A']}{pr.P1['B']}{pr.P1['D']}{pr.P1['E']}인 부모 P1과 {pr.P2['A']}{pr.P2['B']}{pr.P2['D']}{pr.P2['E']}인 부모 P2 사이에서",
            "태어나는 자손의 표현형이 P2와 같을 확률은 0이다.",
            "P1에게서 유전자형이 AbDE인 생식세포가 생성될 수 있다.",
        ])
        ask_line_md = "P1/P2에서 유전자 사이의 관계를 구하고, 자손의 표현형이 P1/P2와 같을 확률을 구하시오."
        answer_md = pr.answer_line
        explanation_md = pr.explanation

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": "",
            "full_table_md": "",
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "data": {
                "P1": pr.P1,
                "P2": pr.P2,
                "solution": pr.sol.__dict__ if hasattr(pr, 'sol') else None,
            },
        }))
        print(f"[진행] {i}/{N_PROBLEMS} 생성 완료")

    doc = Document()
    set_style(doc)

    add_two_col_pages(doc, problems)
    add_answers(doc, problems)

    doc.save(out_path)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print("✅ 저장 완료:", out_path)

if __name__ == "__main__":
    make_docx()