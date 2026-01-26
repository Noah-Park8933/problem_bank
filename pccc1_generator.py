# pccc1_generator.py
# ------------------------------------------------------------
# ✅ PCCC1: Polygenic Cheatkey + Capital Counting 1
#
# 규칙(사용자 최신):
# 1) (A/a), (B/b), (D/d) 3쌍 유전자 모두 독립
# 2) 표현형 (n) = 자손의 "대문자(우성) 대립유전자 총 개수"가 n인 경우
#    - 총 대문자 개수는 0~6 가능 → 표현형 (0)~(6)
# 3) 조건: P2가 P1보다 (d의 수 / A의 수)가 "작다" 또는 "크다"
# 4) 타깃: 자손의 표현형 (k)일 확률 = p  (k, p는 월드에서 자동으로 뽑아 다양화)
#
# 생성 방식(요청 반영: P1,P2 무작위 생성 후 끼워맞추기):
# - P1 랜덤 생성
# - P2 랜덤 생성
# - ratio_dir(< 또는 >)는 P1,P2 비교로 자동 결정(=끼워맞추기)
# - 그 P1×P2에서 나온 분포에서 (k,p)를 랜덤 선택(너무 쉬운 값 제외)
# - 최종적으로 (P1, ratio_dir, k, p) 조건에서 "P2가 유일"인지(간이 솔버)로 검증
# - 유일해만 PACK+DOCX로 저장
#
# 출력:
# - output/PCCC1_YYYYMMDD_HHMMSS.docx  (문제 2단, 정답/해설 뒤에)
# - output/PCCC1_YYYYMMDD_HHMMSS.json  (PACK)
#
# 필요:
#   pip install python-docx
# 실행:
#   python pccc1_generator.py
# ------------------------------------------------------------

import os
import json
import random
import hashlib
import time
from fractions import Fraction
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
ID_PREFIX = "PCCC1_"

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUT_DIR, exist_ok=True)

DOCX_NAME = f"PCCC1_{time.strftime('%Y%m%d_%H%M%S')}.docx"
PACK_NAME = f"PCCC1_{time.strftime('%Y%m%d_%H%M%S')}.json"

# 생성 시도 상한(빡세면 올려)
MAX_TRIES_TOTAL = 50000

# 타깃 선택 시 너무 쉬운 확률 제외(0, 1은 무조건 제외)
EXCLUDE_PROBS = {Fraction(0, 1), Fraction(1, 1)}
# 추가로 1/2까지 빼고 싶으면 아래 주석 해제
# EXCLUDE_PROBS.add(Fraction(1, 2))

# 타깃 k 범위(너무 극단값(0,6) 줄이고 싶으면 조정)
ALLOW_K = set(range(0, 7))  # {0,1,2,3,4,5,6}
# 예: 1~5만 허용하려면 아래처럼:
# ALLOW_K = set(range(1, 6))


# =========================
# Genetics helpers
# =========================
GENO_A = ["AA", "Aa", "aa"]
GENO_B = ["BB", "Bb", "bb"]
GENO_D = ["DD", "Dd", "dd"]

P2_ALL: List[str] = [A + B + D for A in GENO_A for B in GENO_B for D in GENO_D]


def parse_genotype(geno: str) -> Dict[str, str]:
    # "AAbbDd" 같은 형태
    return {"A": geno[0:2], "B": geno[2:4], "D": geno[4:6]}


def uppercase_count_pair(pair: str) -> int:
    return sum(1 for ch in pair if ch.isupper())


def lowercase_count_pair(pair: str) -> int:
    return sum(1 for ch in pair if ch.islower())


def gametes_for_pair(pair: str) -> List[str]:
    # AA -> ["A"], aa -> ["a"], Aa -> ["A","a"]
    if pair[0] == pair[1]:
        return [pair[0]]
    return [pair[0], pair[1]]


def child_pair_dist(p1_pair: str, p2_pair: str) -> Dict[str, Fraction]:
    g1 = gametes_for_pair(p1_pair)
    g2 = gametes_for_pair(p2_pair)

    def probs(g: List[str]) -> Dict[str, Fraction]:
        if len(g) == 1:
            return {g[0]: Fraction(1, 1)}
        # hetero
        return {g[0]: Fraction(1, 2), g[1]: Fraction(1, 2)}

    p_g1 = probs(g1)
    p_g2 = probs(g2)

    dist: Dict[str, Fraction] = {}
    for a, pa in p_g1.items():
        for b, pb in p_g2.items():
            # canonical: uppercase first if possible
            child = "".join(sorted([a, b], key=lambda x: (x.islower(), x)))
            if child in ("aA", "Aa"):
                child = "Aa"
            dist[child] = dist.get(child, Fraction(0, 1)) + pa * pb
    return dist


def offspring_uppercase_total_dist(p1: Dict[str, str], p2: Dict[str, str]) -> Dict[int, Fraction]:
    dA = child_pair_dist(p1["A"], p2["A"])
    dB = child_pair_dist(p1["B"], p2["B"])
    dD = child_pair_dist(p1["D"], p2["D"])

    def map_counts(dist_pair: Dict[str, Fraction]) -> Dict[int, Fraction]:
        out: Dict[int, Fraction] = {}
        for gp, pr in dist_pair.items():
            cnt = uppercase_count_pair(gp)
            out[cnt] = out.get(cnt, Fraction(0, 1)) + pr
        return out

    cA, cB, cD = map_counts(dA), map_counts(dB), map_counts(dD)

    total: Dict[int, Fraction] = {}
    for a_cnt, pa in cA.items():
        for b_cnt, pb in cB.items():
            for d_cnt, pd in cD.items():
                s = a_cnt + b_cnt + d_cnt
                total[s] = total.get(s, Fraction(0, 1)) + pa * pb * pd
    return total


def ratio_d_over_A(geno: Dict[str, str]) -> Optional[Fraction]:
    A_cnt = uppercase_count_pair(geno["A"])      # A의 수
    d_cnt = lowercase_count_pair(geno["D"])      # d의 수
    if A_cnt == 0:
        return None
    return Fraction(d_cnt, A_cnt)


def compare_ratio(r2: Optional[Fraction], r1: Optional[Fraction], op: str) -> bool:
    if r1 is None or r2 is None:
        return False
    if op == "<":
        return r2 < r1
    if op == ">":
        return r2 > r1
    return False


def random_P() -> str:
    return random.choice(GENO_A) + random.choice(GENO_B) + random.choice(GENO_D)


def infer_ratio_dir_from_P1P2(P1: str, P2: str) -> Optional[str]:
    p1 = parse_genotype(P1)
    p2 = parse_genotype(P2)
    r1 = ratio_d_over_A(p1)
    r2 = ratio_d_over_A(p2)
    if r1 is None or r2 is None:
        return None
    if r2 < r1:
        return "<"
    if r2 > r1:
        return ">"
    return None


def pick_target_from_dist(dist: Dict[int, Fraction]) -> Optional[Tuple[int, Fraction]]:
    candidates = []
    for k, p in dist.items():
        if k not in ALLOW_K:
            continue
        if p in EXCLUDE_PROBS:
            continue
        if p == 0:
            continue
        candidates.append((k, p))
    if not candidates:
        # fallback: 0 아닌 것만이라도
        fallback = [(k, p) for k, p in dist.items() if (k in ALLOW_K and p != 0)]
        return random.choice(fallback) if fallback else None
    return random.choice(candidates)


# =========================
# Unique solver
# =========================
def solve_unique_P2(P1: str, ratio_dir: str, target_k: int, target_p: Fraction) -> Tuple[Optional[str], List[str]]:
    p1 = parse_genotype(P1)
    r1 = ratio_d_over_A(p1)
    valid: List[str] = []

    for P2 in P2_ALL:
        p2 = parse_genotype(P2)
        r2 = ratio_d_over_A(p2)
        if not compare_ratio(r2, r1, ratio_dir):
            continue
        dist = offspring_uppercase_total_dist(p1, p2)
        if dist.get(target_k, Fraction(0, 1)) == target_p:
            valid.append(P2)

    if len(valid) == 1:
        return valid[0], valid
    return None, valid


# =========================
# PACK format
# =========================
@dataclass
class PackItem:
    id: str
    module_code: str
    id_prefix: str
    problem_text_md: str
    ask_line_md: str
    table_md: str
    answer_md: str
    explanation_md: str
    payload: dict


class ProblemPack:
    def __init__(self, module_code: str, out_dir: str, id_prefix: str):
        self.module_code = module_code
        self.out_dir = out_dir
        self.id_prefix = id_prefix
        self.items: List[PackItem] = []

    def add(self, item: PackItem):
        self.items.append(item)

    def save_json(self, filename: str) -> str:
        path = os.path.join(self.out_dir, filename)
        data = {
            "module_code": self.module_code,
            "id_prefix": self.id_prefix,
            "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "count": len(self.items),
            "items": [
                {
                    "id": it.id,
                    "module_code": it.module_code,
                    "id_prefix": it.id_prefix,
                    "problem_text_md": it.problem_text_md,
                    "ask_line_md": it.ask_line_md,
                    "table_md": it.table_md,
                    "answer_md": it.answer_md,
                    "explanation_md": it.explanation_md,
                    "payload": it.payload,
                }
                for it in self.items
            ],
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return path


# =========================
# Rendering helpers
# =========================
def make_id(seed: str) -> str:
    h = hashlib.sha1(seed.encode("utf-8")).hexdigest()[:10]
    return ID_PREFIX + h


def build_problem_md(P1: str, ratio_dir: str, target_k: int, target_p: Fraction) -> Tuple[str, str]:
    op_word = "작다" if ratio_dir == "<" else "크다"
    lines = [
        "문제 제목 : Polygenic Cheatkey+Capital Counting 1(PCCC1)",
        "(A/a), (B/b), (D/d) 3set 유전자에 의해 결정되는 다인자유전을 가정하자. 단, (A/a), (B/b), (D/d)는 모두 독립이다.",
        f"P1의 유전자형은 {P1}이다.",
        f"P2가 P1보다 d의 수 / A의 수 가 {op_word}.",
        f"부모 P1과 P2를 교배시킬 때 나오는 자손의 표현형이 ({target_k})일 확률은 {target_p.numerator}/{target_p.denominator}이다.",
        "",
        "※ 표현형 (n)은 자손의 대문자 대립유전자 총 개수가 n인 경우를 의미한다.",
    ]
    ask = "P2의 유전자형을 찾고 자손의 표현형의 종류와 각 표현형의 확률을 구하시오."
    return "\n".join(lines), ask


def dist_to_md_table(dist: Dict[int, Fraction]) -> str:
    rows = [(k, p) for k, p in sorted(dist.items()) if p != 0]
    md = "| 표현형 | 확률 |\n|---|---|\n"
    for k, p in rows:
        md += f"| ({k}) | {p.numerator}/{p.denominator} |\n"
    return md


def make_explanation(P1: str, P2: str, ratio_dir: str, target_k: int, target_p: Fraction,
                     candidates_count: int) -> Tuple[str, str, str]:
    p1 = parse_genotype(P1)
    p2 = parse_genotype(P2)
    r1 = ratio_d_over_A(p1)
    r2 = ratio_d_over_A(p2)

    dist = offspring_uppercase_total_dist(p1, p2)
    table_md = dist_to_md_table(dist)
    answer_md = f"P2 = {P2}"

    op_word = "작다" if ratio_dir == "<" else "크다"
    exp = []
    exp.append(f"- P1 = {P1} → A의 수 = {uppercase_count_pair(p1['A'])}, d의 수 = {lowercase_count_pair(p1['D'])} 이므로 (d/A)_P1 = {r1.numerator}/{r1.denominator}.")
    exp.append(f"- 정답 P2 = {P2} → (d/A)_P2 = {r2.numerator}/{r2.denominator} 이고, 따라서 P2의 d/A는 P1보다 {op_word}.")
    exp.append(f"- 또한 이 교배에서 자손 표현형 ({target_k}) 확률은 {target_p.numerator}/{target_p.denominator}이다.")
    exp.append(f"- 위 조건(비율 + 목표확률)을 만족하는 P2를 전수검사하면 정답이 유일(후보 {candidates_count}개 중 1개)이다.")
    exp.append("")
    exp.append("[자손 표현형 분포(=대문자 총개수)]")
    for k, p in sorted(dist.items()):
        if p != 0:
            exp.append(f"  • ({k}) : {p.numerator}/{p.denominator}")
    explanation_md = "\n".join(exp)

    return answer_md, table_md, explanation_md


# =========================
# DOCX output (2단 = 표 1x2로 구현)
# =========================
def set_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "바탕"
    style.font.size = Pt(9)


def add_problem_cell(cell, pnum: int, it: PackItem):
    p = cell.add_paragraph(f"[문제 {pnum}]")
    p.runs[0].bold = True

    for line in it.problem_text_md.splitlines():
        cell.add_paragraph(line)

    cell.add_paragraph("")
    cell.add_paragraph(f"요구사항: {it.ask_line_md}")


def make_docx(pack: ProblemPack, filename: str) -> str:
    doc = Document()
    set_doc_style(doc)

    items = pack.items
    idx = 0
    pnum = 1

    while idx < len(items):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        add_problem_cell(left, pnum, items[idx]); idx += 1; pnum += 1
        if idx < len(items):
            add_problem_cell(right, pnum, items[idx]); idx += 1; pnum += 1

        if idx < len(items):
            doc.add_page_break()

    # 정답/해설 뒤에 모아서
    doc.add_page_break()
    h = doc.add_paragraph("[정답]")
    h.runs[0].bold = True
    for i, it in enumerate(items, start=1):
        doc.add_paragraph(f"{i}번: {it.answer_md}")

    doc.add_page_break()
    h2 = doc.add_paragraph("[해설]")
    h2.runs[0].bold = True
    for i, it in enumerate(items, start=1):
        t = doc.add_paragraph(f"{i}번 해설 ({it.id})")
        t.runs[0].bold = True
        doc.add_paragraph(it.explanation_md)
        doc.add_paragraph("")
        doc.add_paragraph("[분포표]")
        doc.add_paragraph(it.table_md)
        doc.add_paragraph("")

    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    return path


# =========================
# Generator
# =========================
def generate_pack_item() -> Optional[PackItem]:
    # 1) P1, P2 무작위 생성
    P1 = random_P()
    P2_pick = random_P()
    if P1 == P2_pick:
        return None

    # 2) ratio_dir는 P1,P2에 맞춰 자동 결정(끼워맞추기)
    ratio_dir = infer_ratio_dir_from_P1P2(P1, P2_pick)
    if ratio_dir is None:
        return None

    p1 = parse_genotype(P1)
    p2 = parse_genotype(P2_pick)

    # 3) 이 조합에서 자손 분포 만들고 (k,p) 랜덤 선택
    dist = offspring_uppercase_total_dist(p1, p2)
    picked = pick_target_from_dist(dist)
    if picked is None:
        return None
    target_k, target_p = picked

    # 4) 이제 "유일정답"인지 검사
    sol, candidates = solve_unique_P2(P1, ratio_dir, target_k, target_p)
    if sol is None:
        return None

    # 5) 문제/정답/해설 구성
    problem_text_md, ask_line_md = build_problem_md(P1, ratio_dir, target_k, target_p)
    answer_md, table_md, explanation_md = make_explanation(P1, sol, ratio_dir, target_k, target_p, len(candidates))

    payload = {
        "P1": P1,
        "P2": sol,
        "ratio": "d/A",
        "ratio_dir": ratio_dir,
        "target_pheno": target_k,
        "target_prob": f"{target_p.numerator}/{target_p.denominator}",
        "offspring_dist": {
            str(k): f"{v.numerator}/{v.denominator}"
            for k, v in offspring_uppercase_total_dist(parse_genotype(P1), parse_genotype(sol)).items()
            if v != 0
        },
        "p2_candidate_count": len(candidates),
        "level": "unknown",
    }

    seed = f"{P1}|{ratio_dir}|{target_k}|{target_p}|{sol}"
    pid = make_id(seed)

    return PackItem(
        id=pid,
        module_code="PCCC1",
        id_prefix=ID_PREFIX,
        problem_text_md=problem_text_md,
        ask_line_md=ask_line_md,
        table_md=table_md,
        answer_md=answer_md,
        explanation_md=explanation_md,
        payload=payload,
    )


def main():
    random.seed()

    pack = ProblemPack(module_code="PCCC1", out_dir=OUT_DIR, id_prefix=ID_PREFIX)
    seen_ids = set()

    tries = 0
    while len(pack.items) < N_PROBLEMS and tries < MAX_TRIES_TOTAL:
        tries += 1
        if tries % 300 == 0:
            print(f"[진행] tries={tries} / made={len(pack.items)}")

        it = generate_pack_item()
        if it is None:
            continue
        if it.id in seen_ids:
            continue
        seen_ids.add(it.id)
        pack.add(it)

    if len(pack.items) < N_PROBLEMS:
        raise RuntimeError(
            f"생성 실패: {len(pack.items)}/{N_PROBLEMS} "
            f"(MAX_TRIES_TOTAL 올리거나 EXCLUDE_PROBS/ALLOW_K 완화 필요)"
        )

    pack_path = pack.save_json(PACK_NAME)
    docx_path = make_docx(pack, DOCX_NAME)

    print("✅ 완료")
    print("PACK:", pack_path)
    print("DOCX:", docx_path)


if __name__ == "__main__":
    main()
