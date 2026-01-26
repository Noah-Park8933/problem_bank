# blood_type_calculation_generator.py
# ------------------------------------------------------------
# Blood Type Calculation (ABO only)
# "임의 집단"에서 (가/나/다) 존재 여부(특성) 조합의 인원수가 주어질 때,
# (가,나,다)가 무엇(응집원A / 응집소 α / 응집소 β)인지 찾고,
# A형/B형/AB형/O형 인원을 계산하는 문제 자동 생성기.
#
# 모델(관측 특성 3개만 사용): {응집원A(AgA), 응집소 α, 응집소 β}
#  - AB형: AgA만
#  - A형 : AgA + β
#  - B형 : α만
#  - O형 : α + β
#
# 입력(문항 정보)은 아래 3개 형태를 기본으로 제공:
#  1) "X만 가지는 사람"
#  2) "Y만 가지는 사람"
#  3) "U와 V를 모두 가지는 사람"
# + 보조조건 1개(예: AB형이 O형보다 많다)로 매핑 유일화
#
# 기능:
# 1) 무작위로 혈액형 인원수(A,B,AB,O)와 라벨 매핑(가/나/다)을 생성
# 2) 위 3개 정보 + 보조조건으로 유일정답이 되는지 솔버로 검증
# 3) 유일정답만 채택하여 30문항 생성
# 4) DOCX 출력(2단, 2문제/페이지) + 정답 + 해설(근거 포함)
# 글자 크기: 9pt
# ------------------------------------------------------------
import time
import os
import random
import itertools
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
MAX_TRIES_PER_PROBLEM = 1200

# 총원 범위(원하면 바꾸기)
N_RANGE = (80, 140)

# 최소 각 혈액형 인원(0 허용하면 너무 쉬운/이상한 케이스가 많아짐)
MIN_EACH = 5

# 라벨
LABELS = ["가", "나", "다"]

# 특성(진짜 의미)
FEATURES = ["AgA", "α", "β"]  # 응집원A, 응집소α, 응집소β
FEATURE_KOR = {"AgA": "응집원 A", "α": "응집소 α", "β": "응집소 β"}

# 혈액형별 보유 특성(관측 가능한 3개만)
# AB: AgA
# A : AgA + β
# B : α
# O : α + β
TYPE_FEATURES = {
    "AB": {"AgA"},
    "A":  {"AgA", "β"},
    "B":  {"α"},
    "O":  {"α", "β"},
}

TYPES = ["A", "B", "AB", "O"]

# 보조조건 풀(유일정답 만드는 데 도움)
# 각 조건은 (문장, 판정함수)
def cond_AB_gt_O(cnt): return cnt["AB"] > cnt["O"]
def cond_AB_lt_O(cnt): return cnt["AB"] < cnt["O"]
def cond_A_gt_B(cnt):  return cnt["A"] > cnt["B"]
def cond_A_lt_B(cnt):  return cnt["A"] < cnt["B"]
def cond_B_gt_O(cnt):  return cnt["B"] > cnt["O"]
def cond_O_gt_A(cnt):  return cnt["O"] > cnt["A"]
def cond_max_AB(cnt):  return cnt["AB"] == max(cnt.values())
def cond_max_O(cnt):   return cnt["O"] == max(cnt.values())
def cond_min_B(cnt):   return cnt["B"] == min(cnt.values())

COND_POOL = [
    ("이 집단에서는 AB형이 O형보다 많다.", cond_AB_gt_O),
    ("이 집단에서는 AB형이 O형보다 적다.", cond_AB_lt_O),
    ("이 집단에서는 A형이 B형보다 많다.", cond_A_gt_B),
    ("이 집단에서는 A형이 B형보다 적다.", cond_A_lt_B),
    ("이 집단에서는 B형이 O형보다 많다.", cond_B_gt_O),
    ("이 집단에서는 O형이 A형보다 많다.", cond_O_gt_A),
    ("이 집단에서 가장 많은 혈액형은 AB형이다.", cond_max_AB),
    ("이 집단에서 가장 많은 혈액형은 O형이다.", cond_max_O),
    ("이 집단에서 가장 적은 혈액형은 B형이다.", cond_min_B),
]

# =========================
# Utility: compute observed clue values from true mapping
# =========================
def feature_of_label(true_label_map: Dict[str, str], label: str) -> str:
    """label -> feature"""
    return true_label_map[label]

def only_count(feature: str, counts: Dict[str, int]) -> int:
    """
    feature 하나만 가진 사람 수를 반환 (이 모델에서 가능한 경우만 nonzero가 되도록 counts가 구성됨).
    - only AgA = AB
    - only α   = B
    - only β   = 0 (불가능)
    """
    if feature == "AgA":
        return counts["AB"]
    if feature == "α":
        return counts["B"]
    if feature == "β":
        return 0
    raise ValueError("bad feature")

def both_count(f1: str, f2: str, counts: Dict[str, int]) -> int:
    """
    두 특성을 모두 가진 사람 수:
    - (AgA, β) = A
    - (α, β)   = O
    - (AgA, α) = 0 (동시에 불가)
    """
    s = {f1, f2}
    if s == {"AgA", "β"}:
        return counts["A"]
    if s == {"α", "β"}:
        return counts["O"]
    if s == {"AgA", "α"}:
        return 0
    raise ValueError("bad pair")

# =========================
# Solver: given clues + total N + condition, find unique mapping and counts
# =========================
def solve_unique(
    N_total: int,
    clue_only_1: Tuple[str, int],  # (label, value)
    clue_only_2: Tuple[str, int],  # (label, value)
    clue_both: Tuple[Tuple[str, str], int],  # ((label1,label2), value)
    cond_text: str,
    cond_fn
) -> Tuple[int, Optional[Dict]]:
    """
    returns (n_solutions, witness) where witness has:
      - label_map: {label -> feature}
      - counts: {A,B,AB,O}
    """
    sol = 0
    witness = None

    # try all permutations for label -> feature
    for perm in itertools.permutations(FEATURES, 3):
        label_map = dict(zip(LABELS, perm))  # e.g. 가->AgA, 나->α, 다->β

        # interpret clues under this mapping
        l1, v1 = clue_only_1
        l2, v2 = clue_only_2
        (lb1, lb2), vb = clue_both

        f_l1 = label_map[l1]
        f_l2 = label_map[l2]
        f_b1 = label_map[lb1]
        f_b2 = label_map[lb2]

        # Under our model:
        # only AgA = AB, only α = B, only β = 0
        # both (AgA,β)=A, both(α,β)=O, both(AgA,α)=0
        # So we can compute implied counts:
        AB = None
        B = None

        # only clue 1
        if f_l1 == "AgA":
            AB = v1
        elif f_l1 == "α":
            B = v1
        elif f_l1 == "β":
            if v1 != 0:
                continue  # impossible
        else:
            continue

        # only clue 2
        if f_l2 == "AgA":
            if AB is None: AB = v2
            elif AB != v2: continue
        elif f_l2 == "α":
            if B is None: B = v2
            elif B != v2: continue
        elif f_l2 == "β":
            if v2 != 0:
                continue
        else:
            continue

        if AB is None or B is None:
            # If neither only clue pins AB or B, too weak => skip
            continue

        # both clue gives A or O or impossible-0
        pair = {f_b1, f_b2}
        if pair == {"AgA", "β"}:
            A = vb
            O = N_total - (A + B + AB)
        elif pair == {"α", "β"}:
            O = vb
            A = N_total - (O + B + AB)
        elif pair == {"AgA", "α"}:
            if vb != 0:
                continue
            # then A or O not directly from vb -> underdetermined, skip
            continue
        else:
            continue

        # validity
        if A < 0 or O < 0:
            continue

        counts = {"A": A, "B": B, "AB": AB, "O": O}

        # apply condition
        if not cond_fn(counts):
            continue

        sol += 1
        witness = {"label_map": label_map, "counts": counts, "cond": cond_text}
        if sol >= 2:
            return sol, witness

    return sol, witness

# =========================
# "결정적 단서" 자동 생성(중간 버전)
# =========================
def make_reasoning_text(
    N_total: int,
    clue_only_1, clue_only_2, clue_both,
    cond_text: str,
    witness: Dict
) -> List[str]:
    """
    해설에 넣을 '결정적 근거' 문장 4~8줄 정도(중간 버전).
    핵심: 틀린 가정(라벨 매핑)이 어떤 clue와 모순되는지 1~2개씩.
    """
    lines = []
    label_map = witness["label_map"]
    counts = witness["counts"]

    # 1) "β만 가진 사람은 존재할 수 없다"를 활용한 배제 근거
    # (이 모델에서는 β 단독 혈액형이 없음)
    # only clue에 나온 라벨은 β일 수 없음(값이 0이 아닌 경우).
    (l1, v1) = clue_only_1
    (l2, v2) = clue_only_2

    if v1 != 0:
        lines.append(f"① {l1}만 가지는 사람이 {v1}명(0이 아님)이므로, {l1}={FEATURE_KOR['β']}는 불가능하다(β만 가진 혈액형이 없다).")
    if v2 != 0 and l2 != l1:
        lines.append(f"② {l2}만 가지는 사람이 {v2}명(0이 아님)이므로, {l2}={FEATURE_KOR['β']}는 불가능하다(β만 가진 혈액형이 없다).")

    # 2) only(AgA)=AB, only(α)=B 라는 결정적 대응
    # witness에서 어떤 라벨이 AgA/α인지 찍어줌
    # AgA 라벨 찾기
    inv = {feat: lab for lab, feat in label_map.items()}
    lab_AgA = inv["AgA"]
    lab_alpha = inv["α"]
    lab_beta = inv["β"]

    # Show mapping lock
    lines.append(f"③ 단독 보유는 '응집원 A만=AB형', '응집소 α만=B형'로 대응하므로, {lab_AgA}만 가진 인원=AB형, {lab_alpha}만 가진 인원=B형이다.")

    # 3) both clue가 A형 또는 O형을 고정
    ((b1, b2), vb) = clue_both
    fset = {label_map[b1], label_map[b2]}
    if fset == {"AgA", "β"}:
        lines.append(f"④ {b1}과 {b2}를 모두 가지는 사람은 (응집원 A와 β 동시)이므로 A형이며, 따라서 A형={vb}명이다.")
    elif fset == {"α", "β"}:
        lines.append(f"④ {b1}과 {b2}를 모두 가지는 사람은 (α와 β 동시)이므로 O형이며, 따라서 O형={vb}명이다.")
    else:
        # 이 케이스는 생성 단계에서 거의 나오지 않게 설계했지만, 안전 처리
        lines.append(f"④ {b1}과 {b2}를 모두 가지는 경우는 특정 혈액형에 직접 대응하므로(모순 없이) A형 또는 O형이 결정된다.")

    # 4) 조건(부등식)으로 마지막 애매함 제거했다는 한 줄
    lines.append(f"⑤ 마지막으로 '{cond_text}' 조건을 적용하면 남는 경우가 1개뿐이므로 (가,나,다)의 의미가 유일하게 결정된다.")

    # 5) 결과 요약
    lines.append(
        f"⇒ 결론: {lab_AgA}={FEATURE_KOR['AgA']}, {lab_alpha}={FEATURE_KOR['α']}, {lab_beta}={FEATURE_KOR['β']} / "
        f"A형 {counts['A']}명, B형 {counts['B']}명, AB형 {counts['AB']}명, O형 {counts['O']}명."
    )
    return lines

# =========================
# Problem generator
# =========================
def random_counts(N_total: int) -> Dict[str, int]:
    """
    A,B,AB,O 합=N_total, 각 최소 MIN_EACH 이상
    """
    # sample 4 positive parts
    # allocate remaining after minima
    rem = N_total - 4 * MIN_EACH
    if rem < 0:
        raise ValueError("N too small for minima")

    # random split rem into 4 nonnegative
    cuts = sorted(random.sample(range(rem + 3), 3))
    parts = [cuts[0], cuts[1] - cuts[0], cuts[2] - cuts[1], rem + 3 - cuts[2] - 1]
    # above trick is messy; let's just do simpler
    # We'll do direct:
    a = random.randint(0, rem)
    b = random.randint(0, rem - a)
    c = random.randint(0, rem - a - b)
    d = rem - a - b - c
    parts = [a, b, c, d]
    random.shuffle(parts)
    return {
        "A":  MIN_EACH + parts[0],
        "B":  MIN_EACH + parts[1],
        "AB": MIN_EACH + parts[2],
        "O":  MIN_EACH + parts[3],
    }

def generate_one_problem(problem_index: int):
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        N_total = random.randint(*N_RANGE)
        counts = random_counts(N_total)

        # pick a true label mapping (가/나/다 -> features)
        perm = list(FEATURES)
        random.shuffle(perm)
        true_label_map = dict(zip(LABELS, perm))

        # we will output 3 clues in fixed "shape":
        #  - two "only" clues (two different labels)
        #  - one "both" clue (two labels)
        #
        # To keep problems like your example, we generate:
        #  - one only clue should be AB (only AgA)
        #  - one only clue should be B (only α)
        #  - one both clue should be A (AgA+β)  OR sometimes O (α+β) for variety
        #
        # We still keep labels unknown, so solver must deduce mapping.

        # Decide which both-pair to use: A-type (AgA+β) or O-type (α+β)
        use_both_for = random.choice(["A", "O"])

        # Find which labels correspond to each feature
        inv_true = {feat: lab for lab, feat in true_label_map.items()}
        lab_AgA = inv_true["AgA"]
        lab_alpha = inv_true["α"]
        lab_beta = inv_true["β"]

        # Only clues: AgA-only (=AB) and α-only (=B)
        clue_only_1 = (lab_AgA, only_count("AgA", counts))  # AB
        clue_only_2 = (lab_alpha, only_count("α", counts))  # B

        # Both clue
        if use_both_for == "A":
            pair_labels = (lab_AgA, lab_beta)  # AgA + β -> A
            vb = both_count("AgA", "β", counts)  # A
        else:
            pair_labels = (lab_alpha, lab_beta)  # α + β -> O
            vb = both_count("α", "β", counts)    # O

        # randomize order inside pair for naturalness
        pair_labels = tuple(pair_labels)
        if random.random() < 0.5:
            pair_labels = (pair_labels[1], pair_labels[0])

        clue_both = (pair_labels, vb)

        # Choose a condition and ensure unique solution under solver
        random.shuffle(COND_POOL)
        for cond_text, cond_fn in COND_POOL[:6]:
            sol, wit = solve_unique(
                N_total,
                clue_only_1, clue_only_2, clue_both,
                cond_text, cond_fn
            )
            if sol == 1 and wit is not None:
                # also avoid degenerate O or A etc too small maybe
                if min(wit["counts"].values()) <= 0:
                    continue
                # build reasoning
                reasons = make_reasoning_text(
                    N_total, clue_only_1, clue_only_2, clue_both, cond_text, wit
                )
                return {
                    "N_total": N_total,
                    "clue_only_1": clue_only_1,
                    "clue_only_2": clue_only_2,
                    "clue_both": clue_both,
                    "cond_text": cond_text,
                    "witness": wit,
                    "reasons": reasons,
                }

    raise RuntimeError(f"[P{problem_index:02d}] 생성 실패(유일정답 확보 불가)")

# =========================
# DOCX rendering
# =========================
def make_docx():
    # output 폴더 생성
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="BLOODC", out_dir=out_dir, id_prefix="BLOODC_")

    os.makedirs(out_dir, exist_ok=True)
    # 저장 파일명 (시간 붙여서 중복 방지)
    filename = os.path.join(out_dir, f"Blood_Type_Calculation_{int(time.time())}.docx")
    doc = Document()
    style = doc.styles["Normal"]
    set_page_margins(doc)
    style.font.name = "바탕"
    style.font.size = Pt(9)

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        pdata = generate_one_problem(i)
        problems.append(pdata)
        # PACK record
        # PACK record (표시용 텍스트/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            f"아래는 {pdata['N_total']}명 집단에서 가/나/다의 유무로 분류한 정보이다.",
            f"{pdata['clue_only_1']['label']}만 가지는 사람: {pdata['clue_only_1']['count']}",
            f"{pdata['clue_only_2']['label']}만 가지는 사람: {pdata['clue_only_2']['count']}",
            f"{pdata['clue_both']['label1']}와 {pdata['clue_both']['label2']}를 모두 가지는 사람: {pdata['clue_both']['count']}",
            f"조건: {pdata['cond_text']}",
        ])
        ask_line_md = "표에서 가/나/다가 무엇인지 찾고, A형/B형/AB형/O형 사람 수를 구하시오."
        answer_md = pdata['witness'].get('final_answer_line', '')
        explanation_md = "\n".join(pdata.get("reasons", [])) if isinstance(pdata.get("reasons"), list) else str(pdata.get("reasons"))

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": "",
            "full_table_md": "",
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "data": pdata,
        }))
        if i % 5 == 0:
            print(f"진행: {i}/{N_PROBLEMS}")

    # 문제(2단, 2문제/페이지)
    idx = 0
    pnum = 1
    while idx < N_PROBLEMS:
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        def put(cell, num, pdata):
            cell.add_paragraph(f"[문제 {num}]").runs[0].bold = True
            cell.add_paragraph(f"아래는 {pdata['N_total']}명으로 구성되는 집단에서 (가), (나), (다)의 유무를 나타낸 것이다.")
            cell.add_paragraph("가, 나, 다는 (응집원 A, 응집소 α, 응집소 β)를 순서 없이 나타낸 것이다.")
            cell.add_paragraph(f"단, {pdata['cond_text']}")
            cell.add_paragraph("")

            (l1, v1) = pdata["clue_only_1"]
            (l2, v2) = pdata["clue_only_2"]
            ((b1, b2), vb) = pdata["clue_both"]

            cell.add_paragraph(f"{l1}만 가지는 사람 : {v1}")
            cell.add_paragraph(f"{l2}만 가지는 사람 : {v2}")
            cell.add_paragraph(f"{b1}와 {b2}를 모두 가지는 사람 : {vb}")
            cell.add_paragraph("")
            cell.add_paragraph("표에서 가, 나, 다를 찾고 이 집단에서 A형, B형, AB형, O형인 사람의 수를 구하시오.")

        put(left, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            put(right, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            doc.add_page_break()

    # 정답
    doc.add_page_break()
    doc.add_paragraph("[정답]").runs[0].bold = True
    for i, pdata in enumerate(problems, start=1):
        cnt = pdata["witness"]["counts"]
        lm = pdata["witness"]["label_map"]
        doc.add_paragraph(
            f"{i}번: 가={FEATURE_KOR[lm['가']]}, 나={FEATURE_KOR[lm['나']]}, 다={FEATURE_KOR[lm['다']]} / "
            f"A형 {cnt['A']}명, B형 {cnt['B']}명, AB형 {cnt['AB']}명, O형 {cnt['O']}명"
        )

    # 해설(근거)
    doc.add_page_break()
    doc.add_paragraph("[해설(결정적 근거 포함)]").runs[0].bold = True
    for i, pdata in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번").runs[0].bold = True

        # 조건/정보 재제시(짧게)
        (l1, v1) = pdata["clue_only_1"]
        (l2, v2) = pdata["clue_only_2"]
        ((b1, b2), vb) = pdata["clue_both"]

        doc.add_paragraph(f"총원: {pdata['N_total']}명 / 조건: {pdata['cond_text']}")
        doc.add_paragraph(f"{l1}만: {v1}, {l2}만: {v2}, {b1}와 {b2} 모두: {vb}")
        doc.add_paragraph("")

        doc.add_paragraph("결정적 근거")
        for s in pdata["reasons"]:
            doc.add_paragraph(f"- {s}")

        cnt = pdata["witness"]["counts"]
        lm = pdata["witness"]["label_map"]
        doc.add_paragraph("")
        doc.add_paragraph(
            f"정답: 가={FEATURE_KOR[lm['가']]}, 나={FEATURE_KOR[lm['나']]}, 다={FEATURE_KOR[lm['다']]} / "
            f"A형 {cnt['A']}명, B형 {cnt['B']}명, AB형 {cnt['AB']}명, O형 {cnt['O']}명"
        )
        doc.add_paragraph("")

    doc.save(filename)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {filename}")


if __name__ == "__main__":
    make_docx()