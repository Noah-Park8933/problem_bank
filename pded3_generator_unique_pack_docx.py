# pded3_generator_unique_pack_docx.py
# ------------------------------------------------------------
# Polygenic Deduction 3 (PDED3) — 완전 연관(r=0) 버전
# ✅ 형식 고정 + 레버(A,B,C,D) 전부 랜덤(=문제 조건 랜덤)
# ✅ 유일정답(=해 1개)만 통과
# ✅ 30문제 DOCX 출력 + PACK JSON 저장(problem_text_md/ask_line_md 포함)
#
# 필요 설치:
#   pip install python-docx
#
# 실행:
#   python pded3_generator_unique_pack_docx.py
#
# 출력:
#   output/PDED3_YYYYMMDD_HHMMSS.docx
#   output/PDED3_YYYYMMDD_HHMMSS.pack.json
# ------------------------------------------------------------

import os, json, time, random, hashlib
from dataclasses import dataclass
from typing import Dict, Tuple, List
from fractions import Fraction

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUT_DIR, exist_ok=True)

MODULE_CODE = "PDED3"
ID_PREFIX = "PDED3_"

# A 레버(타겟) 후보: 실제로 나올 법한 분수만(필터 통과 빨라짐)
A_TARGETS = [Fraction(1, 16), Fraction(1, 8), Fraction(3, 16), Fraction(1, 4), Fraction(3, 8)]

# B 레버(표현형 종류 개수) 후보 + 가중치(5가 제일 잘 나옴)
B_TARGETS = [(4, 0.20), (5, 0.55), (6, 0.25)]

# C 레버(A수/D수 관계)
C_RELATIONS = ["EQ", "GT", "LT"]  # 같음/큼/작음

# D 레버(비교 문장 종류)
D_COMPARATORS = ["B_COUNT", "A_COUNT", "HET_COUNT", "DOM_TOTAL"]

MAX_TRIES_PER_PROBLEM = 5000


# =========================
# ProblemPack(의존성 최소화: 단독 파일로 동작)
# =========================
class ProblemPack:
    def __init__(self, module_code: str, out_dir: str, id_prefix: str):
        self.module_code = module_code
        self.out_dir = out_dir
        self.id_prefix = id_prefix
        self.items = []

    def _make_id(self) -> str:
        # 짧고 충돌 거의 없는 id
        raw = f"{self.module_code}-{time.time_ns()}-{random.random()}"
        h = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:10]
        return f"{self.id_prefix}{h}"

    def add(self, payload: dict, problem_text_md: str, ask_line_md: str,
            answer_text_md: str, solution_md: str, difficulty: int):
        pid = self._make_id()
        self.items.append({
            "id": pid,
            "module": self.module_code,
            "id_prefix": self.id_prefix,
            "difficulty": difficulty,
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "answer_text_md": answer_text_md,
            "solution_md": solution_md,
            "payload": payload,
        })
        return pid

    def save_json(self, filename: str) -> str:
        path = os.path.join(self.out_dir, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump({
                "module": self.module_code,
                "id_prefix": self.id_prefix,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                "items": self.items
            }, f, ensure_ascii=False, indent=2)
        return path


# =========================
# Genetics core
# =========================
HAPS = ["AB", "Ab", "aB", "ab"]  # 완전 연관 => gamete는 hap 자체

def hap_pair_all() -> List[Tuple[str, str]]:
    # (h1,h2) with replacement (순서 무관)
    out = []
    for i in range(len(HAPS)):
        for j in range(i, len(HAPS)):
            out.append((HAPS[i], HAPS[j]))
    return out

HAP_PAIRS = hap_pair_all()

def hap_pair_to_AaBb(h1: str, h2: str) -> str:
    # hap pair -> genotype string "AA", "Aa", "aa" / "BB","Bb","bb"
    A_alleles = [h1[0], h2[0]]  # 'A'/'a'
    B_alleles = [h1[1], h2[1]]  # 'B'/'b'

    def norm2(x, y, up, low):
        s = "".join(sorted([x, y], key=lambda t: 0 if t == up else 1))
        if s == up+up: return up+up
        if s == low+low: return low+low
        return up+low

    Ageno = norm2(A_alleles[0], A_alleles[1], "A", "a")
    Bgeno = norm2(B_alleles[0], B_alleles[1], "B", "b")
    return Ageno + Bgeno  # "AaBb" 형태

def Dgeno_all() -> List[str]:
    return ["DD", "Dd", "dd"]

def gametes_from_hap_pair(h1: str, h2: str) -> Dict[str, Fraction]:
    # 완전 연관(r=0): 부모 AB gamete 분포
    if h1 == h2:
        return {h1: Fraction(1, 1)}
    return {h1: Fraction(1, 2), h2: Fraction(1, 2)}

def gametes_from_D(geno: str) -> Dict[str, Fraction]:
    if geno == "DD": return {"D": Fraction(1, 1)}
    if geno == "dd": return {"d": Fraction(1, 1)}
    return {"D": Fraction(1, 2), "d": Fraction(1, 2)}

def cross_dist(gam1: Dict[str, Fraction], gam2: Dict[str, Fraction], combine_fn):
    out = {}
    for g1, p1 in gam1.items():
        for g2, p2 in gam2.items():
            z = combine_fn(g1, g2)
            out[z] = out.get(z, Fraction(0, 1)) + p1 * p2
    return out

def combine_haps_to_AaBb(h1: str, h2: str) -> Tuple[str, str, str]:
    # return (Ageno, Bgeno, full "AaBb")
    s = hap_pair_to_AaBb(h1, h2)  # "AaBb"
    return (s[:2], s[2:], s)

def combine_D(a: str, b: str) -> str:
    # allele + allele => genotype normalized
    pair = "".join(sorted([a, b], key=lambda t: 0 if t.isupper() else 1))
    if pair == "DD": return "DD"
    if pair == "dd": return "dd"
    return "Dd"

def count_caps_in_geno(AB: str, D: str) -> int:
    # AB = "AaBb" , D = "Dd" etc
    cnt = 0
    # A
    cnt += 1 if AB[0] == "A" else 0
    cnt += 1 if AB[1] == "A" else 0
    # B
    cnt += 1 if AB[2] == "B" else 0
    cnt += 1 if AB[3] == "B" else 0
    # D
    cnt += 1 if D[0] == "D" else 0
    cnt += 1 if D[1] == "D" else 0
    return cnt

def parent_stats(h1: str, h2: str, D: str):
    AB = hap_pair_to_AaBb(h1, h2)  # "AaBb"
    A_cnt = (1 if AB[0] == "A" else 0) + (1 if AB[1] == "A" else 0)
    B_cnt = (1 if AB[2] == "B" else 0) + (1 if AB[3] == "B" else 0)
    D_cnt = (1 if D[0] == "D" else 0) + (1 if D[1] == "D" else 0)
    het_cnt = (1 if AB[:2] == "Aa" else 0) + (1 if AB[2:] == "Bb" else 0) + (1 if D == "Dd" else 0)
    dom_total = A_cnt + B_cnt + D_cnt
    return AB, A_cnt, B_cnt, D_cnt, het_cnt, dom_total

def ratio_compare(A1, D1, A2, D2, rel: str) -> bool:
    # D=0이면 금지(나눗셈/비교 불가) => false
    if D1 == 0 or D2 == 0:
        return False
    r1 = Fraction(A1, D1)
    r2 = Fraction(A2, D2)
    if rel == "EQ": return r1 == r2
    if rel == "GT": return r1 > r2
    return r1 < r2


# =========================
# Build offspring distributions
# =========================
def offspring_distribution(hp1: Tuple[str, str], D1: str, hp2: Tuple[str, str], D2: str):
    (h1a, h1b) = hp1
    (h2a, h2b) = hp2

    g1_ab = gametes_from_hap_pair(h1a, h1b)
    g2_ab = gametes_from_hap_pair(h2a, h2b)
    ab_dist_haps = cross_dist(g1_ab, g2_ab, lambda x, y: tuple(sorted([x, y])))
    # ab_dist_haps: key = (hapX, hapY) (unordered), value=prob

    # convert to AB genotype (AaBb)
    ab_dist = {}
    for (ha, hb), p in ab_dist_haps.items():
        _, _, AB = combine_haps_to_AaBb(ha, hb)
        ab_dist[AB] = ab_dist.get(AB, Fraction(0, 1)) + p

    g1_d = gametes_from_D(D1)
    g2_d = gametes_from_D(D2)
    d_dist = cross_dist(g1_d, g2_d, combine_D)

    # full genotype dist over (AB, D)
    full = {}
    for AB, pAB in ab_dist.items():
        for Dg, pD in d_dist.items():
            full[(AB, Dg)] = full.get((AB, Dg), Fraction(0, 1)) + pAB * pD

    # phenotype dist
    ph = {}
    for (AB, Dg), p in full.items():
        k = count_caps_in_geno(AB, Dg)
        ph[k] = ph.get(k, Fraction(0, 1)) + p

    return full, ph


# =========================
# Lever sampling
# =========================
def weighted_choice(options):
    # options = [(value, weight), ...]
    r = random.random() * sum(w for _, w in options)
    s = 0.0
    for v, w in options:
        s += w
        if r <= s:
            return v
    return options[-1][0]

def frac_to_str(f: Fraction) -> str:
    if f.denominator == 1:
        return str(f.numerator)
    return f"{f.numerator}/{f.denominator}"


# =========================
# Unique solver for one problem spec
# =========================
@dataclass
class Spec:
    A_prob: Fraction      # target P(AaBbDd)
    B_n_ph: int           # target number of phenotype categories
    C_rel: str            # EQ/GT/LT for A/D ratio
    D_cmp: str            # comparator type for "P1 is more than P2"

def comparator_ok(D_cmp: str, s1, s2) -> bool:
    # s = (AB, A_cnt,B_cnt,D_cnt,het_cnt,dom_total)
    if D_cmp == "B_COUNT": return s1[2] > s2[2]
    if D_cmp == "A_COUNT": return s1[1] > s2[1]
    if D_cmp == "HET_COUNT": return s1[4] > s2[4]
    return s1[5] > s2[5]

def comparator_sentence(D_cmp: str) -> str:
    if D_cmp == "B_COUNT": return "P1이 P2보다 B의 수가 많다."
    if D_cmp == "A_COUNT": return "P1이 P2보다 A의 수가 많다."
    if D_cmp == "HET_COUNT": return "P1이 P2보다 이형접합성 수가 많다."
    return "P1이 P2보다 우성 대립유전자 총수가 많다."

def ratio_sentence(C_rel: str) -> str:
    if C_rel == "EQ": return "P1과 P2에서 A의 수/D의 수는 같다."
    if C_rel == "GT": return "P1의 A의 수/D의 수가 P2보다 크다."
    return "P1의 A의 수/D의 수가 P2보다 작다."

def solve_unique(spec: Spec):
    sols = []

    # P1 fixed condition in your format:
    # "P1에서 유전자형이 ABD인 생식세포가 형성될 수 있다."
    # => P1 AB gamete includes 'AB' + D gamete includes 'D'
    # => P1 hap-pair must include "AB" and D genotype != "dd"
    for hp1 in HAP_PAIRS:
        for D1 in Dgeno_all():
            AB1, A1, B1, D1c, het1, dom1 = parent_stats(hp1[0], hp1[1], D1)
            if "AB" not in hp1:  # hap pair contains AB?
                continue
            if D1 == "dd":
                continue

            for hp2 in HAP_PAIRS:
                for D2 in Dgeno_all():
                    AB2, A2, B2, D2c, het2, dom2 = parent_stats(hp2[0], hp2[1], D2)

                    # C: ratio constraint
                    if not ratio_compare(A1, D1c, A2, D2c, spec.C_rel):
                        continue

                    # D: comparator constraint
                    if not comparator_ok(spec.D_cmp, (AB1,A1,B1,D1c,het1,dom1), (AB2,A2,B2,D2c,het2,dom2)):
                        continue

                    # offspring distributions
                    full, ph = offspring_distribution(hp1, D1, hp2, D2)

                    # A: target probability
                    target_key = ("AaBb", "Dd")  # AB="AaBb", D="Dd"
                    p_target = full.get(target_key, Fraction(0, 1))
                    if p_target != spec.A_prob:
                        continue

                    # B: phenotype category count
                    ph_keys = sorted([k for k, p in ph.items() if p > 0])
                    if len(ph_keys) != spec.B_n_ph:
                        continue

                    sols.append({
                        "P1": {"hap": hp1, "D": D1, "AB": AB1},
                        "P2": {"hap": hp2, "D": D2, "AB": AB2},
                        "full": full,
                        "ph": ph,
                        "ph_keys": ph_keys,
                    })

                    if len(sols) > 1:
                        return sols  # not unique

    return sols


# =========================
# Problem build & docs
# =========================
def fmt_parent(AB: str, D: str) -> str:
    # "AaBb" + "Dd" => "AaBbDd"
    return f"{AB}{D}"

def phenotype_distribution_to_lines(ph: Dict[int, Fraction]) -> List[str]:
    keys = sorted([k for k, p in ph.items() if p > 0])
    lines = []
    for k in keys:
        lines.append(f"({k}) : {frac_to_str(ph[k])}")
    return lines

def make_solution_text(sol) -> Tuple[str, str]:
    P1 = sol["P1"]; P2 = sol["P2"]
    ph = sol["ph"]

    p1_g = fmt_parent(P1["AB"], P1["D"])
    p2_g = fmt_parent(P2["AB"], P2["D"])

    # answer
    ans = []
    ans.append(f"P1 = {p1_g}")
    ans.append(f"P2 = {p2_g}")
    ans.append("")
    ans.append("자손 표현형(대문자 개수) 분포:")
    for line in phenotype_distribution_to_lines(ph):
        ans.append("  " + line)
    answer_text = "\n".join(ans)

    # solution (핵심 근거)
    soln = []
    soln.append("핵심 계산:")
    soln.append("- (A/a),(B/b)는 완전 연관이므로, 각 부모의 AB 생식세포는 1종(동형) 또는 2종(이형, 각 1/2)만 나온다.")
    soln.append("- (D/d)는 독립 분리로 일반적인 멘델 분리(DD/Dd/dd)에 따른다.")
    soln.append(f"- P(AaBbDd) = {frac_to_str(sol['full'].get(('AaBb','Dd'), Fraction(0,1)))} 조건을 만족하는 (P1,P2)만 남긴다.")
    soln.append(f"- 표현형 종류 개수(=대문자 개수 종류)가 {len(sol['ph_keys'])}가지가 되도록 필터링한다.")
    soln.append("")
    soln.append("자손 표현형 확률(대문자 개수):")
    for line in phenotype_distribution_to_lines(sol["ph"]):
        soln.append("  " + line)
    solution_md = "\n".join(soln)
    return answer_text, solution_md

def difficulty_score(spec: Spec) -> int:
    # 간단 점수(웹 분류용): 조건이 빡셀수록 높게
    score = 1
    if spec.B_n_ph == 4: score += 2
    if spec.A_prob in [Fraction(1,16), Fraction(3,16)]: score += 2
    if spec.C_rel != "EQ": score += 1
    if spec.D_cmp in ["HET_COUNT","DOM_TOTAL"]: score += 1
    return min(10, score)

def make_problem_text(spec: Spec) -> Tuple[str, str]:
    # 형식 고정(네가 준 문장 틀 유지)
    # - 표현형은 (숫자)=대문자 개수
    # - (A/a),(B/b)는 연관(완전 연관)
    problem = []
    problem.append("문제 제목 : Polygenic Deduction 3(PDED3)")
    problem.append("(A/a), (B/b), (D/d) 3set 유전자에 의해 결정되는 다인자유전을 가정하자. 단, (A/a), (B/b)는 연관이다(완전 연관).")
    problem.append(f"부모 P1과 P2를 교배시킬 때 나오는 자손의 유전자형이 AaBbDd일 확률은 {frac_to_str(spec.A_prob)}이다.")
    problem.append(f"부모 P1과 P2를 교배시킬 때 나오는 자손의 표현형은 {spec.B_n_ph}가지이다. (표현형 (k) = 대문자 대립유전자 개수)")
    problem.append(ratio_sentence(spec.C_rel))
    problem.append(comparator_sentence(spec.D_cmp))
    problem.append("P1에서 유전자형이 ABD인 생식세포가 형성될 수 있다.")
    problem_text_md = "\n".join(problem)

    ask = "부모 P1, P2의 유전자형을 찾고 자손의 표현형의 종류와 각 표현형의 확률을 구하시오."
    return problem_text_md, ask


# =========================
# Main generator
# =========================
def generate_one_unique_problem() -> dict:
    for _ in range(MAX_TRIES_PER_PROBLEM):
        spec = Spec(
            A_prob=random.choice(A_TARGETS),
            B_n_ph=weighted_choice(B_TARGETS),
            C_rel=random.choice(C_RELATIONS),
            D_cmp=random.choice(D_COMPARATORS),
        )

        sols = solve_unique(spec)
        if len(sols) != 1:
            continue

        sol = sols[0]
        problem_text_md, ask_line_md = make_problem_text(spec)
        answer_text_md, solution_md = make_solution_text(sol)

        payload = {
            "spec": {
                "A_prob": frac_to_str(spec.A_prob),
                "B_n_ph": spec.B_n_ph,
                "C_rel": spec.C_rel,
                "D_cmp": spec.D_cmp,
            },
            "solution": {
                "P1": {"AB": sol["P1"]["AB"], "D": sol["P1"]["D"], "hap": list(sol["P1"]["hap"])},
                "P2": {"AB": sol["P2"]["AB"], "D": sol["P2"]["D"], "hap": list(sol["P2"]["hap"])},
            },
            "phenotype_dist": {str(k): frac_to_str(v) for k, v in sol["ph"].items() if v > 0},
        }

        return {
            "spec": spec,
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "answer_text_md": answer_text_md,
            "solution_md": solution_md,
            "payload": payload,
        }

    raise RuntimeError("유일정답 문제 생성 실패: MAX_TRIES_PER_PROBLEM 늘리거나 레버 후보 조정 필요")


def docx_add_problem(doc: Document, idx: int, p: dict):
    doc.add_paragraph(f"[{idx}번]").runs[0].bold = True

    for line in p["problem_text_md"].split("\n"):
        doc.add_paragraph(line)

    doc.add_paragraph("")  # spacing
    ask = doc.add_paragraph("요구사항: " + p["ask_line_md"])
    ask.runs[0].bold = True
    doc.add_paragraph("")


def make_docx_and_pack():
    ts = time.strftime("%Y%m%d_%H%M%S")
    docx_name = f"{MODULE_CODE}_{ts}.docx"
    pack_name = f"{MODULE_CODE}_{ts}.pack.json"

    # DOCX
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "바탕"
    style.font.size = Pt(10)

    title = doc.add_paragraph("PDED3 자동 생성 문항")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    doc.add_paragraph("")

    pack = ProblemPack(module_code=MODULE_CODE, out_dir=OUT_DIR, id_prefix=ID_PREFIX)

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        p = generate_one_unique_problem()
        problems.append(p)

        # pack
        pid = pack.add(
            payload=p["payload"],
            problem_text_md=p["problem_text_md"],
            ask_line_md=p["ask_line_md"],
            answer_text_md=p["answer_text_md"],
            solution_md=p["solution_md"],
            difficulty=difficulty_score(p["spec"]),
        )

        # docx
        docx_add_problem(doc, i, p)

        # 진행 상황
        print(f"[{i:02d}/{N_PROBLEMS}] OK  id={pid}  A={p['payload']['spec']['A_prob']}  B={p['payload']['spec']['B_n_ph']}")

        if i != N_PROBLEMS:
            doc.add_page_break()

    # 정답/해설 모아서
    doc.add_page_break()
    h = doc.add_paragraph("[정답 및 해설]")
    h.runs[0].bold = True
    doc.add_paragraph("")

    for i, p in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번 정답")
        ans = doc.add_paragraph(p["answer_text_md"])
        ans.paragraph_format.space_after = Pt(6)
        doc.add_paragraph("해설(핵심)")
        sol = doc.add_paragraph(p["solution_md"])
        sol.paragraph_format.space_after = Pt(12)

    docx_path = os.path.join(OUT_DIR, docx_name)
    doc.save(docx_path)

    pack_path = pack.save_json(pack_name)

    print("")
    print("✅ 저장 완료")
    print(" - DOCX:", docx_path)
    print(" - PACK:", pack_path)


if __name__ == "__main__":
    random.seed()  # 시스템 시드
    make_docx_and_pack()
