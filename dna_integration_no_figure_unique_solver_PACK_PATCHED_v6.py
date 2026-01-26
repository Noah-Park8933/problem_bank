# dna_integration_no_figure_unique_solver.py
# ------------------------------------------------------------
# ✅ DNA Integration (그림 X 버전) 자동 출제기 + 유일정답 솔버 + DOCX 출력
#
# 문제 형식(사용자 예제 고정):
#  - 4개 세포: (가)(나)(다)(라)
#  - 2개는 남성 I의 세포, 2개는 여성 II의 세포
#  - (A/a), (B/b), (D/d) 대립유전자쌍
#  - 그 중 2개는 X염색체, 1개는 상염색체
#  - 표는 3개 부분합 열만 제시:
#      1) A + B + D
#      2) a + b + D
#      3) a + B + d
#  - ?는 일부 칸에만 배치 (부등식 조건 없음)
#
# 고정(요청 반영):
#  - 각 사람은 2개의 세포만 등장:
#      2n(2) 1개 + (n(2) 또는 n(1)) 1개   (총 4세포)
#
# 출력:
#  - 30문제 DOCX
#  - 맨 뒤 해설: 정답(각 세포의 성별/상태), ? 채운 표, 완성표(A,a,B,b,D,d),
#               X염색체 유전자쌍 2개/상염색체 1개 표시,
#               "결정적 단서" 문장 자동 생성(가능한 범위에서)
#
# 실행:
#   pip install python-docx
#   python dna_integration_no_figure_unique_solver.py
# ------------------------------------------------------------

import os
import time
import random
from dataclasses import dataclass
from typing import Dict, List, Tuple, Any, Optional

from docx import Document
from docx.shared import Pt
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)




# =========================
# CONFIG
# =========================
N_PROBLEMS = 30

FONT_NAME = "바탕"
FONT_SIZE_PT = 9

OUT_DIR = "output"
OUT_FILE_PREFIX = "DNA_Integration_NoFigure_Unique"

ROW_LABELS = ["가", "나", "다", "라"]
COL_LABELS = ["A+B+D", "a+b+D", "a+B+d"]

# ? 배치 설정 (부등식 없음)
MASK_MIN = 2
MASK_MAX = 5
MIN_KNOWN_PER_ROW = 1
MIN_KNOWN_PER_COL = 2

# 생성 시도
MAX_TRIES_PER_PROBLEM = 2500
LOG_EVERY = 100


# =========================
# Helpers: allele DNA amounts by state/sex/chromosome type
# =========================
def pair_options_autosome_2n2():
    # 2n(2)에서 상염색체 대립유전자쌍의 (대문자, 소문자) DNA량 가능한 값
    # AA: (2,0), Aa: (1,1), aa: (0,2)
    return [(2, 0), (1, 1), (0, 2)]

def pair_options_femaleX_2n2():
    # 여성(XX)의 X유전자도 "쌍"으로 존재하므로 상염색체와 동일하게 취급
    return pair_options_autosome_2n2()

def pair_options_maleX_2n2():
    # 남성(XY)의 X연관 유전자는 X가 1개이므로 2n(2)에서 합이 1:
    # A만 있거나 a만 있는 형태 -> (1,0) 또는 (0,1)
    return [(1, 0), (0, 1)]

def n2_from_allele(allele_is_upper: bool) -> Tuple[int, int]:
    # n(2): (2,0) 또는 (0,2)만 가능
    return (2, 0) if allele_is_upper else (0, 2)

def n1_from_n2(p: Tuple[int, int]) -> Tuple[int, int]:
    # n(2)->n(1): 절반
    if p == (2, 0):
        return (1, 0)
    if p == (0, 2):
        return (0, 1)
    if p == (0, 0):
        return (0, 0)
    # 방어
    return (0, 0)


# =========================
# World model
# =========================
@dataclass
class World:
    # which pair is autosome? index 0:(A/a), 1:(B/b), 2:(D/d)
    autosome_index: int
    x_indices: List[int]  # two loci indices on X

    # each person has two cells: 2n(2) and n-state (n(2) or n(1))
    male_n_state: str     # "n2" or "n1"
    female_n_state: str   # "n2" or "n1"
    male_gamete: str      # "X" or "Y" (male n-cell only; affects X loci)

    # genotypes at 2n(2) level (pair amounts) for each locus:
    # dict locus_index -> (Upper, lower) DNA amounts in 2n(2)
    male_2n2_pairs: Dict[int, Tuple[int, int]]
    female_2n2_pairs: Dict[int, Tuple[int, int]]

    # transmitted allele choice for n-cell when locus is diploid (autosome or female X)
    # dict locus_index -> bool (True=Upper allele transmitted, False=lower transmitted)
    male_n_transmit: Dict[int, bool]    # for autosome only
    female_n_transmit: Dict[int, bool]  # for autosome + female X

    # built full allele table for each row label: dict row -> dict allele -> int
    # alleles are: "A","a","B","b","D","d"
    full_alleles_by_row: Dict[str, Dict[str, int]]

    # state/sex per row
    row_meta: Dict[str, Dict[str, str]]  # row -> {"person":"I/II", "sex":"male/female", "state":"2n(2)/n(2)/n(1)"}

    # integration sums full and masked
    sums_full: Dict[str, Dict[str, int]]   # row -> col -> value
    sums_masked: Dict[str, Dict[str, Any]] # row -> col -> value or "?"


# =========================
# Build full allele amounts for a specific cell type
# =========================
def pairs_to_alleles(pairs: Dict[int, Tuple[int, int]]) -> Dict[str, int]:
    # locus index mapping: 0->A/a, 1->B/b, 2->D/d
    out = {}
    for li, (U, L) in pairs.items():
        if li == 0:
            out["A"], out["a"] = U, L
        elif li == 1:
            out["B"], out["b"] = U, L
        else:
            out["D"], out["d"] = U, L
    return out

def build_n2_pairs_from_2n2_pairs_for_person(sex: str,
                                            autosome_index: int,
                                            x_indices: List[int],
                                            pairs_2n2: Dict[int, Tuple[int, int]],
                                            transmit: Dict[int, bool],
                                            male_gamete: str = "X") -> Dict[int, Tuple[int, int]]:
    """
    sex:
      - "male": autosome is diploid -> choose transmitted allele by transmit[autosome_index]
                X loci: if gamete Y -> (0,0), else from male 2n2 X locus -> (2,0) or (0,2)
      - "female": autosome diploid choose transmit; X loci are also diploid choose transmit
    """
    out = {}

    for li in [0, 1, 2]:
        is_autosome = (li == autosome_index)
        is_x = (li in x_indices)

        if sex == "male":
            if is_autosome:
                # autosome genotype encoded in pairs_2n2: (2,0)/(1,1)/(0,2)
                U2, L2 = pairs_2n2[li]
                # if hetero (1,1), can transmit either; else fixed
                if (U2, L2) == (2, 0):
                    out[li] = (2, 0)
                elif (U2, L2) == (0, 2):
                    out[li] = (0, 2)
                else:
                    out[li] = n2_from_allele(transmit[li])
            else:
                # X locus (male) or (shouldn't be) because exactly 2 X loci anyway
                if is_x:
                    if male_gamete == "Y":
                        out[li] = (0, 0)
                    else:
                        # male X 2n2 is (1,0) or (0,1) -> n2 becomes (2,0) or (0,2)
                        U2, L2 = pairs_2n2[li]
                        if (U2, L2) == (1, 0):
                            out[li] = (2, 0)
                        else:
                            out[li] = (0, 2)
                else:
                    # 이론상 없음(요구: X 2개 + autosome 1개)
                    out[li] = (0, 0)

        else:
            # female
            if is_autosome or is_x:
                U2, L2 = pairs_2n2[li]
                if (U2, L2) == (2, 0):
                    out[li] = (2, 0)
                elif (U2, L2) == (0, 2):
                    out[li] = (0, 2)
                else:
                    out[li] = n2_from_allele(transmit[li])
            else:
                out[li] = (0, 0)

    return out

def n_pairs_to_n1_pairs(pairs_n2: Dict[int, Tuple[int, int]]) -> Dict[int, Tuple[int, int]]:
    return {li: n1_from_n2(p) for li, p in pairs_n2.items()}


# =========================
# Integration sums (fixed columns)
# =========================
def compute_sums_from_alleles(alleles: Dict[str, int]) -> Dict[str, int]:
    # columns fixed:
    # 1) A+B+D
    # 2) a+b+D
    # 3) a+B+d
    return {
        "A+B+D": alleles["A"] + alleles["B"] + alleles["D"],
        "a+b+D": alleles["a"] + alleles["b"] + alleles["D"],
        "a+B+d": alleles["a"] + alleles["B"] + alleles["d"],
    }


# =========================
# Masking
# =========================
def mask_table(sums_full: Dict[str, Dict[str, int]]) -> Dict[str, Dict[str, Any]]:
    # choose positions to mask subject to per-row/per-col minimums
    rows = ROW_LABELS[:]
    cols = COL_LABELS[:]
    all_pos = [(r, c) for r in rows for c in cols]
    random.shuffle(all_pos)

    m = random.randint(MASK_MIN, MASK_MAX)
    masked = {r: {c: sums_full[r][c] for c in cols} for r in rows}

    # track known counts
    known_row = {r: len(cols) for r in rows}
    known_col = {c: len(rows) for c in cols}

    masked_count = 0
    for r, c in all_pos:
        if masked_count >= m:
            break
        # if already masked, skip
        if masked[r][c] == "?":
            continue
        # ensure row/col keep minimum known
        if known_row[r] - 1 < MIN_KNOWN_PER_ROW:
            continue
        if known_col[c] - 1 < MIN_KNOWN_PER_COL:
            continue
        masked[r][c] = "?"
        known_row[r] -= 1
        known_col[c] -= 1
        masked_count += 1

    return masked


# =========================
# Solver: enumerate all worlds consistent with masked sums
# =========================
def masked_match(masked: Dict[str, Dict[str, Any]], candidate: Dict[str, Dict[str, int]]) -> bool:
    for r in ROW_LABELS:
        for c in COL_LABELS:
            v = masked[r][c]
            if v == "?":
                continue
            if int(v) != int(candidate[r][c]):
                return False
    return True

def choose_autosome_and_x() -> List[Tuple[int, List[int]]]:
    # autosome index is one of 0,1,2; X indices are the other two
    out = []
    for auto in [0, 1, 2]:
        xs = [i for i in [0, 1, 2] if i != auto]
        out.append((auto, xs))
    return out

def genotype_choices_for_2n2(sex: str, li: int, autosome_index: int, x_indices: List[int]) -> List[Tuple[int, int]]:
    is_autosome = (li == autosome_index)
    is_x = (li in x_indices)
    if sex == "male" and is_x:
        return pair_options_maleX_2n2()
    # female X or any autosome
    if is_autosome or (sex == "female" and is_x):
        return pair_options_autosome_2n2()
    # shouldn't happen
    return [(0, 0)]

def possible_transmit_for_pair_2n2(p: Tuple[int, int]) -> List[bool]:
    # for (1,1) -> can transmit upper or lower
    if p == (1, 1):
        return [True, False]
    if p == (2, 0):
        return [True]
    if p == (0, 2):
        return [False]
    return [True]

def build_candidate_rows_from_params(autosome_index: int,
                                    x_indices: List[int],
                                    male_2n2_pairs: Dict[int, Tuple[int, int]],
                                    female_2n2_pairs: Dict[int, Tuple[int, int]],
                                    male_n_state: str,
                                    female_n_state: str,
                                    male_gamete: str,
                                    male_n_transmit: Dict[int, bool],
                                    female_n_transmit: Dict[int, bool],
                                    row_assignment: Dict[str, Tuple[str, str]]) -> Tuple[Dict[str, Dict[str, int]], Dict[str, Dict[str, str]]]:
    """
    row_assignment: row -> ("male"/"female", "2n2"/"n")
    Build full alleles dict for each row.
    """
    full_by_row = {}
    meta_by_row = {}

    # build male n pairs
    male_n2_pairs = build_n2_pairs_from_2n2_pairs_for_person(
        sex="male",
        autosome_index=autosome_index,
        x_indices=x_indices,
        pairs_2n2=male_2n2_pairs,
        transmit=male_n_transmit,
        male_gamete=male_gamete,
    )
    male_pairs_n = male_n2_pairs if male_n_state == "n2" else n_pairs_to_n1_pairs(male_n2_pairs)

    # build female n pairs
    female_n2_pairs = build_n2_pairs_from_2n2_pairs_for_person(
        sex="female",
        autosome_index=autosome_index,
        x_indices=x_indices,
        pairs_2n2=female_2n2_pairs,
        transmit=female_n_transmit,
        male_gamete="X",
    )
    female_pairs_n = female_n2_pairs if female_n_state == "n2" else n_pairs_to_n1_pairs(female_n2_pairs)

    for r in ROW_LABELS:
        sex, kind = row_assignment[r]  # kind: "2n2" or "n"
        if sex == "male":
            pairs = male_2n2_pairs if kind == "2n2" else male_pairs_n
            alleles = pairs_to_alleles(pairs)
            full_by_row[r] = alleles
            meta_by_row[r] = {
                "person": "I",
                "sex": "남성",
                "state": "2n(2)" if kind == "2n2" else ("n(2)" if male_n_state == "n2" else "n(1)")
            }
        else:
            pairs = female_2n2_pairs if kind == "2n2" else female_pairs_n
            alleles = pairs_to_alleles(pairs)
            full_by_row[r] = alleles
            meta_by_row[r] = {
                "person": "II",
                "sex": "여성",
                "state": "2n(2)" if kind == "2n2" else ("n(2)" if female_n_state == "n2" else "n(1)")
            }

    return full_by_row, meta_by_row

def enumerate_row_assignments():
    """
    Row assignment must satisfy:
      - exactly 2 rows are male (I), 2 rows are female (II)
      - for each sex: one is 2n2, one is n
    So rows partitioned into: male_2n2, male_n, female_2n2, female_n
    """
    rows = ROW_LABELS[:]
    for male_2n2 in rows:
        rem1 = [r for r in rows if r != male_2n2]
        for male_n in rem1:
            rem2 = [r for r in rem1 if r != male_n]
            for female_2n2 in rem2:
                female_n = [r for r in rem2 if r != female_2n2][0]
                ra = {}
                ra[male_2n2] = ("male", "2n2")
                ra[male_n] = ("male", "n")
                ra[female_2n2] = ("female", "2n2")
                ra[female_n] = ("female", "n")
                yield ra

def solve_all(masked: Dict[str, Dict[str, Any]]) -> List[World]:
    sols: List[World] = []

    for autosome_index, x_indices in choose_autosome_and_x():
        # enumerate genotypes at 2n2 for male/female
        # male: autosome locus uses autosome options, X loci use maleX options
        # female: autosome locus + X loci all use autosome-like options
        male_pairs_choices = {}
        female_pairs_choices = {}
        for li in [0, 1, 2]:
            male_pairs_choices[li] = genotype_choices_for_2n2("male", li, autosome_index, x_indices)
            female_pairs_choices[li] = genotype_choices_for_2n2("female", li, autosome_index, x_indices)

        for m0 in male_pairs_choices[0]:
            for m1 in male_pairs_choices[1]:
                for m2 in male_pairs_choices[2]:
                    male_2n2_pairs = {0: m0, 1: m1, 2: m2}

                    # male transmit choices only needed for autosome locus when hetero
                    male_trans_choices = {autosome_index: possible_transmit_for_pair_2n2(male_2n2_pairs[autosome_index])}

                    for f0 in female_pairs_choices[0]:
                        for f1 in female_pairs_choices[1]:
                            for f2 in female_pairs_choices[2]:
                                female_2n2_pairs = {0: f0, 1: f1, 2: f2}

                                # female transmit choices for autosome + X loci where hetero
                                female_trans_choices = {}
                                for li in [0, 1, 2]:
                                    if (li == autosome_index) or (li in x_indices):
                                        female_trans_choices[li] = possible_transmit_for_pair_2n2(female_2n2_pairs[li])

                                for male_n_state in ["n2", "n1"]:
                                    for female_n_state in ["n2", "n1"]:
                                        for male_gamete in ["X", "Y"]:
                                            # build transmit loops
                                            for m_tr in male_trans_choices[autosome_index]:
                                                male_n_transmit = {autosome_index: m_tr}

                                                # female transmit cartesian
                                                def iter_f_trans():
                                                    keys = list(female_trans_choices.keys())
                                                    def rec(i, cur):
                                                        if i == len(keys):
                                                            yield dict(cur)
                                                            return
                                                        k = keys[i]
                                                        for v in female_trans_choices[k]:
                                                            cur[k] = v
                                                            yield from rec(i+1, cur)
                                                    yield from rec(0, {})
                                                for female_n_transmit in iter_f_trans():
                                                    # fill missing keys with defaults (for safety)
                                                    if autosome_index not in female_n_transmit:
                                                        female_n_transmit[autosome_index] = True

                                                    for row_assignment in enumerate_row_assignments():
                                                        # build row alleles & sums
                                                        full_by_row, meta_by_row = build_candidate_rows_from_params(
                                                            autosome_index=autosome_index,
                                                            x_indices=x_indices,
                                                            male_2n2_pairs=male_2n2_pairs,
                                                            female_2n2_pairs=female_2n2_pairs,
                                                            male_n_state=male_n_state,
                                                            female_n_state=female_n_state,
                                                            male_gamete=male_gamete,
                                                            male_n_transmit=male_n_transmit,
                                                            female_n_transmit=female_n_transmit,
                                                            row_assignment=row_assignment,
                                                        )

                                                        sums_full = {r: compute_sums_from_alleles(full_by_row[r]) for r in ROW_LABELS}
                                                        if not masked_match(masked, sums_full):
                                                            continue

                                                        # masked can be explained by this world
                                                        # store candidate world
                                                        sols.append(World(
                                                            autosome_index=autosome_index,
                                                            x_indices=x_indices,
                                                            male_n_state=male_n_state,
                                                            female_n_state=female_n_state,
                                                            male_gamete=male_gamete,
                                                            male_2n2_pairs=male_2n2_pairs,
                                                            female_2n2_pairs=female_2n2_pairs,
                                                            male_n_transmit=male_n_transmit,
                                                            female_n_transmit=female_n_transmit,
                                                            full_alleles_by_row=full_by_row,
                                                            row_meta=meta_by_row,
                                                            sums_full=sums_full,
                                                            sums_masked=masked,
                                                        ))

                                                        # early stop: if too many solutions, we can cut
                                                        if len(sols) > 2:
                                                            return sols
    return sols


# =========================
# Generate a unique problem
# =========================
def loci_name(li: int) -> str:
    return ["(A/a)", "(B/b)", "(D/d)"][li]

def chromosome_location_text(w: World) -> str:
    # two X loci and one autosome
    auto = w.autosome_index
    xs = w.x_indices
    parts = []
    parts.append(f"{loci_name(auto)}는 상염색체에 존재한다.")
    parts.append(f"{loci_name(xs[0])}, {loci_name(xs[1])}는 X염색체에 존재한다.")
    return " ".join(parts)

def clue_sentences_from_known(masked: Dict[str, Dict[str, Any]]) -> List[str]:
    # "결정적 단서" 자동문장: 값이 0인 항목이 공개되어 있으면 3개 항이 모두 0임을 말해줌
    clues = []
    for r in ROW_LABELS:
        for c in COL_LABELS:
            v = masked[r][c]
            if v == "?":
                continue
            if int(v) == 0:
                if c == "A+B+D":
                    clues.append(f"({r})에서 A+B+D=0이므로 A=B=D=0이다.")
                elif c == "a+b+D":
                    clues.append(f"({r})에서 a+b+D=0이므로 a=b=D=0이다.")
                else:  # a+B+d
                    clues.append(f"({r})에서 a+B+d=0이므로 a=B=d=0이다.")
    # 중복 제거
    out = []
    seen = set()
    for s in clues:
        if s not in seen:
            seen.add(s)
            out.append(s)
    return out[:5]  # 너무 길어지지 않게


def generate_unique_problem(pnum: int) -> World:
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        # choose autosome locus (then 2 X loci fixed)
        autosome_index, x_indices = random.choice(choose_autosome_and_x())

        # choose 2n2 genotypes (as pair amounts) for male/female
        male_2n2_pairs = {}
        female_2n2_pairs = {}
        for li in [0, 1, 2]:
            male_2n2_pairs[li] = random.choice(genotype_choices_for_2n2("male", li, autosome_index, x_indices))
            female_2n2_pairs[li] = random.choice(genotype_choices_for_2n2("female", li, autosome_index, x_indices))

        # choose n states
        male_n_state = random.choice(["n2", "n1"])
        female_n_state = random.choice(["n2", "n1"])

        # male gamete type for n cell
        male_gamete = random.choice(["X", "Y"])

        # transmit choices
        male_n_transmit = {}
        # autosome locus only
        auto_pair = male_2n2_pairs[autosome_index]
        male_n_transmit[autosome_index] = random.choice(possible_transmit_for_pair_2n2(auto_pair))

        female_n_transmit = {}
        for li in [0, 1, 2]:
            if (li == autosome_index) or (li in x_indices):
                female_n_transmit[li] = random.choice(possible_transmit_for_pair_2n2(female_2n2_pairs[li]))

        # assign rows to 4 cell types randomly (subject to 2 each sex + 2n2/n per sex)
        row_assignment = random.choice(list(enumerate_row_assignments()))

        # build full alleles
        full_by_row, meta_by_row = build_candidate_rows_from_params(
            autosome_index=autosome_index,
            x_indices=x_indices,
            male_2n2_pairs=male_2n2_pairs,
            female_2n2_pairs=female_2n2_pairs,
            male_n_state=male_n_state,
            female_n_state=female_n_state,
            male_gamete=male_gamete,
            male_n_transmit=male_n_transmit,
            female_n_transmit=female_n_transmit,
            row_assignment=row_assignment,
        )

        sums_full = {r: compute_sums_from_alleles(full_by_row[r]) for r in ROW_LABELS}

        # mask
        masked = mask_table(sums_full)

        # solve
        sols = solve_all(masked)
        if len(sols) == 1:
            # unique!
            w = sols[0]
            # overwrite with our masked (solver already matches it)
            w.sums_masked = masked
            # and keep full sums (from the unique solution)
            return w

        if attempt % LOG_EVERY == 0:
            print(f"[P{pnum:02d}] attempt={attempt}, sols={len(sols)} (searching...)")

    raise RuntimeError(f"[P{pnum:02d}] 유일정답 생성 실패: MAX_TRIES_PER_PROBLEM 증가 또는 MASK 설정 조정 필요")


# =========================
# DOCX output
# =========================
def set_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)

def add_problem_to_cell(cell, pnum: int, w: World):
    p = cell.add_paragraph(f"[문제 {pnum}]")
    if p.runs:
        p.runs[0].bold = True

    cell.add_paragraph("아래의 표는 세포 (가), (나), (다), (라)에 있는 유전자 A, a, B, b, D, d 중 3개의 DNA의 양의 합을 나타낸 것이다.")
    cell.add_paragraph("세포 (가), (나), (다), (라) 중 2개는 남성 I의 세포이고 나머지 2개는 여성 II의 세포이다.")
    cell.add_paragraph("단, (A/a), (B/b), (D/d)는 각각 대립유전자이고, 세 유전자쌍 중 2가지는 X염색체에, 나머지 1가지는 상염색체에 있다.")
    cell.add_paragraph("")

    # table: rows=5, cols=4 (row label + 3 columns)
    t = cell.add_table(rows=1 + len(ROW_LABELS), cols=1 + len(COL_LABELS))
    t.style = "Table Grid"
    t.cell(0, 0).text = ""
    for j, c in enumerate(COL_LABELS, start=1):
        t.cell(0, j).text = c

    for i, r in enumerate(ROW_LABELS, start=1):
        t.cell(i, 0).text = f"({r})"
        for j, c in enumerate(COL_LABELS, start=1):
            t.cell(i, j).text = str(w.sums_masked[r][c])

    cell.add_paragraph("")
    cell.add_paragraph("표에서 각 세포의 성별(I/II)과 상태(2n(2), n(2), n(1))를 찾고,")
    cell.add_paragraph("A, a, B, b, D, d의 DNA 상대량을 구한 다음, (A/a), (B/b), (D/d)가 있는 염색체(X/상염색체)를 찾으시오.")

def add_two_column_pages(doc: Document, worlds: List[World]):
    idx = 0
    pnum = 1
    while idx < len(worlds):
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        add_problem_to_cell(left, pnum, worlds[idx])
        idx += 1
        pnum += 1

        if idx < len(worlds):
            add_problem_to_cell(right, pnum, worlds[idx])
            idx += 1
            pnum += 1

        if idx < len(worlds):
            doc.add_page_break()

def add_solution_section(doc: Document, worlds: List[World]):
    doc.add_page_break()
    h = doc.add_paragraph("[정답·해설]")
    if h.runs:
        h.runs[0].bold = True

    for i, w in enumerate(worlds, start=1):
        title = doc.add_paragraph(f"{i}번")
        if title.runs:
            title.runs[0].bold = True

        # 1) row meta
        doc.add_paragraph("① 각 세포의 성별/상태")
        for r in ROW_LABELS:
            m = w.row_meta[r]
            doc.add_paragraph(f" - ({r}): {m['person']}({m['sex']}), {m['state']}")

        # 2) ? 채운 부분합 표
        doc.add_paragraph("② ? 채운 표(부분합)")
        t = doc.add_table(rows=1 + len(ROW_LABELS), cols=1 + len(COL_LABELS))
        t.style = "Table Grid"
        t.cell(0, 0).text = ""
        for j, c in enumerate(COL_LABELS, start=1):
            t.cell(0, j).text = c
        for rr, r in enumerate(ROW_LABELS, start=1):
            t.cell(rr, 0).text = f"({r})"
            for j, c in enumerate(COL_LABELS, start=1):
                t.cell(rr, j).text = str(w.sums_full[r][c])

        # 3) 완성표(A,a,B,b,D,d)
        doc.add_paragraph("③ A, a, B, b, D, d의 DNA 상대량(완성표)")
        alleles = ["A", "a", "B", "b", "D", "d"]
        t2 = doc.add_table(rows=1 + len(ROW_LABELS), cols=1 + len(alleles))
        t2.style = "Table Grid"
        t2.cell(0, 0).text = ""
        for j, a in enumerate(alleles, start=1):
            t2.cell(0, j).text = a
        for rr, r in enumerate(ROW_LABELS, start=1):
            t2.cell(rr, 0).text = f"({r})"
            for j, a in enumerate(alleles, start=1):
                t2.cell(rr, j).text = str(w.full_alleles_by_row[r][a])

        # 4) 염색체 위치
        doc.add_paragraph("④ 염색체 위치")
        doc.add_paragraph(" - " + chromosome_location_text(w))

        # 5) 결정적 단서 문장
        clues = clue_sentences_from_known(w.sums_masked)
        if clues:
            doc.add_paragraph("⑤ 결정적 단서")
            for s in clues:
                doc.add_paragraph(" - " + s)

        doc.add_paragraph("")


def make_output_path() -> str:
    os.makedirs(OUT_DIR, exist_ok=True)
    ts = int(time.time())
    return os.path.join(OUT_DIR, f"{OUT_FILE_PREFIX}_{ts}.docx")



def sums_to_md(sums: Dict[str, Dict[str, Any]], rows: List[str], cols: List[str]) -> str:
    header = [""] + [str(c) for c in cols]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for r in rows:
        row = [f"({r})"] + [str(sums[r][c]) for c in cols]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

def count_masks_in_sums(sums: Dict[str, Dict[str, Any]]) -> int:
    n = 0
    for r in sums:
        for c in sums[r]:
            if sums[r][c] == "?":
                n += 1
    return n

def main():
    out_path = make_output_path()
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(out_dir, exist_ok=True)
    pack_writer = ProblemPack(module_code="DNAIND", out_dir=out_dir, id_prefix="DNAIND_")


    doc = Document()
    set_doc_style(doc)

    worlds: List[World] = []
    for pnum in range(1, N_PROBLEMS + 1):
        w = generate_unique_problem(pnum)
        worlds.append(w)
        # PACK record (표시용 텍스트/표/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "아래의 표는 세포 (가), (나), (다), (라)에 있는 유전자 A, a, B, b, D, d 중 3개의 DNA 양의 합을 나타낸 것이다.",
            "가, 나, 다, 라 중 2개는 남성 I의 세포이고 나머지 2개는 여성 II의 세포이다.",
            "단, (A/a), (B/b), (D/d) 중 2가지는 X염색체에 있고 나머지 1가지는 상염색체에 있다.",
        ])
        ask_line_md = "표에서 각 세포의 성별과 상태를 찾고, A, a, B, b, D, d의 DNA 상대량을 구하고 염색체 위치를 찾으시오."

        table_md = sums_to_md(w.sums_masked, ROW_LABELS, COL_LABELS)
        full_table_md = sums_to_md(w.sums_full, ROW_LABELS, COL_LABELS)

        answer_md = "\n".join([
            "① 각 세포의 성별/상태: " + "; ".join([f"({r})={w.row_meta[r]['person']}({w.row_meta[r]['sex']}),{w.row_meta[r]['state']}" for r in ROW_LABELS]),
            "② 염색체 위치: " + chromosome_location_text(w),
        ])
        clues = clue_sentences_from_known(w.sums_masked)
        explanation_md = "\n".join(["- " + s for s in clues]) if clues else "(결정적 단서 자동생성 없음)"

        pack_writer.new_problem(qnum=pnum, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": table_md,
            "full_table_md": full_table_md,
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "mask_used": count_masks_in_sums(w.sums_masked),
            "data": {
                "world": w,
            },
        }))
        print(f"[진행] {pnum}/{N_PROBLEMS} 생성 완료")

    add_two_column_pages(doc, worlds)
    add_solution_section(doc, worlds)

    doc.save(out_path)
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {out_path}")


if __name__ == "__main__":
    main()