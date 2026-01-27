
# problem_bank_web_REWRITE_v1.py
# ------------------------------------------------------------
# Streamlit 문제은행 웹 (표준화(normalize) 기반, UI 안정화)
#
# 실행:
#   pip install streamlit python-docx pandas openpyxl
#   streamlit run problem_bank_web.py
#
# 폴더 구조(권장):
#   프로젝트/
#     problem_bank_web.py
#     output/               <- 각 생성기가 만든 PACK JSON들이 모이는 폴더(재귀 탐색)
#     tree_base_A.png (선택)
#     tree_base_B.png (선택)
#
# 기능:
# - PACK JSON 자동 탐색/로드(깨진 JSON 스킵)
# - payload 자동 표준화: problem_text_md, ask_line_md, given_table, full_table, answer_md, explain_md
# - ID prefix/모듈 그룹 접기/펼치기, 현재 필터 기준 전체선택/해제
# - 선택한 것만 보기, 선택 표시 유지
# - 난이도 드롭다운(로컬 저장: difficulty_overrides.json)
# - DOCX 다운로드(2단/표는 "진짜 표"로 삽입, 정답/해설은 뒤에 모아서)
# ------------------------------------------------------------

from __future__ import annotations
import os
import re
import io
import json
import base64
import hashlib
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st
import pandas as pd


# ============================================================
# Fallback text extractors (문제 본문/요구사항이 PACK에 없을 때 표시용)
# ============================================================

FALLBACK_PROBLEM_TEXT_KEYS = [
    "problem_text_md", "problem_text", "problem_md", "qtext", "stem", "prompt", "question", "desc", "description",
    "problem", "text"
]
FALLBACK_ASK_TEXT_KEYS = [
    "ask_line_md", "ask_line", "ask_md", "ask", "task", "requirement", "requirements", "what_to_do", "query"
]

def _deep_find_first(obj, keys):
    """Recursively find first value for any key in keys within dict/list structures."""
    if isinstance(obj, dict):
        # direct hit
        for k in keys:
            if k in obj and obj[k] not in (None, "", [], {}):
                return obj[k]
        # recurse
        for v in obj.values():
            got = _deep_find_first(v, keys)
            if got not in (None, "", [], {}):
                return got
    elif isinstance(obj, list):
        for it in obj:
            got = _deep_find_first(it, keys)
            if got not in (None, "", [], {}):
                return got
    return None

def _as_md_text(v):
    """Convert arbitrary payload value to displayable markdown-ish text (safe, compact)."""
    if v is None:
        return ""
    if isinstance(v, (str, int, float, bool)):
        return str(v)
    # pretty JSON for dict/list
    try:
        return json.dumps(v, ensure_ascii=False, indent=2)
    except Exception:
        return str(v)

def get_display_texts(payload: dict):
    """
    Returns (problem_text_md, ask_line_md).
    If pack doesn't contain them, fallback searches payload recursively for reasonable keys.
    """
    # First, try existing canonical fields from get_fields (if present)
    try:
        ptxt, atxt, ans, expl = get_fields(payload)  # type: ignore
        if ptxt or atxt:
            return (ptxt or "", atxt or "")
    except Exception:
        pass

    p = _deep_find_first(payload, FALLBACK_PROBLEM_TEXT_KEYS)
    a = _deep_find_first(payload, FALLBACK_ASK_TEXT_KEYS)

    ptxt = _as_md_text(p) if p is not None else ""
    atxt = _as_md_text(a) if a is not None else ""

    return (ptxt, atxt)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

# =========================
# CONFIG
# =========================
APP_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_PACK_ROOT = os.path.join(APP_DIR, "output")  # 재귀 탐색
DIFF_OVERRIDES_PATH = os.path.join(APP_DIR, "difficulty_overrides.json")

# 표준 키 후보(모듈별 차이를 흡수)
PROBLEM_TEXT_KEYS = [
    "problem_text_md", "problem_md", "problem_text", "question_md", "question",
    "stem", "prompt", "text_md", "text", "problem"
]
ASK_LINE_KEYS = [
    "ask_line_md", "ask_md", "ask", "task", "requirement", "requirements",
    "what_to_do", "query"
]
ANSWER_KEYS = ["answer_md", "answer", "ans", "correct_answer", "solution_answer"]
EXPLAIN_KEYS = ["explain_md", "explanation_md", "explanation", "explain", "reasoning", "reasons"]
GIVEN_TABLE_KEYS = [
    "given_table", "masked_table", "table_given", "table", "grid", "masked", "given",
    "presented_table", "question_table"
]
FULL_TABLE_KEYS = [
    "full_table", "complete_table", "solution_table", "answer_table", "table_full",
    "filled_table"
]
# 이미지 키(있으면 첨부)
IMAGE_KEYS = ["image_path", "figure_path", "fig_path", "png_path", "tree_png_path", "img_path", "image_b64", "img_b64"]

DIFFICULTY_LEVELS = ["미분류", "하", "중", "상", "극상"]

# =========================
# DATA MODEL
# =========================
@dataclass
class ProblemItem:
    pid: str
    module: str
    prefix: str
    payload: Dict[str, Any]
    norm: Dict[str, Any]
    source_file: str
    uid: str = ""

    def __post_init__(self):
        # uid: 화면 렌더링/체크박스 key 충돌 방지용 내부 고유키
        if not self.uid:
            base = f"{self.source_file}::{self.pid}"
            self.uid = hashlib.sha1(base.encode("utf-8")).hexdigest()[:16]

# =========================
# UTILS
# =========================
def safe_str(x: Any) -> str:
    try:
        return json.dumps(x, ensure_ascii=False, indent=2)
    except Exception:
        try:
            return str(x)
        except Exception:
            return "<unprintable>"

def find_first(d: Dict[str, Any], keys: List[str]) -> Any:
    for k in keys:
        if k in d and d[k] is not None and d[k] != "":
            return d[k]
    return None

def parse_obj_maybe(x: Any) -> Any:
    """문자열로 들어온 dict/list를 실제 객체로 변환."""
    if isinstance(x, str):
        s = x.strip()
        # 너무 긴 문자열 eval 방지용(그래도 필요하면 늘려)
        if len(s) > 200000:
            return x
        # json 먼저
        if (s.startswith("{") and s.endswith("}")) or (s.startswith("[") and s.endswith("]")):
            try:
                return json.loads(s)
            except Exception:
                pass
            # python dict repr
            try:
                return eval(s, {"__builtins__": {}})
            except Exception:
                return x
        return x
    return x

def infer_prefix(pid: str) -> str:
    # 1) ABCD_xxx
    if "_" in pid:
        return pid.split("_", 1)[0]
    # 2) ABCD-2026...
    if "-" in pid:
        return pid.split("-", 1)[0]
    # 3) fallback
    m = re.match(r"([A-Za-z0-9]+)", pid)
    return m.group(1) if m else "UNKNOWN"

def dict_table_to_df(obj: Any) -> Optional[pd.DataFrame]:
    obj = parse_obj_maybe(obj)
    if obj is None:
        return None
    if isinstance(obj, pd.DataFrame):
        return obj
    # list-of-list
    if isinstance(obj, list):
        if len(obj) == 0:
            return pd.DataFrame()
        if all(isinstance(r, list) for r in obj):
            return pd.DataFrame(obj)
        # list of dict
        if all(isinstance(r, dict) for r in obj):
            return pd.DataFrame(obj)
        return None
    # dict-of-dict
    if isinstance(obj, dict):
        # if values are dict -> columns keyed by outer keys
        if all(isinstance(v, dict) for v in obj.values()):
            # try stable order
            outer_keys = list(obj.keys())
            inner_keys = set()
            for v in obj.values():
                inner_keys |= set(v.keys())
            inner_keys = list(sorted(inner_keys, key=lambda x: (str(x))))
            data = []
            for ik in inner_keys:
                row = []
                for ok in outer_keys:
                    row.append(obj.get(ok, {}).get(ik, ""))
                data.append(row)
            df = pd.DataFrame(data, columns=[str(k) for k in outer_keys])
            df.insert(0, "row", [str(k) for k in inner_keys])
            return df
        # plain dict
        return pd.DataFrame([obj])
    return None

def safe_dataframe(df: pd.DataFrame):
    """pyarrow 오류(리스트/스칼라 혼합) 방지: 전부 문자열화"""
    def norm_cell(v):
        if isinstance(v, (list, tuple, dict)):
            return safe_str(v)
        return v
    df2 = df.copy()
    for c in df2.columns:
        df2[c] = df2[c].map(norm_cell)
    st.dataframe(df2, use_container_width=True)

def load_diff_overrides() -> Dict[str, str]:
    if os.path.exists(DIFF_OVERRIDES_PATH):
        try:
            with open(DIFF_OVERRIDES_PATH, "r", encoding="utf-8") as f:
                d = json.load(f)
            if isinstance(d, dict):
                return {str(k): str(v) for k, v in d.items()}
        except Exception:
            pass
    return {}

def save_diff_overrides(d: Dict[str, str]) -> None:
    try:
        with open(DIFF_OVERRIDES_PATH, "w", encoding="utf-8") as f:
            json.dump(d, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.warning(f"난이도 저장 실패: {e}")

def normalize_payload(payload: Dict[str, Any], pid: str, module: str) -> Dict[str, Any]:
    norm: Dict[str, Any] = {}
    norm["id"] = pid
    norm["module"] = module
    norm["prefix"] = infer_prefix(pid)

    # 본문/요구
    norm["problem_text_md"] = find_first(payload, PROBLEM_TEXT_KEYS) or ""
    norm["ask_line_md"] = find_first(payload, ASK_LINE_KEYS) or ""

    # 표
    norm["given_table"] = parse_obj_maybe(find_first(payload, GIVEN_TABLE_KEYS))
    norm["full_table"] = parse_obj_maybe(find_first(payload, FULL_TABLE_KEYS))

    # 정답/해설
    norm["answer_md"] = find_first(payload, ANSWER_KEYS) or ""
    norm["explain_md"] = find_first(payload, EXPLAIN_KEYS) or ""

    # 난이도
    norm["difficulty"] = payload.get("difficulty", "") or ""
    # 태그
    tags = payload.get("tags", [])
    norm["tags"] = tags if isinstance(tags, list) else []

    # 이미지
    img = find_first(payload, IMAGE_KEYS)
    norm["image"] = img

    # 요약용
    norm["payload_keys"] = list(payload.keys())
    return norm


def make_uid(source_path: str, pid: str) -> str:
    """Stable unique key for Streamlit widgets (avoids duplicate element keys)."""
    h = hashlib.sha1(f"{source_path}|{pid}".encode("utf-8")).hexdigest()
    return h[:12]
# =========================
# PACK LOADING
# =========================

def dedupe_items_by_uid(items: List[ProblemItem]) -> List[ProblemItem]:
    """Streamlit 위젯 key 충돌 방지: 같은 uid(=같은 문제)가 여러 번 로드되면 1개만 남김."""
    seen = set()
    out: List[ProblemItem] = []
    for it in items:
        if it.uid in seen:
            continue
        seen.add(it.uid)
        out.append(it)
    return out

def iter_json_files(root: str) -> List[str]:
    out = []
    if not root or not os.path.exists(root):
        return out
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            if fn.lower().endswith(".json"):
                out.append(os.path.join(dirpath, fn))
    out.sort()
    return out

def load_one_pack_file(path: str) -> List[ProblemItem]:
    items: List[ProblemItem] = []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return items  # 깨진 json은 스킵

    module = "UNKNOWN"
    raw_items = None

    if isinstance(data, dict):
        module = str(data.get("module_code") or data.get("module") or data.get("moduleCode") or "UNKNOWN")
        raw_items = data.get("items") or data.get("problems") or data.get("data")
        # 어떤 생성기는 {"items":[{"pid":..,"payload":..}]} 형태
    elif isinstance(data, list):
        raw_items = data
    else:
        return items

    if raw_items is None:
        return items
    if not isinstance(raw_items, list):
        return items

    for it in raw_items:
        if not isinstance(it, dict):
            continue
        pid = str(it.get("id") or it.get("pid") or it.get("problem_id") or it.get("problemId") or "")
        payload = it.get("payload") if isinstance(it.get("payload"), dict) else None
        if not payload:
            # 어떤 파일은 it 자체가 payload일 수 있음
            payload = {k: v for k, v in it.items() if k not in ("id", "pid", "problem_id", "problemId")}
        if not pid:
            # 해시 기반 임시 id
            h = hashlib.sha1(safe_str(payload).encode("utf-8")).hexdigest()[:10]
            pid = f"{module}_{h}"
        prefix = infer_prefix(pid)
        norm = normalize_payload(payload, pid, module)
        items.append(ProblemItem(pid=pid, module=module, prefix=prefix, payload=payload, norm=norm, source_file=path))

    return items

@st.cache_data(show_spinner=False)
def load_all_packs(pack_root: str, max_files: int = 5000) -> Tuple[List[ProblemItem], Dict[str, Any]]:
    files = iter_json_files(pack_root)[:max_files]
    all_items: List[ProblemItem] = []
    bad = 0
    for fp in files:
        got = load_one_pack_file(fp)
        if not got:
            bad += 1
        all_items.extend(got)
    meta = {"files": len(files), "bad_or_empty": bad, "items": len(all_items)}
    return all_items, meta

# =========================
# RENDER (Streamlit)
# =========================
def render_markdown_block(title: str, md: str):
    if not md:
        return
    st.markdown(f"**{title}**")
    st.markdown(md)

def render_table_block(title: str, tbl_obj: Any):
    st.markdown(f"**{title}**")
    if tbl_obj is None or tbl_obj == "":
        st.info("표 없음")
        return
    df = dict_table_to_df(tbl_obj)
    if df is not None:
        safe_dataframe(df)
    else:
        st.code(safe_str(tbl_obj))

def try_resolve_image(norm: Dict[str, Any], source_file: str) -> Optional[bytes]:
    img = norm.get("image")
    if not img:
        return None

    # base64
    if isinstance(img, str) and (img.startswith("data:image") or len(img) > 2000):
        # data url
        if img.startswith("data:image"):
            try:
                b64 = img.split(",", 1)[1]
                return base64.b64decode(b64)
            except Exception:
                return None
        # raw b64
        try:
            return base64.b64decode(img)
        except Exception:
            return None

    # path
    if isinstance(img, str):
        cand = []
        # 1) absolute or relative to app dir
        cand.append(img)
        cand.append(os.path.join(APP_DIR, img))
        # 2) relative to source pack dir
        cand.append(os.path.join(os.path.dirname(source_file), img))
        for p in cand:
            if os.path.exists(p) and os.path.isfile(p):
                try:
                    with open(p, "rb") as f:
                        return f.read()
                except Exception:
                    pass
    return None

# =========================
# DOCX EXPORT
# =========================
def set_doc_font(doc: Document, font_name: str = "바탕", font_pt: int = 9):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_pt)

def add_par(container, text: str, bold: bool = False):
    p = container.add_paragraph(text)
    if p.runs:
        p.runs[0].bold = bold
    return p

def _set_cell_text(cell, text: str, bold: bool = False, center: bool = True):
    # Clear existing paragraphs
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "바탕"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass


def add_docx_table(container, table_obj: Any):
    """
    table_obj가 dict/list 형태(=표)면 '진짜 표'로 예쁘게 출력.
    - 2단 레이아웃에서도 깨지지 않게 폭을 균등 분배
    - 글자 9pt, 가운데 정렬(기본), 줄간격/여백 최소화
    """
    df = dict_table_to_df(table_obj)
    if df is None:
        # fallback: text
        add_par(container, safe_str(table_obj))
        return

    rows = df.shape[0] + 1
    cols = df.shape[1]
    t = container.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    try:
        t.autofit = False
    except Exception:
        pass

    # 2단 셀 안에서 너무 넓게 잡히지 않도록, '셀 기준'으로 적당한 총 폭을 가정해 균등 분배
    # (한 단 폭은 대략 3.1~3.3 inch 정도라서 3.2로 고정)
    total_w = 3.2
    col_w = max(0.35, total_w / max(1, cols))  # 너무 얇아지는 것 방지

    # header
    for j, col in enumerate(df.columns):
        c = t.cell(0, j)
        _set_cell_text(c, str(col), bold=True, center=True)
        try:
            c.width = Inches(col_w)
        except Exception:
            pass

    # body
    for i in range(df.shape[0]):
        for j in range(cols):
            val = df.iat[i, j]
            if isinstance(val, (list, dict, tuple)):
                val = safe_str(val)
            text = "" if val is None else str(val)
            c = t.cell(i + 1, j)
            _set_cell_text(c, text, bold=False, center=True)
            try:
                c.width = Inches(col_w)
            except Exception:
                pass


def export_docx_bytes(selected: List[ProblemItem], include_expl: bool, include_full: bool) -> bytes:
    doc = Document()
    set_doc_font(doc, "바탕", 9)

    # 문제 파트: 2단(페이지당 2문항)
    idx = 0
    pnum = 1
    while idx < len(selected):
        outer = doc.add_table(rows=1, cols=2)
        left = outer.cell(0, 0)
        right = outer.cell(0, 1)

        def fill(cell, it: ProblemItem, num: int):
            n = it.norm
            add_par(cell, f"[문제 {num}]  ID: {it.pid}  ({it.prefix})", bold=True)

            if n.get("problem_text_md"):
                add_par(cell, "문제", bold=True)
                add_par(cell, n["problem_text_md"])
            if n.get("ask_line_md"):
                add_par(cell, "요구사항", bold=True)
                add_par(cell, n["ask_line_md"])

            # 이미지(있으면)
            img_bytes = try_resolve_image(n, it.source_file)
            if img_bytes:
                try:
                    cell.add_paragraph("")
                    run = cell.add_paragraph().add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(2.2))
                except Exception:
                    pass

            # 제시표
            add_par(cell, "제시표", bold=True)
            given = n.get("given_table")
            if given is None or given == "":
                add_par(cell, "(표 없음)")
            else:
                add_docx_table(cell, given)
            cell.add_paragraph("")

        fill(left, selected[idx], pnum)
        idx += 1
        pnum += 1

        if idx < len(selected):
            fill(right, selected[idx], pnum)
            idx += 1
            pnum += 1

        if idx < len(selected):
            doc.add_page_break()

    # 정답/해설 파트
    doc.add_page_break()
    add_par(doc, "[정답/해설]", bold=True)

    for i, it in enumerate(selected, start=1):
        n = it.norm
        add_par(doc, f"{i}. ID: {it.pid}  ({it.prefix})", bold=True)

        ans = n.get("answer_md") or "(없음)"
        add_par(doc, "정답", bold=True)
        add_par(doc, ans)

        if include_full:
            ft = n.get("full_table")
            if ft not in (None, ""):
                add_par(doc, "완성표", bold=True)
                add_docx_table(doc, ft)

        if include_expl:
            expl = n.get("explain_md") or "(없음)"
            add_par(doc, "해설", bold=True)
            add_par(doc, expl)

        add_par(doc, "-" * 40)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# =========================
# MAIN UI
# =========================
def main():
    st.set_page_config(page_title="문제은행", layout="wide")

    st.title("문제은행 웹 (안정화 리라이트)")

    # Sidebar: load
    st.sidebar.header("로드")
    pack_root = st.sidebar.text_input("PACK 폴더", DEFAULT_PACK_ROOT)
    max_files = st.sidebar.number_input("최대 JSON 파일 수", min_value=10, max_value=20000, value=5000, step=10)

    if st.sidebar.button("새로고침(재로드)", use_container_width=True):
        st.cache_data.clear()

    items, meta = load_all_packs(pack_root, max_files=int(max_files))
    items = dedupe_items_by_uid(items)
    diff_over = load_diff_overrides()

    # inject difficulty overrides
    for it in items:
        if it.pid in diff_over:
            it.norm["difficulty"] = diff_over[it.pid]

    st.sidebar.caption(f"파일 {meta['files']}개 / 빈·깨짐 {meta['bad_or_empty']}개 / 문항 {meta['items']}개")

    if not items:
        st.warning("문항을 찾지 못했습니다. PACK 폴더 경로를 확인하세요.")
        st.stop()

    # Sidebar: filters
    st.sidebar.header("필터")
    all_modules = sorted(set(i.module for i in items))
    all_prefix = sorted(set(i.prefix for i in items))
    all_diff = DIFFICULTY_LEVELS

    f_module = st.sidebar.multiselect("모듈", options=all_modules, default=all_modules)
    f_prefix = st.sidebar.multiselect("ID prefix", options=all_prefix, default=all_prefix)
    f_diff = st.sidebar.multiselect("난이도", options=all_diff, default=all_diff)
    query = st.sidebar.text_input("검색(본문/요구/ID)", "")
    selected_only = st.sidebar.checkbox("선택한 것만 보기", value=False)

    # selection state
    if "selected_uids" not in st.session_state:
        st.session_state.selected_uids = set()

    # filter items
    def match(it: ProblemItem) -> bool:
        if it.module not in f_module:
            return False
        if it.prefix not in f_prefix:
            return False
        d = it.norm.get("difficulty") or "미분류"
        if d not in f_diff:
            return False
        if query:
            q = query.lower()
            blob = (it.pid + " " + (it.norm.get("problem_text_md") or "") + " " + (it.norm.get("ask_line_md") or "")).lower()
            if q not in blob:
                return False
        if selected_only and (it.uid not in st.session_state.selected_uids):
            return False
        return True

    view_items = [it for it in items if match(it)]

    # top controls
    c1, c2, c3, c4 = st.columns([1, 1, 1, 2])
    with c1:
        if st.button("현재 화면 전체 선택"):
            for it in view_items:
                st.session_state.selected_uids.add(it.uid)
    with c2:
        if st.button("현재 화면 전체 해제"):
            for it in view_items:
                st.session_state.selected_uids.discard(it.uid)
    with c3:
        st.write(f"표시 {len(view_items)}개 / 전체 {len(items)}개")
    with c4:
        st.write(f"선택 {len(st.session_state.selected_uids)}개")

    # layout: list + preview
    left, right = st.columns([1.1, 1.4], gap="large")

    # LEFT: grouped list
    with left:
        st.subheader("목록")
        group_mode = st.radio("그룹 기준", ["prefix", "module"], horizontal=True)
        groups: Dict[str, List[ProblemItem]] = {}
        for it in view_items:
            key = it.prefix if group_mode == "prefix" else it.module
            groups.setdefault(key, []).append(it)

        for gkey in sorted(groups.keys()):
            gitems = groups[gkey]
            with st.expander(f"{gkey}  ({len(gitems)}개)", expanded=False):
                # group buttons (current group only)
                bc1, bc2, bc3 = st.columns([1, 1, 2])
                with bc1:
                    if st.button("그룹 선택", key=f"sel_{group_mode}_{gkey}"):
                        for row_i, it in enumerate(gitems):
                            st.session_state.selected_uids.add(it.uid)
                with bc2:
                    if st.button("그룹 해제", key=f"clr_{group_mode}_{gkey}"):
                        for row_i, it in enumerate(gitems):
                            st.session_state.selected_uids.discard(it.uid)
                with bc3:
                    st.caption("체크박스/난이도는 즉시 반영")

                for row_i, it in enumerate(gitems):
                    checked = it.uid in st.session_state.selected_uids
                    row = st.container()
                    cc1, cc2, cc3 = row.columns([0.15, 0.55, 0.3])
                    with cc1:
                        new_checked = st.checkbox("", value=checked, key=f"chk_{it.uid}")
                        if new_checked:
                            st.session_state.selected_uids.add(it.uid)
                        else:
                            st.session_state.selected_uids.discard(it.uid)
                    with cc2:
                        st.write(f"**{it.pid}**")
                        st.caption(os.path.basename(it.source_file))
                    with cc3:
                        cur = it.norm.get("difficulty") or "미분류"
                        new = st.selectbox("난이도", DIFFICULTY_LEVELS, index=DIFFICULTY_LEVELS.index(cur) if cur in DIFFICULTY_LEVELS else 0, key=f"diff_{it.uid}")
                        if new != cur:
                            it.norm["difficulty"] = new
                            diff_over[it.pid] = new
                            save_diff_overrides(diff_over)

        st.divider()
        st.caption("※ JSON이 일부 깨져 있어도 로딩은 계속됩니다(스킵 처리).")

    # RIGHT: preview + export
    with right:
        st.subheader("미리보기 / 내보내기")

        # pick preview item
        if "preview_id" not in st.session_state:
            st.session_state.preview_id = view_items[0].pid if view_items else items[0].pid

        pid_to_item = {it.pid: it for it in view_items} | {it.pid: it for it in items}
        preview_options = [it.pid for it in view_items] if view_items else [it.pid for it in items]
        if st.session_state.preview_id not in preview_options and preview_options:
            st.session_state.preview_id = preview_options[0]

        st.session_state.preview_id = st.selectbox("미리볼 ID", options=preview_options, index=preview_options.index(st.session_state.preview_id))

        it = pid_to_item.get(st.session_state.preview_id)
        if it:
            n = it.norm
            st.markdown(f"### {it.pid}  ({it.prefix} / {it.module})")
            st.caption(f"source: {it.source_file}")

            render_markdown_block("문제", n.get("problem_text_md", ""))
            render_markdown_block("요구사항", n.get("ask_line_md", ""))

            img_bytes = try_resolve_image(n, it.source_file)
            if img_bytes:
                st.image(img_bytes, caption="그림", use_column_width=True)

            render_table_block("제시표", n.get("given_table"))
            if n.get("full_table") not in (None, ""):
                with st.expander("완성표(있을 때만)", expanded=False):
                    render_table_block("완성표", n.get("full_table"))

            if n.get("answer_md"):
                with st.expander("정답", expanded=False):
                    st.markdown(n["answer_md"])
            if n.get("explain_md"):
                with st.expander("해설", expanded=False):
                    st.markdown(n["explain_md"])

            with st.expander("payload 키", expanded=False):
                st.code(", ".join(n.get("payload_keys", [])))

        st.divider()

        # export selection
        selected_items = [x for x in items if x.uid in st.session_state.selected_uids]
        st.write(f"선택 문항: **{len(selected_items)}개**")

        ec1, ec2, ec3 = st.columns([1, 1, 1])
        with ec1:
            include_full = st.checkbox("완성표 포함", value=True)
        with ec2:
            include_expl = st.checkbox("해설 포함", value=True)
        with ec3:
            st.caption("DOCX는 2단 + 정답/해설 뒤에 모음")

        if selected_items:
            docx_bytes = export_docx_bytes(selected_items, include_expl=include_expl, include_full=include_full)
            st.download_button(
                "DOCX 다운로드",
                data=docx_bytes,
                file_name="export_selected.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            # json export(선택)
            pack = {
                "module_code": "MERGED",
                "created_at": "",
                "items": [{"id": it.pid, "payload": it.payload} for it in selected_items],
            }
            j = json.dumps(pack, ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button(
                "선택 PACK(JSON) 다운로드",
                data=j,
                file_name="export_selected_pack.json",
                mime="application/json",
                use_container_width=True,
            )
        else:
            st.info("왼쪽에서 문항을 체크해 선택하세요.")

if __name__ == "__main__":
    main()
