# division_generator_PACK.py
# ------------------------------------------------------------
# ✅ Division 4-cells (Templates A/B) + UNIQUE MAPPING SOLVER
# ✅ PACK(JSON) 저장 + problem_bank_web 호환(problem_text_md/ask_line_md 포함)
#
# 템플릿(사용자 확정):
#   Template A (짝수 문제번호):
#     I  : 2n(2)
#     II : n(2)
#     III: n(1)  (II에서 분열)
#     IV : n(1)  (II에서 분열된 것이 아님)
#
#   Template B (홀수 문제번호):
#     I  : 2n(2)
#     II : 2n(4)
#     III: n(2)
#     IV : n(1)  (III에서 분열된 것이 아님)
#
# 추가 설정(사용자 확정):
#   - X연관 유전자쌍: 2개 (LOCI 3개 중 2개가 X, 1개가 상염색체)
#   - ? 마스킹: 문제당 5~6개
#   - 30문항 생성
#   - 난이도 자동태그는 기본 off
#
# 출력:
#   1) DOCX(문제 2단/정답 모음)
#   2) PACK JSON(output/division_pack/...)
#
# 준비:
#   pip install python-docx
#   같은 폴더에 problem_pack.py, tree_base_A.png, tree_base_B.png 두기
#
# 실행:
#   python division_generator_PACK.py
# ------------------------------------------------------------

from __future__ import annotations

import os
import time
import json
import random
import itertools
from dataclasses import dataclass
from typing import Dict, Tuple, Optional, List, Set, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from problem_pack import ProblemPack

# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
MASK_RANGE = (7, 8)  # per problem

MODULE_CODE = "DIV4C"
ID_PREFIX = "DIV4C_"

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
PACK_DIR = os.path.join(OUT_DIR, "division_pack")
DOCX_DIR = os.path.join(OUT_DIR, "division_docx")
os.makedirs(PACK_DIR, exist_ok=True)
os.makedirs(DOCX_DIR, exist_ok=True)

IMG_A = "tree_base_A.png"
IMG_B = "tree_base_B.png"

# 이미지 크기(너무 크면 줄이기)
IMG_WIDTH_IN = 1.2

# 유일정답(=가능한 매핑 1개)만 통과
REQUIRE_UNIQUE_MAPPING = True

# -------------------------
# DATA TYPES
# -------------------------
AlleleAmount = Tuple[int, int]  # (대문자 대립유전자량, 소문자 대립유전자량)

LOCI = ["E", "F", "G"]
ALLELES = {"E": ("E", "e"), "F": ("F", "f"), "G": ("G", "g")}

KOREAN_LABELS = ["가", "나", "다", "라"]  # 표의 행 라벨(=I~IV의 순서없는 표기)

@dataclass
class WorldSpec:
    """
    linked_pair:
      - None: 3독립
      - ("E","F") 등: 2연관 1독립 (연관은 '완전 연관'으로 처리)
    x_loci:
      - X염색체에 있는 locus 집합 (사용자 확정: size=2)
    """
    linked_pair: Optional[Tuple[str, str]]
    x_loci: Set[str]

@dataclass
class Cell:
    stage: str                        # "2n(2)" / "2n(4)" / "n(2)" / "n(1)"
    amounts: Dict[str, AlleleAmount]  # locus -> (upper, lower)
    parent: Optional[str] = None      # "II" / "III" 등
    note: str = ""

# -------------------------
# AMOUNT DOMAINS (male model with Y-logic baked-in by X-locus rules)
# -------------------------
def rand_autosome_2n2() -> AlleleAmount:
    # 2n(2)에서 상염색체: 합 2
    return random.choice([(2,0),(1,1),(0,2)])

def rand_autosome_2n4() -> AlleleAmount:
    # 2n(4)에서 상염색체: 합 4, (3,1) 같은 것은 금지 -> (4,0)(2,2)(0,4)
    return random.choice([(4,0),(2,2),(0,4)])

def rand_autosome_n2() -> AlleleAmount:
    # n(2)에서 상염색체: (2,0) or (0,2)
    return random.choice([(2,0),(0,2)])

def rand_autosome_n1() -> AlleleAmount:
    # n(1)에서 상염색체: (1,0) or (0,1)
    return random.choice([(1,0),(0,1)])

def rand_x_2n2() -> AlleleAmount:
    # 남성 XY에서 X연관: 2n(2)에서도 합 1 (X 하나만 존재)
    # (1,0) 또는 (0,1) 또는 (0,0)?? -> (0,0)은 Y만 있을 때의 논리인데
    # 여기서는 'Y유전자와 매칭하지 않는다' 설정이므로 (0,0)은 n에서만 등장시키는 쪽이 안정적
    return random.choice([(1,0),(0,1)])

def rand_x_2n4() -> AlleleAmount:
    # X연관 2n(4): 합 2 (복제) -> (2,0) or (0,2)
    return random.choice([(2,0),(0,2)])

def rand_x_n2() -> AlleleAmount:
    # n(2)에서 X연관: X정자(합2) or Y정자(합0)
    return random.choice([(2,0),(0,2),(0,0)])

def rand_x_n1() -> AlleleAmount:
    # n(1)에서 X연관: X정자(합1) or Y정자(합0)
    return random.choice([(1,0),(0,1),(0,0)])

def domain_for(stage: str, locus: str, spec: WorldSpec) -> List[AlleleAmount]:
    is_x = (locus in spec.x_loci)
    if stage == "2n(2)":
        return [rand_x_2n2()] if is_x else [rand_autosome_2n2()]
    if stage == "2n(4)":
        return [rand_x_2n4()] if is_x else [rand_autosome_2n4()]
    if stage == "n(2)":
        return [rand_x_n2()] if is_x else [rand_autosome_n2()]
    if stage == "n(1)":
        return [rand_x_n1()] if is_x else [rand_autosome_n1()]
    raise ValueError(stage)

# -------------------------
# LINKAGE: 완전연관(2연관 1독립)
#   - 2연관이면 두 locus가 "함께 이동"하도록 n세포에서 같은 방향으로 고정
# -------------------------
def enforce_linkage_n_cell(amounts: Dict[str, AlleleAmount], linked_pair: Tuple[str,str]) -> None:
    a,b = linked_pair
    # n(2)/n(1)에서 (a,b) 둘 다 'upper 쪽' 또는 둘 다 'lower 쪽' 또는 둘 다 0(=Y)로 맞추는 방식
    # 간단 규칙: a의 상태(upper>0? lower>0? all0?)를 b에 복제
    def state(x: AlleleAmount) -> str:
        if x==(0,0): return "0"
        if x[0]>0 and x[1]==0: return "U"
        if x[1]>0 and x[0]==0: return "L"
        # (1,1) 같은 건 여기서 나오지 않음
        return "M"
    st = state(amounts[a])
    if st=="U":
        amounts[b] = (amounts[b][0] if amounts[b][0]>0 else (1 if amounts[a][0]==1 else 2), 0)
        # 위는 stage별 크기가 달라질 수 있어 그냥 방향만 유지
    elif st=="L":
        amounts[b] = (0, amounts[b][1] if amounts[b][1]>0 else (1 if amounts[a][1]==1 else 2))
    elif st=="0":
        amounts[b] = (0,0)

# -------------------------
# TEMPLATE BUILDERS
# -------------------------
def build_template_A(spec: WorldSpec) -> Dict[str, Cell]:
    """
    A:
      I : 2n(2)
      II: n(2)
      III: n(1) from II
      IV: n(1) not from II
    """
    I = Cell(stage="2n(2)", amounts={})
    II = Cell(stage="n(2)", amounts={}, parent=None, note="II")
    III = Cell(stage="n(1)", amounts={}, parent="II", note="III(from II)")
    # I
    for L in LOCI:
        I.amounts[L] = domain_for("2n(2)", L, spec)[0]
    # II
    for L in LOCI:
        II.amounts[L] = domain_for("n(2)", L, spec)[0]
    # linkage enforce on II if applicable
    if spec.linked_pair:
        enforce_linkage_n_cell(II.amounts, spec.linked_pair)

    # III derived from II: halve non-zero entries in II where possible
    for L in LOCI:
        u,l = II.amounts[L]
        if (u,l)==(0,0):
            III.amounts[L] = (0,0)
        else:
            # n(2)->n(1): halve 2->1, keep direction
            if u>0 and l==0:
                III.amounts[L] = (1,0)
            elif l>0 and u==0:
                III.amounts[L] = (0,1)
            else:
                # shouldn't happen
                III.amounts[L] = (0,0)

    if spec.linked_pair:
        enforce_linkage_n_cell(III.amounts, spec.linked_pair)

    # IV: n(1) independent, but must be "not from II" meaning amounts not identical to III
    banned = dict(III.amounts)
    for _ in range(5000):
        cand = {}
        for L in LOCI:
            cand[L] = domain_for("n(1)", L, spec)[0]
        if spec.linked_pair:
            enforce_linkage_n_cell(cand, spec.linked_pair)
        if cand != banned:
            IV = Cell(stage="n(1)", amounts=cand, parent=None, note="IV(not from II)")
            return {"I": I, "II": II, "III": III, "IV": IV}
    raise RuntimeError("Template A: IV 생성 실패")

def build_template_B(spec: WorldSpec) -> Dict[str, Cell]:
    """
    B:
      I : 2n(2)
      II: 2n(4)
      III: n(2)
      IV: n(1) not from III
    """
    I = Cell(stage="2n(2)", amounts={})
    II = Cell(stage="2n(4)", amounts={}, parent="I", note="II(from I)")
    III = Cell(stage="n(2)", amounts={}, parent="II", note="III(from II)")
    # I
    for L in LOCI:
        I.amounts[L] = domain_for("2n(2)", L, spec)[0]
    # II doubles I (2n(2)->2n(4))
    for L in LOCI:
        u,l = I.amounts[L]
        II.amounts[L] = (u*2, l*2)
        # for autosome, u+l==2 => 4 good; for X, u+l==1 => 2 good
    # III splits II (2n(4)->n(2))
    for L in LOCI:
        u,l = II.amounts[L]
        if u==0 and l==0:
            III.amounts[L] = (0,0)
        else:
            # choose one homologue / chromatid set: keep one side
            # For autosome: (4,0)->(2,0), (2,2)->(2,0) or (0,2) ??? but our II autosome never (2,2) unless I was (1,1)
            if u>0 and l==0:
                III.amounts[L] = (2,0)
            elif l>0 and u==0:
                III.amounts[L] = (0,2)
            else:
                # (2,2) case: pick (2,0) or (0,2)
                III.amounts[L] = random.choice([(2,0),(0,2)])
    if spec.linked_pair:
        enforce_linkage_n_cell(III.amounts, spec.linked_pair)

    # IV: n(1) not from III
    banned = {}
    for L in LOCI:
        u,l = III.amounts[L]
        if (u,l)==(0,0):
            banned[L]=(0,0)
        else:
            banned[L]=(1,0) if u>0 else (0,1)

    for _ in range(5000):
        cand={}
        for L in LOCI:
            cand[L]=domain_for("n(1)", L, spec)[0]
        if spec.linked_pair:
            enforce_linkage_n_cell(cand, spec.linked_pair)
        if cand != banned:
            IV = Cell(stage="n(1)", amounts=cand, parent=None, note="IV(not from III)")
            return {"I": I, "II": II, "III": III, "IV": IV}
    raise RuntimeError("Template B: IV 생성 실패")

# -------------------------
# WORLD SPEC GENERATION
# -------------------------
def make_spec() -> WorldSpec:
    # ✅ X-연관 유전자쌍 개수: 1개 또는 2개 (랜덤)
    k = random.choice([1, 2])
    x_loci = set(random.sample(LOCI, k))

    # linkage: None(3독립) or one pair
    pairs = [("E","F"),("E","G"),("F","G")]
    linked_pair = random.choice([None, random.choice(pairs)])
    return WorldSpec(linked_pair=linked_pair, x_loci=x_loci)

# -------------------------
# TABLE REPRESENTATION
#  - Present as rows = 가/나/다/라, cols = E e F f G g
# -------------------------
COLS = ["E","e","F","f","G","g"]

def cell_to_row(cell: Cell) -> Dict[str, int]:
    row={}
    for L in LOCI:
        U,lo = ALLELES[L]
        u_amt, l_amt = cell.amounts[L]
        row[U]=u_amt
        row[lo]=l_amt
    return row

def build_full_table(cells: Dict[str,Cell], lab_to_stage: Dict[str,str]) -> Dict[str, Dict[str,int]]:
    # lab_to_stage: "가"->"I" etc
    table={}
    for lab, stg in lab_to_stage.items():
        table[lab]=cell_to_row(cells[stg])
    return table

# -------------------------
# MASKING
# -------------------------
def mask_table(full: Dict[str,Dict[str,int]], nmask: int) -> Dict[str,Dict[str,Any]]:
    masked={lab: dict(full[lab]) for lab in full}
    candidates=[(lab,c) for lab in KOREAN_LABELS for c in COLS]
    random.shuffle(candidates)
    done=0
    for lab,c in candidates:
        if done>=nmask: break
        # avoid masking all 6 in a row
        known=sum(1 for v in masked[lab].values() if v!="?")
        if known<=2: 
            continue
        masked[lab][c]="?"
        done+=1
    return masked

# -------------------------
# SOLVER: count possible mappings
#   - brute force 24 mappings between (가나다라) and (I..IV)
#   - check if there exists completion of "?" that matches template rules/spec
# -------------------------
def stage_set_for_template(template: str) -> List[str]:
    if template=="A":
        return ["I","II","III","IV"]
    return ["I","II","III","IV"]

def enumerate_amounts_for_stage(stage: str, spec: WorldSpec) -> List[Dict[str,AlleleAmount]]:
    # generate possible per-locus allele amounts combinations for a stage
    domains={}
    for L in LOCI:
        is_x = (L in spec.x_loci)
        if stage=="2n(2)":
            domains[L]=[(1,0),(0,1)] if is_x else [(2,0),(1,1),(0,2)]
        elif stage=="2n(4)":
            domains[L]=[(2,0),(0,2)] if is_x else [(4,0),(2,2),(0,4)]
        elif stage=="n(2)":
            domains[L]=[(2,0),(0,2),(0,0)] if is_x else [(2,0),(0,2)]
        elif stage=="n(1)":
            domains[L]=[(1,0),(0,1),(0,0)] if is_x else [(1,0),(0,1)]
        else:
            raise ValueError(stage)
    out=[]
    for combo in itertools.product(domains["E"], domains["F"], domains["G"]):
        cand={"E":combo[0],"F":combo[1],"G":combo[2]}
        if spec.linked_pair and stage in ("n(2)","n(1)"):
            enforce_linkage_n_cell(cand, spec.linked_pair)
        out.append(cand)
    return out

def row_matches_amounts(masked_row: Dict[str,Any], amounts: Dict[str,AlleleAmount]) -> bool:
    # masked_row has keys E,e,F,f,G,g with int or "?"
    for L in LOCI:
        U,lo = ALLELES[L]
        u_amt,l_amt = amounts[L]
        if masked_row[U]!="?" and masked_row[U]!=u_amt: return False
        if masked_row[lo]!="?" and masked_row[lo]!=l_amt: return False
    return True

def template_constraints_ok(template: str, spec: WorldSpec, amts: Dict[str,Dict[str,AlleleAmount]]) -> bool:
    """
    amts: stage-> amounts dict for E/F/G
    """
    if template=="A":
        # II is n(2), III is n(1) derived from II
        II = amts["II"]; III=amts["III"]
        for L in LOCI:
            u2,l2 = II[L]
            u1,l1 = III[L]
            if (u2,l2)==(0,0):
                if (u1,l1)!=(0,0): return False
            else:
                # must halve 2->1 direction preserved
                if u2>0 and l2==0 and (u1,l1)!=(1,0): return False
                if l2>0 and u2==0 and (u1,l1)!=(0,1): return False
                # if somehow (2,2) not expected in n(2) domain
        # IV not from II => IV amounts != III amounts
        if amts["IV"]==amts["III"]: return False
        # (중요) IV는 II에서 분열된 세포가 아니더라도 같은 개체이므로 유전자형과 모순되면 안 된다.
        I = amts["I"]  # 2n(2) 기준
        for L in LOCI:
            ref = I[L]
            iv = amts["IV"][L]
            if (L in spec.x_loci) and iv==(0,0):
                continue
            if ref==(2,0):
                if iv!=(1,0): return False
            elif ref==(0,2):
                if iv!=(0,1): return False
            elif ref==(1,1):
                if iv not in [(1,0),(0,1)]: return False
            elif ref==(1,0):
                if iv not in ([(1,0)] + ([(0,0)] if (L in spec.x_loci) else [])): return False
            elif ref==(0,1):
                if iv not in ([(0,1)] + ([(0,0)] if (L in spec.x_loci) else [])): return False
            elif ref==(0,0):
                if iv!=(0,0): return False
        return True

    # template B
    I=amts["I"]; II=amts["II"]; III=amts["III"]
    # II doubles I
    for L in LOCI:
        if II[L]!=(I[L][0]*2, I[L][1]*2): return False
    # III from II by selecting one side
    for L in LOCI:
        u4,l4 = II[L]
        u2,l2 = III[L]
        if (u4,l4)==(0,0):
            if (u2,l2)!=(0,0): return False
        else:
            # for X: (2,0)-> (2,0) OR (0,2)->(0,2) OR (0,0)? but n(2) domain allows 0,0. from II should not produce 0,0
            if (u2,l2)==(0,0): 
                return False
            # autosome: (4,0)->(2,0), (0,4)->(0,2), (2,2)->(2,0) or (0,2)
            if u4==4 and l4==0 and (u2,l2)!=(2,0): return False
            if l4==4 and u4==0 and (u2,l2)!=(0,2): return False
            if (u4,l4)==(2,2) and (u2,l2) not in [(2,0),(0,2)]: return False
            if (u4,l4) in [(2,0),(0,2)] and (u2,l2)!=(u4,l4): 
                # X locus in II is (2,0) or (0,2) and III should match
                return False
    # IV not from III
    if amts["IV"]=={L: ((1,0) if III[L][0]>0 else (0,1) if III[L][1]>0 else (0,0)) for L in LOCI}:
        return False
    # (중요) 같은 개체의 독립 n(1) 세포(IV)도 원래 유전자형과 모순되면 안 된다.
    # 예: 2n(4)에서 (4,0)이면 해당 대립유전자만 존재하므로 n(1)은 (1,0)이어야 하고 (0,1)은 불가.
    for L in LOCI:
        ref = I[L]  # 2n(2) 단계(I)에서의 정보로 유전자형 범위를 고정
        iv = amts["IV"][L]
        # X-연관에서 Y 정자(0,0)는 예외적으로 허용
        if (L in spec.x_loci) and iv==(0,0):
            continue
        if ref==(2,0):
            if iv!=(1,0): return False
        elif ref==(0,2):
            if iv!=(0,1): return False
        elif ref==(1,1):
            if iv not in [(1,0),(0,1)]: return False
        elif ref==(1,0):  # 남성 X(2n(2))
            if iv not in ([(1,0)] + ([(0,0)] if (L in spec.x_loci) else [])): return False
        elif ref==(0,1):
            if iv not in ([(0,1)] + ([(0,0)] if (L in spec.x_loci) else [])): return False
        elif ref==(0,0):
            if iv!=(0,0): return False
    return True

def _count_possible_mappings_given_spec(masked: Dict[str,Dict[str,Any]], template: str, spec: WorldSpec) -> Tuple[int, List[Dict[str,str]]]:
    stages = ["I","II","III","IV"]
    stage_to_required = {
        "A": {"I":"2n(2)","II":"n(2)","III":"n(1)","IV":"n(1)"},
        "B": {"I":"2n(2)","II":"2n(4)","III":"n(2)","IV":"n(1)"},
    }[template]

    # pre-enumerate domains per stage
    stage_domains = {}
    for stg in stages:
        stage_domains[stg]=enumerate_amounts_for_stage(stage_to_required[stg], spec)

    solutions=[]
    cnt=0
    for perm in itertools.permutations(stages, 4):
        lab_to_stage = dict(zip(KOREAN_LABELS, perm))
        # early row match pruning:
        ok=True
        stage_choice={}
        for lab, stg in lab_to_stage.items():
            req_stage = stage_to_required[stg]
            # filter domain by row match
            candidates=[a for a in stage_domains[stg] if row_matches_amounts(masked[lab], a)]
            if not candidates:
                ok=False; break
            stage_choice[stg]=candidates
        if not ok:
            continue

        # Now need pick one candidate for each stage such that template constraints ok
        # brute force nested loops (domains are small: <=27 each)
        for aI in stage_choice["I"]:
            for aII in stage_choice["II"]:
                for aIII in stage_choice["III"]:
                    for aIV in stage_choice["IV"]:
                        amts={"I":aI,"II":aII,"III":aIII,"IV":aIV}
                        if template_constraints_ok(template, spec, amts):
                            cnt += 1
                            solutions.append(lab_to_stage)
                            # unique만 필요하면 조기중단
                            if REQUIRE_UNIQUE_MAPPING and cnt>1:
                                return cnt, solutions
                            break
                    if REQUIRE_UNIQUE_MAPPING and cnt>1: break
                if REQUIRE_UNIQUE_MAPPING and cnt>1: break
            if REQUIRE_UNIQUE_MAPPING and cnt>1: break

    return cnt, solutions



def count_possible_mappings_hidden_x(masked: Dict[str,Dict[str,Any]],
                                    template: str,
                                    linked_pair: Optional[Tuple[str,str]]) -> Tuple[int, List[Tuple[Dict[str,str], List[str]]]]:
    """학생 입장에서 X-연관 유전자쌍이 무엇인지(그리고 몇 개인지)를 모르도록 출제하므로,
    솔버(유일정답 판정)는 X-연관 후보(1개/2개)를 모두 열어두고 가능한 해를 합쳐서 센다.

    Returns:
      (총 해 개수, [ (lab_to_stage, x_loci_sorted_list) ... ])
    """
    all_solutions: List[Tuple[Dict[str,str], List[str]]] = []
    total = 0

    x_candidates: List[Tuple[str,...]] = []
    x_candidates += list(itertools.combinations(LOCI, 1))
    x_candidates += list(itertools.combinations(LOCI, 2))

    for xset in x_candidates:
        spec = WorldSpec(linked_pair=linked_pair, x_loci=set(xset))
        cnt, sols = _count_possible_mappings_given_spec(masked, template, spec)
        if cnt:
            total += cnt
            for s in sols:
                all_solutions.append((s, list(xset)))
        if REQUIRE_UNIQUE_MAPPING and total > 1:
            return total, all_solutions

    return total, all_solutions

    return total, all_solutions

# -------------------------
# PROBLEM TEXT (Markdown) for web preview
# -------------------------
def table_to_md(masked: Dict[str,Dict[str,Any]]) -> str:
    header = "|세포| " + " | ".join(COLS) + " |\n"
    sep = "|" + "---|"*(len(COLS)+1) + "\n"
    rows=""
    for lab in KOREAN_LABELS:
        vals=[str(masked[lab][c]) for c in COLS]
        rows += f"|{lab}| " + " | ".join(vals) + " |\n"
    return header+sep+rows

def linkage_text(spec: WorldSpec) -> str:
    if spec.linked_pair is None:
        return "(E/e), (F/f), (G/g)는 서로 독립일 수도 있고 아닐 수도 있을 것이다."
    a,b = spec.linked_pair
    others = [x for x in LOCI if x not in spec.linked_pair][0]
    return f"({a}/{a.lower()}), ({b}/{b.lower()})는 완전 연관(같이 이동)이고, ({others}/{others.lower()})는 독립이다."

def x_text(spec: WorldSpec) -> str:
    xs=sorted(list(spec.x_loci))
    autos=[x for x in LOCI if x not in spec.x_loci]
    return f"X염색체 연관 유전자쌍은 {len(xs)}개이며 {', '.join([f'({x}/{x.lower()})' for x in xs])} 이고, 나머지 {', '.join([f'({x}/{x.lower()})' for x in autos])} 는 상염색체에 있다."

def make_problem_text(template: str, masked: Dict[str,Dict[str,Any]], spec: WorldSpec) -> Tuple[str,str]:
    """
    Returns (problem_text_md, ask_line_md)
    """
    tname = "A" if template=="A" else "B"
    problem_md = (
        f"### Division 4-cells (Template {tname})\n"
        f"- {linkage_text(spec)}\n"
        f"- 아래 표의 (가),(나),(다),(라)는 그림의 세포 I, II, III, IV를 순서 없이 나타낸 것이다.\n\n"
        f"{table_to_md(masked)}\n"
    )
    ask_md = "표에서 (가),(나),(다),(라)에 해당하는 세포(I~IV)를 찾고, 물음표를 모두 채우시오."
    return problem_md, ask_md

# -------------------------
# DOCX UTIL
# -------------------------
def set_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "바탕"
    style.font.size = Pt(9)

def add_md_as_plain(par, md: str):
    # 최소 구현: markdown을 그대로 문단으로 넣음(표는 docx table로 따로 처리)
    for line in md.splitlines():
        par.add_run(line)
        par.add_run("\n")

def add_table_docx(container, masked: Dict[str,Dict[str,Any]]):
    t = container.add_table(rows=1+len(KOREAN_LABELS), cols=1+len(COLS))
    t.style="Table Grid"
    t.cell(0,0).text="세포"
    for j,c in enumerate(COLS):
        t.cell(0,j+1).text=c
    for i,lab in enumerate(KOREAN_LABELS):
        t.cell(i+1,0).text=lab
        for j,c in enumerate(COLS):
            t.cell(i+1,j+1).text=str(masked[lab][c])

def add_problem_cell(cell, pnum: int, template: str, spec: WorldSpec, masked: Dict[str,Dict[str,Any]], img_path: str):
    title = cell.add_paragraph(f"[문제 {pnum}]")
    title.runs[0].bold=True

    # 그림
    if os.path.exists(img_path):
        p = cell.add_paragraph()
        run = p.add_run()
        run.add_picture(img_path, width=Inches(IMG_WIDTH_IN))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell.add_paragraph(linkage_text(spec))
    cell.add_paragraph("다음 표의 (가),(나),(다),(라)는 그림의 세포 I, II, III, IV를 순서 없이 나타낸 것이다.")
    add_table_docx(cell, masked)
    cell.add_paragraph("표에서 (가),(나),(다),(라)에 해당하는 세포(I~IV)를 찾고, 물음표를 모두 채우시오.")

# -------------------------
# GENERATION (with progress)
# -------------------------
def generate_one(pnum: int) -> Dict[str,Any]:
    template = "A" if (pnum % 2 == 0) else "B"
    img_path = IMG_A if template=="A" else IMG_B

    for attempt in range(1, 4000+1):
        spec = make_spec()
        # build cells
        try:
            cells = build_template_A(spec) if template=="A" else build_template_B(spec)
        except RuntimeError:
            continue

        # shuffle mapping between labs and stages
        perm = list(itertools.permutations(["I","II","III","IV"], 4))
        lab_to_stage = dict(zip(KOREAN_LABELS, random.choice(perm)))

        full = build_full_table(cells, lab_to_stage)
        nmask = random.randint(MASK_RANGE[0], MASK_RANGE[1])
        masked = mask_table(full, nmask)

        # ✅ 유일정답 판정: X-연관 정보는 문제에서 숨기므로(개수/무엇인지 미제시),
        #    솔버는 X-연관 후보(1개/2개)를 모두 열어두고 가능한 해를 합쳐서 센다.
        cnt, sols = count_possible_mappings_hidden_x(masked, template, spec.linked_pair)
        if REQUIRE_UNIQUE_MAPPING and cnt != 1:
            continue

        # unique라면 첫 해가 정답(라벨→단계 + 추론된 X-loci)
        sol_map, inferred_x = (sols[0] if sols else (None, []))

        problem_text_md, ask_line_md = make_problem_text(template, masked, spec)

        payload = {
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "template": template,
            "img": img_path,
            # 문제에는 X정보를 숨기지만, DB/해설용으로는 저장
            "x_loci_true": sorted(list(spec.x_loci)),
            "x_loci_inferred": inferred_x,
            "linked_pair": list(spec.linked_pair) if spec.linked_pair else None,
            "masked_table": masked,
            "full_table": full,
            "lab_to_stage_answer": sol_map,
            "mask_used": nmask,
            "unique_mapping_count": cnt,
        }

        return {
            "template": template,
            "img": img_path,
            "spec": spec,
            "masked": masked,
            "full": full,
            "masked_table": masked,
            "full_table": full,
            "answer_map": sol_map,
            "inferred_x": inferred_x,
            "payload": payload,
        }

    raise RuntimeError(f"[P{pnum:02d}] 문제 생성 실패: 시도횟수 부족/조건 과함")

def make_docx_and_pack():
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    out_docx = os.path.join(DOCX_DIR, f"Division_4cells_{timestamp}.docx")
    pack = ProblemPack(module_code=MODULE_CODE, out_dir=PACK_DIR, id_prefix=ID_PREFIX)

    doc = Document()
    set_doc_style(doc)

    problems=[]
    for pnum in range(1, N_PROBLEMS+1):
        print(f"[{pnum}/{N_PROBLEMS}] 생성 중...")
        w = generate_one(pnum)
        problems.append(w)

        # PACK 저장(문항 단위)
        pack.new_problem(qnum=pnum, payload=w["payload"])

    # DOCX: 2단 2문제/페이지
    idx=0
    pnum=1
    while idx < N_PROBLEMS:
        tbl = doc.add_table(rows=1, cols=2)
        left, right = tbl.rows[0].cells[0], tbl.rows[0].cells[1]

        w = problems[idx]
        add_problem_cell(left, pnum, w["template"], w["spec"], w["masked"], w["img"])
        idx+=1; pnum+=1

        if idx < N_PROBLEMS:
            w = problems[idx]
            add_problem_cell(right, pnum, w["template"], w["spec"], w["masked"], w["img"])
            idx+=1; pnum+=1

        if idx < N_PROBLEMS:
            doc.add_page_break()

    # 정답 모음
    doc.add_page_break()
    h = doc.add_paragraph("[정답]")
    h.runs[0].bold=True
    for i,w in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번: {w['answer_map']} (template {w['template']}, ?={w['payload']['mask_used']})")
        # ✅ 완성표(물음표 없는 원본)
        add_table_docx(doc, w["full_table"])
        doc.add_paragraph("")

    doc.save(out_docx)
    pack_path = pack.save_json()
    print("✅ DOCX:", out_docx)
    print("✅ PACK:", pack_path)

if __name__ == "__main__":
    make_docx_and_pack()