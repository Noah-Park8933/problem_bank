# pded1_generator_PACK.py
# ------------------------------------------------------------
# PDED1: Polygenic Deduction 1 (연관 + 조건추론) 30문제 딸깍 생성기
#
# - (X/x),(Y/y)는 연관(완전연관 r=0 가정; 문제/해설에서 교차 없음 명시)
# - (Z/z)는 독립
# - 표현형 숫자 = 대문자(우성) 대립유전자 개수 총합 (0~6)
#
# 고정 논리(유일해):
# 1) P1: XxYyZZ (X-Y는 결합상 XY/xy)
# 2) P2: XxYyZz (X-Y는 반발상 Xy/xY)
# 3) P( xxYyZz ) = 1/8,  P( XxYyZz ) = 0
# 4) 자손 표현형 4가지, 각 1/4 (2,3,4,5)
#
# 출력:
# - output/PDED1_PACK_<batch>.json
# - output/PDED1_<timestamp>.docx
#
# 실행:
#   pip install python-docx
#   python pded1_generator_PACK.py
# ------------------------------------------------------------

import os, time, random
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from problem_pack import ProblemPack


# ------------------------------
# CONFIG
# ------------------------------
MODULE_CODE = "PDED1"
ID_PREFIX = "PDED1_"
N_PROBLEMS = 30

OUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUT_DIR, exist_ok=True)


# ------------------------------
# Utilities
# ------------------------------
def pick_gene_triplet() -> Tuple[str, str, str]:
    """
    대문자 유전자 3개를 뽑아 (X,Y,Z)로 사용.
    보기 좋게 E~Z에서 뽑고, 서로 겹치지 않게.
    """
    pool = list("EFGHIJKLMNPQRSTUVWXZY")  # O 제외, 보기용
    X = random.choice(pool); pool.remove(X)
    Y = random.choice(pool); pool.remove(Y)
    Z = random.choice(pool); pool.remove(Z)
    return X, Y, Z


def geno_str(X: str, Y: str, Z: str, phase_xy: str, z: str) -> str:
    """
    예: X='E',Y='F',Z='G', phase_xy='EF/ef', z='GG' -> 'EeFfGG' 같이 붙여쓰기
    """
    x = X.lower()
    y = Y.lower()
    zlo = Z.lower()

    # XY의 개별 유전자형은 항상 XxYy로 고정이므로 문자열은 XxYy부터 구성
    xy_part = f"{X}{x}{Y}{y}"  # 예: EeFf
    z_part = z  # 'GG' or 'Gg'

    return f"{xy_part}{z_part}"


def phenotype_number_from_genotype(g: Dict[str, str]) -> int:
    """
    g: {'X':'Xx'/'xx', 'Y':'Yy'/'yy', 'Z':'ZZ'/'Zz'/'zz'} 형태를 받아
    대문자 대립유전자 총개수 반환(0~6)
    """
    total = 0
    for key, pair in g.items():
        for ch in pair:
            if ch.isupper():
                total += 1
    return total


def md_table_kv(rows: List[Tuple[str, str]]) -> str:
    """
    간단 마크다운 표
    """
    out = ["|항목|값|", "|---|---|"]
    for k, v in rows:
        out.append(f"|{k}|{v}|")
    return "\n".join(out)


# ------------------------------
# Core logic (유일해 고정)
# ------------------------------
def build_one_problem(pnum: int) -> Dict:
    """
    문제 1개 생성(겉모양만 랜덤: 유전자 기호 X,Y,Z)
    """
    X, Y, Z = pick_gene_triplet()
    x, y, z = X.lower(), Y.lower(), Z.lower()

    # P1 고정: XxYyZZ, 결합상 XY/xy
    P1_phase = f"{X}{Y}/{x}{y}"      # 예: EF/ef
    P1_Z = f"{Z}{Z}"                 # ZZ

    # P2 고정: XxYyZz, 반발상 Xy/xY, Z는 이형
    P2_phase = f"{X}{y}/{x}{Y}"      # 예: Ef/eF
    P2_Z = f"{Z}{z}"                 # Zz

    P1_geno_txt = geno_str(X, Y, Z, P1_phase, P1_Z)  # 예: EeFfGG
    P2_geno_txt = geno_str(X, Y, Z, P2_phase, P2_Z)  # 예: EeFfGg

    # 조건에 등장할 특정 자손 유전자형 2개:
    #   xxYyZz 확률 1/8
    #   XxYyZz 확률 0
    target1 = f"{x}{x}{Y}{y}{Z}{z}"  # xxYyZz 붙여쓰기
    target2 = f"{X}{x}{Y}{y}{Z}{z}"  # XxYyZz 붙여쓰기

    # 해설/정답(고정)
    # 자손 표현형 분포: (2),(3),(4),(5) 각 1/4
    # (XY쪽에서 1 또는 3 대문자, Z쪽에서 1 또는 2 대문자)
    phenos = [(2, "1(연관부) + 1(Z부)"), (3, "1 + 2"), (4, "3 + 1"), (5, "3 + 2")]
    pheno_lines = "\n".join([f"- ({k}) : 1/4" for k, _ in phenos])

    # 문제 본문/요구
    problem_text_md = "\n".join([
        f"문제 제목 : Polygenic Deduction 1 (PDED1-{pnum:02d})",
        f"({X}/{x}), ({Y}/{y}), ({Z}/{z}) 3set 유전자에 의해 결정되는 다인자유전을 가정하자.",
        f"단, ({X}/{x}), ({Y}/{y})는 연관이며 **교차는 일어나지 않는다(완전 연관)**.",
        "",
        f"- P2가 P1보다 이형접합성 수가 많다.",
        f"- P1에서 유전자형이 {X}{Y}{Z} 인 생식세포가 형성될 수 있다.",
        f"- 부모 P1과 P2를 교배시킬 때 나오는 자손의 표현형은 4가지이다.",
        f"- 자손의 유전자형이 **{target1}** 일 확률은 **1/8** 이고, **{target2}** 일 확률은 **0** 이다.",
        "",
        "부모 P1, P2의 유전자형을 찾고 자손의 표현형의 종류와 각 표현형의 확률을 구하시오.",
        "",
        "※ 표현형 (k)는 대문자 대립유전자 개수의 합이 k인 것을 의미한다.",
    ])

    ask_line_md = "P1, P2의 유전자형 + 자손 표현형 분포(종류/확률)을 구하시오."

    answer_text_md = "\n".join([
        f"- P1 = {P1_geno_txt}  (연관상: {P1_phase})",
        f"- P2 = {P2_geno_txt}  (연관상: {P2_phase})",
        "",
        "자손 표현형 분포:",
        pheno_lines
    ])

    solution_md = "\n".join([
        "### 풀이 핵심",
        f"1) P1에서 {X}{Y}{Z} 생식세포가 가능하려면, 연관부에 {X}{Y} 하플로타입이 존재해야 하므로 P1은 결합상({P1_phase})이어야 한다.",
        f"2) 자손에서 {target2}(=XxYyZz)가 **0**이 되려면, 연관부에서 AaBb가 나오지 않아야 하므로 P2는 P1과 반대 위상(반발상 {P2_phase})이어야 한다.",
        f"3) {target1}(=xxYyZz)이 되려면 연관부에서 xxYy가 나올 확률이 1/4이고, Z부에서 Zz가 1/2여야 하므로 P1은 {Z}{Z}, P2는 {Z}{z}로 결정된다(1/4×1/2=1/8).",
        "",
        "### 자손 표현형",
        "- 연관부(완전연관)에서 자손은 AABB형(대문자 3개) 또는 Aabb형(대문자 1개)만 나온다(각 1/2).",
        "- Z부(ZZ×Zz)에서 자손은 ZZ(대문자 2개) 또는 Zz(대문자 1개)만 나온다(각 1/2).",
        "- 따라서 총합은 2,3,4,5의 4가지이며 각각 1/4."
    ])

    payload = {
        "problem_code": f"PDED1-{pnum:02d}",
        "genes": {"X": X, "Y": Y, "Z": Z},
        "linked_pair": [X, Y],
        "no_crossover": True,
        "P1": {"genotype": P1_geno_txt, "phase": P1_phase},
        "P2": {"genotype": P2_geno_txt, "phase": P2_phase},
        "targets": {
            "p_is_1_8": {"genotype": target1, "prob": "1/8"},
            "p_is_0": {"genotype": target2, "prob": "0"}
        },
        "offspring_phenotypes": {
            "2": "1/4", "3": "1/4", "4": "1/4", "5": "1/4"
        },
        # 웹 미리보기용 핵심
        "problem_text_md": problem_text_md,
        "ask_line_md": ask_line_md,
        "answer_text_md": answer_text_md,
        "solution_md": solution_md,
        "difficulty_tag": "mid",  # UI 분류용(원하면 규칙으로 바꿔도 됨)
    }

    return payload


# ------------------------------
# DOCX rendering
# ------------------------------
def set_normal_style(doc: Document, font_name="바탕", font_size_pt=10):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)


def add_md_like_paragraphs(doc: Document, md_text: str):
    """
    아주 단순한 md 렌더링:
    - 빈 줄 -> 공백 문단
    - "- "로 시작 -> 불릿 느낌
    - 나머지 -> 일반 문단
    """
    for line in md_text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(12)
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:])
            p.runs[0].bold = True
        else:
            doc.add_paragraph(line)


def make_docx_and_pack():
    pack = ProblemPack(module_code=MODULE_CODE, out_dir=OUT_DIR, id_prefix=ID_PREFIX)

    # DOCX
    doc = Document()
    set_normal_style(doc, font_size_pt=10)

    title = doc.add_paragraph("PDED1 문제 세트")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 문제 생성
    items = []
    for pnum in range(1, N_PROBLEMS + 1):
        payload = build_one_problem(pnum)
        pp = pack.new_problem(qnum=pnum, payload=payload)
        items.append((pp.id, payload))

    # 문제 파트
    for idx, (pid, payload) in enumerate(items, start=1):
        h = doc.add_paragraph(f"[문제 {idx}]  ID: {pid}")
        h.runs[0].bold = True
        add_md_like_paragraphs(doc, payload["problem_text_md"])
        doc.add_paragraph("")
        doc.add_paragraph("—")
        doc.add_paragraph("")

    # 정답/해설 파트
    doc.add_page_break()
    h2 = doc.add_paragraph("[정답 및 해설]")
    h2.runs[0].bold = True
    doc.add_paragraph("")

    for idx, (pid, payload) in enumerate(items, start=1):
        hh = doc.add_paragraph(f"{idx}번  ID: {pid}")
        hh.runs[0].bold = True

        doc.add_paragraph("정답(부모 유전자형 + 표현형 분포)")
        add_md_like_paragraphs(doc, payload["answer_text_md"])
        doc.add_paragraph("")
        doc.add_paragraph("해설")
        add_md_like_paragraphs(doc, payload["solution_md"])
        doc.add_paragraph("")
        doc.add_paragraph("—")
        doc.add_paragraph("")

    # 저장
    ts = int(time.time())
    docx_path = os.path.join(OUT_DIR, f"{MODULE_CODE}_{ts}.docx")
    doc.save(docx_path)

    pack_path = pack.save_json()

    print("✅ DOCX 저장:", docx_path)
    print("✅ PACK 저장:", pack_path)


if __name__ == "__main__":
    make_docx_and_pack()
