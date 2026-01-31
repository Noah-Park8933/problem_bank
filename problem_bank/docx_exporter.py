# docx_exporter.py
print("DOCX_EXPORTER VERSION = NEW_PATCH_20260128")
import os
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .loader import ProblemItem
from .table_renderer import normalize_table_to_grid, try_find_table
from .config import AppConfig


# ------------------------------------------------------------
# 기본 paragraph 생성
# ------------------------------------------------------------
def _set_eastasia_font(style, font_name: str):
    # 한글 폰트 적용(중요)
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def _setup_document(doc: Document):
    # 여백(시험지 느낌)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    # 기본 스타일
    style = doc.styles["Normal"]
    _set_eastasia_font(style, "바탕")
    style.font.size = Pt(9)

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")

def _set_table_cell_margins(table, top=0, bottom=0, left=40, right=40):
    # 단위 dxa(1/20 pt). left/right 40이면 대략 2pt 정도.
    tblPr = table._tbl.tblPr
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)

    def _set(tag, val):
        el = mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            mar.append(el)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")

    _set("top", top); _set("bottom", bottom); _set("left", left); _set("right", right)

def _style_cell_text(cell, size_pt=7, bold=False):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.bold = bold
            r.font.size = Pt(size_pt)

def _add_label(container, text: str):
    # 섹션 라벨(문제/요구사항/제시표/정답/해설) 통일
    p = container.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    return p

def _add_problem_header(container, num: int, it: ProblemItem):
    p = container.add_paragraph()
    r1 = p.add_run(f"{num}. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(f"ID: {it.pid}  ({it.prefix})")
    r2.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_par(doc_or_cell, text: str, bold: bool = False):
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# ------------------------------------------------------------
# 안정적인 표 생성 (오류 방지 패치 완전 적용)
# ------------------------------------------------------------
def _add_grid_table(doc_or_cell, table_obj: Any, total_width_in: float = 2.2):
    """
    - normalize_table_to_grid 결과 기반 표 만든다.
    - headers/rows 비어 있으면 자동 문구 처리.
    - 열수/행수 mismatch 자동 보정.
    - DOCX autofit 끄고 폭 고정해서 깨짐 방지.
    """

    # 파싱 (여기서 headers, rows 비거나 깨져 있을 수도 있음)
    headers, rows = normalize_table_to_grid(table_obj)

    # -----------------------------
    # 방어 1: 헤더가 없음 → 표 생성하지 않고 메시지만
    # -----------------------------
    if not headers or len(headers) == 0:
        _add_par(doc_or_cell, "(표 데이터 없음 / 파싱 불가)")
        return

    n_cols = len(headers)

    # -----------------------------
    # rows 비어있으면 헤더만 표로
    # -----------------------------
    if rows is None:
        rows = []
    if not isinstance(rows, list):
        rows = []

    # -----------------------------
    # 방어 2: 행 길이 정규화
    # -----------------------------
    fixed_rows = []
    for r in rows:
        if r is None:
            r = []
        rr = list(r)
        if len(rr) < n_cols:
            rr = rr + [""] * (n_cols - len(rr))
        elif len(rr) > n_cols:
            rr = rr[:n_cols]
        fixed_rows.append(rr)

    n_rows = 1 + len(fixed_rows)

    # -----------------------------
    # 안전한 DOCX 표 생성
    # -----------------------------
    t = doc_or_cell.add_table(rows=n_rows, cols=n_cols)
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_cell_margins(t, top=0, bottom=0, left=30, right=30)  # ✅ 표가 작고 단정해짐

    total_width = Inches(total_width_in)
    col_w = total_width / n_cols

    # ---- header ----
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.width = col_w
        cell.text = str(h)
        _style_cell_text(cell, size_pt=7, bold=True)

        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    # ---- body ----
    for i, row in enumerate(fixed_rows):
        for j in range(n_cols):
            cell = t.rows[i + 1].cells[j]
            cell.width = col_w
            v = row[j] if row[j] is not None else ""
            cell.text = str(v)
            _style_cell_text(cell, size_pt=7, bold=False)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0


# ------------------------------------------------------------
# 텍스트 찾기 함수
# ------------------------------------------------------------
def first_text(payload: Dict[str, Any], keys: Tuple[str, ...] | List[str]) -> Optional[str]:
    for k in keys:
        v = payload.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return None


# ------------------------------------------------------------
# 이미지 처리 패치 (안정화 완료)
# ------------------------------------------------------------
def _try_add_image(container, payload: Dict[str, Any]):
    """
    Division 등에서 tree_base png 넣기
    - cell(_Cell) / doc(Document) 모두에서 동작하도록 run.add_picture 사용
    """
    img_path = (
        payload.get("_image_path")
        or payload.get("image_path")
        or payload.get("tree_img")
        or payload.get("figure_path")
        or payload.get("fig_path")
        or payload.get("img")
    )

    if not isinstance(img_path, str) or not img_path.strip():
        return

    ip = img_path.strip()

    def resolve_path(p: str) -> Optional[str]:
        # 1) 절대/상대 그대로 존재하면 OK
        if os.path.exists(p):
            return p

        # 2) repo 내부 상대경로 후보
        base = os.path.dirname(__file__)
        candidates = [
            os.path.join(base, p),
            os.path.join(base, "..", p),
            os.path.join(base, "..", "..", p),
        ]
        found = next((x for x in candidates if os.path.exists(x)), None)
        return found

    found = resolve_path(ip)
    if not found:
        _add_par(container, f"(이미지 경로 없음: {ip})")
        return

    try:
        # ✅ 셀에서도 되는 방식: 문단 -> run -> run.add_picture
        container.add_paragraph("")  # 위 여백
        p = container.add_paragraph()
        r = p.add_run()
        r.add_picture(found, width=Inches(1.2))
        container.add_paragraph("")  # 아래 여백
    except Exception as e:
        _add_par(container, f"(이미지 삽입 실패: {found} / {type(e).__name__})")

# ------------------------------------------------------------
# DOCX EXPORT 메인
# ------------------------------------------------------------
def export_docx_bytes(
    cfg: AppConfig,
    selected: List[ProblemItem],
    include_explanations: bool = True,
    include_full_table: bool = True,
    two_columns: bool = True,
) -> bytes:

    doc = Document()
    style = doc.styles["Normal"]
    _setup_document(doc)


    idx = 0
    pnum = 1

    # --------------------------------------------------------
    # 문제 본문 2단 출력
    # --------------------------------------------------------
    while idx < len(selected):

        if two_columns:
            outer = doc.add_table(rows=1, cols=2)
            outer.autofit = False
            _remove_table_borders(outer)           # ✅ 2단이 진짜 2단처럼 보임(테두리 숨김)
            _set_table_cell_margins(outer, left=0, right=0, top=0, bottom=0)  # ✅ 바깥 여백 최소화

            outer.columns[0].width = Inches(3.4)
            outer.columns[1].width = Inches(3.4)

            left = outer.rows[0].cells[0]
            right = outer.rows[0].cells[1]

            targets = [(left, selected[idx], pnum)]
            idx += 1; pnum += 1
            if idx < len(selected):
                targets.append((right, selected[idx], pnum))
                idx += 1; pnum += 1
        else:
            targets = [(doc, selected[idx], pnum)]
            idx += 1; pnum += 1

        for container, it, num in targets:
            payload = it.payload or {}

            _add_par(container, f"[문제 {num}]  ID: {it.pid}  ({it.prefix})", bold=True)

            # 문제/요구사항
            ptxt = first_text(payload, cfg.problem_text_keys)
            atxt = first_text(payload, cfg.ask_line_keys)
            if ptxt:
                _add_par(container, "문제", bold=True)
                _add_par(container, ptxt)
            if atxt:
                _add_par(container, "요구사항", bold=True)
                _add_par(container, atxt)

            # 이미지
            _try_add_image(container, payload)

            # 제시표
            given = try_find_table(payload, list(cfg.given_table_keys)) or payload.get("_given_table")
            if given is not None:
                _add_par(container, "제시표", bold=True)

                try:
                    _add_grid_table(container, given, total_width_in=2.5)
                except Exception:
                    _add_par(container, "(표 변환 실패)")

                _add_par(container, "")

            _add_par(container, "")

        if idx < len(selected):
            doc.add_page_break()

    # --------------------------------------------------------
    # 정답/해설 파트
    # --------------------------------------------------------
    doc.add_page_break()
    _add_par(doc, "[정답/해설]", bold=True)

    for i, it in enumerate(selected, start=1):
        payload = it.payload or {}
        _add_par(doc, f"{i}. ID: {it.pid}  ({it.prefix})", bold=True)

        ans = first_text(payload, cfg.answer_keys)
        expl = first_text(payload, cfg.explanation_keys)

        _add_par(doc, "정답", bold=True)
        _add_par(doc, ans or "(없음)")

        if include_full_table:
            full = try_find_table(payload, list(cfg.full_table_keys)) or payload.get("_full_table")
            if full is not None:
                _add_par(doc, "완성표", bold=True)
                try:
                    _add_grid_table(doc, full, total_width_in=6.4)
                except Exception:
                    _add_par(doc, "(완성표 변환 실패)")

        if include_explanations:
            _add_par(doc, "해설", bold=True)
            _add_par(doc, expl or "(없음)")

        _add_par(doc, "-" * 40)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
