# agglutination_generator_with_reasons.py
# ------------------------------------------------------------
# Agglutination (ABO 응집원/응집소 표) 자동 출제기 + 해설 근거(중간 버전) 자동 생성
#
# 형식:
#  - 열: [BETA(=β) 고정] + [가, 나, 다] (가/나/다는 {응집원A, 응집원B, 응집소α}의 임의 순서)
#  - 행: [A형 고정] + [X, Y, Z] (X/Y/Z는 {B형, AB형, O형}의 임의 순서)
#  - 값: O / X / ? (존재=O, 없음=X)
#
# 기능:
# 1) 랜덤 매핑으로 완성표 생성
# 2) 일부 칸을 ?로 마스킹
# 3) 솔버(전수조사 36)로 유일정답만 통과
# 4) 해설에 "왜 가=α인지 / 왜 X=B형인지" 같은 결정적 단서 문장 자동 생성(중간 버전)
# 5) DOCX 출력: 문제(2단) + 정답 + 해설(완성표+제시표+근거)
# ------------------------------------------------------------
import os
import time
import random
import itertools
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)


def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30


START_MASK_RANGE = (7, 9)
MIN_MASK = 3
MAX_TRIES_PER_PROBLEM = 700

ROW_LABELS = ["A형", "X", "Y", "Z"]
COL_LABELS = ["BETA", "가", "나", "다"]

TYPES = ["A", "B", "AB", "O"]

PROPS_ALL = ["AgA", "AgB", "α", "β"]  # antigen A, antigen B, antibody alpha, antibody beta
PROP_DISPLAY = {"AgA": "응집원 A", "AgB": "응집원 B", "α": "응집소 α", "β": "응집소 β"}

ABO_TABLE = {
    "A":  {"AgA": "O", "AgB": "X", "α": "X", "β": "O"},
    "B":  {"AgA": "X", "AgB": "O", "α": "O", "β": "X"},
    "AB": {"AgA": "O", "AgB": "O", "α": "X", "β": "X"},
    "O":  {"AgA": "X", "AgB": "X", "α": "O", "β": "O"},
}


# =========================
# World build
# =========================
def build_full_grid() -> Tuple[Dict[str, Dict[str, str]], Dict]:
    others = ["B", "AB", "O"]
    random.shuffle(others)
    row_map = {"A형": "A", "X": others[0], "Y": others[1], "Z": others[2]}

    rest_props = ["AgA", "AgB", "α"]
    random.shuffle(rest_props)
    col_map = {"BETA": "β", "가": rest_props[0], "나": rest_props[1], "다": rest_props[2]}

    full = {r: {} for r in ROW_LABELS}
    for r in ROW_LABELS:
        bt = row_map[r]
        for c in COL_LABELS:
            prop = col_map[c]
            full[r][c] = ABO_TABLE[bt][prop]

    meta = {"row_map": row_map, "col_map": col_map}
    return full, meta


# =========================
# Masking
# =========================
def mask_grid(full: Dict[str, Dict[str, str]], n_mask: int) -> Dict[str, Dict[str, str]]:
    masked = {r: dict(full[r]) for r in ROW_LABELS}
    cells = [(r, c) for r in ROW_LABELS for c in COL_LABELS]
    random.shuffle(cells)
    for i in range(min(n_mask, len(cells))):
        r, c = cells[i]
        masked[r][c] = "?"
    return masked


# =========================
# Solver + "탈락 근거" 수집
# =========================
def first_contradiction(masked: Dict[str, Dict[str, str]], row_map: Dict[str, str], col_map: Dict[str, str]) -> Optional[Tuple[str, str, str, str]]:
    """
    returns first contradiction as (row_label, col_label, expected, observed)
    if none -> None
    """
    for r in ROW_LABELS:
        bt = row_map[r]
        for c in COL_LABELS:
            obs = masked[r][c]
            if obs == "?":
                continue
            prop = col_map[c]
            exp = ABO_TABLE[bt][prop]
            if exp != obs:
                return (r, c, exp, obs)
    return None


def count_solutions(masked: Dict[str, Dict[str, str]]) -> Tuple[int, Optional[Dict]]:
    sol = 0
    witness = None

    for perm_rows in itertools.permutations(["B", "AB", "O"], 3):
        row_map = {"A형": "A", "X": perm_rows[0], "Y": perm_rows[1], "Z": perm_rows[2]}
        for perm_cols in itertools.permutations(["AgA", "AgB", "α"], 3):
            col_map = {"BETA": "β", "가": perm_cols[0], "나": perm_cols[1], "다": perm_cols[2]}

            if first_contradiction(masked, row_map, col_map) is None:
                sol += 1
                witness = {"row_map": row_map, "col_map": col_map}
                if sol >= 2:
                    return sol, witness

    return sol, witness


def fill_full_from_witness(w: Dict) -> Dict[str, Dict[str, str]]:
    row_map = w["row_map"]
    col_map = w["col_map"]
    full = {r: {} for r in ROW_LABELS}
    for r in ROW_LABELS:
        bt = row_map[r]
        for c in COL_LABELS:
            prop = col_map[c]
            full[r][c] = ABO_TABLE[bt][prop]
    return full


def gather_reasons(masked: Dict[str, Dict[str, str]], witness: Dict, max_each: int = 2) -> Dict[str, List[str]]:
    """
    중간 버전 해설용 근거 문장 생성:
      - 가/나/다 각각에 대해: "다른 후보를 가정하면 특정 칸에서 모순" 1~2개
      - X/Y/Z 각각에 대해: "다른 혈액형을 가정하면 특정 칸에서 모순" 1~2개
    반환:
      reasons["col:가"] = [...]
      reasons["row:X"] = [...]
    """
    w_row = witness["row_map"]
    w_col = witness["col_map"]

    reasons: Dict[str, List[str]] = {}

    # ----- column reasons: for each label in 가/나/다
    for label in ["가", "나", "다"]:
        true_prop = w_col[label]
        key = f"col:{label}"
        reasons[key] = []

        for alt in ["AgA", "AgB", "α"]:
            if alt == true_prop:
                continue

            # keep other labels fixed, swap this one to alt and adjust the remaining to keep bijection
            # easiest: try all valid col permutations where label=alt, and others are any remaining
            candidates = []
            others = [p for p in ["AgA", "AgB", "α"] if p != alt]
            other_labels = [x for x in ["가", "나", "다"] if x != label]
            for perm in itertools.permutations(others, 2):
                col_map_try = {"BETA": "β"}
                col_map_try[label] = alt
                col_map_try[other_labels[0]] = perm[0]
                col_map_try[other_labels[1]] = perm[1]

                # rows fixed to witness row_map
                contra = first_contradiction(masked, w_row, col_map_try)
                if contra is not None:
                    candidates.append((col_map_try, contra))

            # pick one sharp contradiction
            if candidates:
                _, (r, c, exp, obs) = random.choice(candidates)
                # 자연어(중간 버전)
                alt_name = PROP_DISPLAY[alt]
                exp_txt = "O" if exp == "O" else "X"
                obs_txt = "O" if obs == "O" else "X"
                reasons[key].append(
                    f"{label}={alt_name}로 가정하면, {r}×{c}에서 예상은 {exp_txt}인데 제시값이 {obs_txt}이므로 모순이다."
                )
            if len(reasons[key]) >= max_each:
                break

    # ----- row reasons: for each label in X/Y/Z
    for label in ["X", "Y", "Z"]:
        true_type = w_row[label]
        key = f"row:{label}"
        reasons[key] = []

        for alt in ["B", "AB", "O"]:
            if alt == true_type:
                continue

            # try all valid row permutations with label=alt
            others = [t for t in ["B", "AB", "O"] if t != alt]
            other_labels = [x for x in ["X", "Y", "Z"] if x != label]
            candidates = []
            for perm in itertools.permutations(others, 2):
                row_map_try = {"A형": "A"}
                row_map_try[label] = alt
                row_map_try[other_labels[0]] = perm[0]
                row_map_try[other_labels[1]] = perm[1]

                # cols fixed to witness col_map
                contra = first_contradiction(masked, row_map_try, w_col)
                if contra is not None:
                    candidates.append((row_map_try, contra))

            if candidates:
                _, (r, c, exp, obs) = random.choice(candidates)
                exp_txt = "O" if exp == "O" else "X"
                obs_txt = "O" if obs == "O" else "X"
                reasons[key].append(
                    f"{label}={alt}형으로 가정하면, {r}×{c}에서 예상은 {exp_txt}인데 제시값이 {obs_txt}이므로 모순이다."
                )
            if len(reasons[key]) >= max_each:
                break

    # 근거가 하나도 못 잡힌 경우(마스킹이 너무 심할 때) 대비: 빈 리스트면 이후 출력 시 자동 스킵
    return reasons


# =========================
# Generate one problem (unique only)
# =========================
def generate_one_problem(problem_index: int):
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        full_true, _ = build_full_grid()

        n_mask = random.randint(*START_MASK_RANGE)
        while n_mask >= MIN_MASK:
            masked = mask_grid(full_true, n_mask)

            sol, w = count_solutions(masked)
            if sol == 1 and w is not None:
                full_solved = fill_full_from_witness(w)
                reasons = gather_reasons(masked, w, max_each=2)
                return {
                    "masked": masked,
                    "full": full_solved,
                    "witness": w,
                    "mask_used": n_mask,
                    "reasons": reasons,
                }

            n_mask -= 1

    raise RuntimeError(f"[P{problem_index:02d}] 생성 실패(유일정답 확보 불가)")


# =========================
# DOCX rendering
# =========================
def add_table(container, grid: Dict[str, Dict[str, str]]):
    t = container.add_table(rows=1 + len(ROW_LABELS), cols=1 + len(COL_LABELS))
    t.style = "Table Grid"

    t.cell(0, 0).text = ""
    for j, c in enumerate(COL_LABELS):
        t.cell(0, j + 1).text = c

    for i, r in enumerate(ROW_LABELS):
        t.cell(i + 1, 0).text = r
        for j, c in enumerate(COL_LABELS):
            t.cell(i + 1, j + 1).text = grid[r][c]


def fmt_mapping_line(witness: Dict) -> str:
    row_map = witness["row_map"]
    col_map = witness["col_map"]

    row_txt = f"A형=A형, X={row_map['X']}형, Y={row_map['Y']}형, Z={row_map['Z']}형"
    col_txt = (
        f"BETA={PROP_DISPLAY[col_map['BETA']]}, "
        f"가={PROP_DISPLAY[col_map['가']]}, 나={PROP_DISPLAY[col_map['나']]}, 다={PROP_DISPLAY[col_map['다']]}"
    )
    return f"[행 매핑] {row_txt} / [열 매핑] {col_txt}"



def grid_to_md(grid: Dict[str, Dict[str, str]], rows: List[str], cols: List[str]) -> str:
    header = [""] + [str(c) for c in cols]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for r in rows:
        row = [str(r)] + [str(grid[r][c]) for c in cols]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

def make_docx() :
    # output 폴더 생성
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="AGGLUT", out_dir=out_dir, id_prefix="AGGLUT_")

    os.makedirs(out_dir, exist_ok=True)
    # 저장 파일명 (시간 붙여서 중복 방지)
    filename = os.path.join(out_dir, f"Agglutination_{int(time.time())}.docx")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "바탕"
    set_page_margins(doc)
    style.font.size = Pt(9)

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        pdata = generate_one_problem(i)
        problems.append(pdata)
        # PACK record
        # PACK record (표시용 텍스트/표/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "아래의 표에서 가, 나, 다는 (응집원 A, 응집원 B, 응집소 α)를 순서 없이 나타낸 것이다.",
            "X, Y, Z는 (B형, AB형, O형)을 순서 없이 나타낸 것이다.",
        ])
        ask_line_md = "표에서 빈칸을 채우고, 가/나/다와 X/Y/Z를 모두 규명하시오."

        table_md = grid_to_md(pdata["masked"], ROW_LABELS, COL_LABELS)
        full_table_md = grid_to_md(pdata["full"], ROW_LABELS, COL_LABELS)

        answer_md = fmt_mapping_line(pdata["witness"])
        reasons = pdata.get("reasons", {})
        reason_lines = []
        for k, lst in reasons.items():
            for s in lst:
                reason_lines.append(f"- {k}: {s}")
        explanation_md = "\n".join(reason_lines) if reason_lines else "(근거 자동추출 실패: 마스킹이 강했을 수 있음)"

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": table_md,
            "full_table_md": full_table_md,
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "mask_used": pdata.get("mask_used"),
            "data": pdata,
        }))
        if i % 5 == 0:
            print(f"진행: {i}/{N_PROBLEMS}")

    # 문제(2단, 2문제/페이지)
    idx = 0
    pnum = 1
    while idx < N_PROBLEMS:
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        def put(cell, num, pdata):
            cell.add_paragraph(f"[문제 {num}]").runs[0].bold = True
            cell.add_paragraph("ABO식 혈액형에서 응집원과 응집소의 존재 여부를 O/X로 나타낸 것이다.")
            cell.add_paragraph("가, 나, 다는 (응집원 A, 응집원 B, 응집소 α)를 순서 없이 나타낸 것이다.")
            cell.add_paragraph("X, Y, Z는 (B형, AB형, O형)을 순서 없이 나타낸 것이다.")
            cell.add_paragraph("")
            add_table(cell, pdata["masked"])
            cell.add_paragraph("")
            cell.add_paragraph("표에서 빈칸(?)을 채우고 가, 나, 다와 X, Y, Z가 무엇을 의미하는지 찾으시오.")

        put(left, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            put(right, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            doc.add_page_break()

    # 정답
    doc.add_page_break()
    doc.add_paragraph("[정답]").runs[0].bold = True
    for i, pdata in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번: ?={pdata['mask_used']} | {fmt_mapping_line(pdata['witness'])}")

    # 해설(완성표 + 근거)
    doc.add_page_break()
    doc.add_paragraph("[해설(완성표+근거)]").runs[0].bold = True
    for i, pdata in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번").runs[0].bold = True
        doc.add_paragraph(fmt_mapping_line(pdata["witness"]))
        doc.add_paragraph("")

        # 근거 파트(중간 버전)
        doc.add_paragraph("결정적 근거(후보 가정 시 모순)")

        # col reasons
        for lab in ["가", "나", "다"]:
            true_prop = pdata["witness"]["col_map"][lab]
            doc.add_paragraph(f"- {lab}={PROP_DISPLAY[true_prop]}")
            lines = pdata["reasons"].get(f"col:{lab}", [])
            if lines:
                for s in lines[:2]:
                    doc.add_paragraph(f"  · {s}")
            else:
                doc.add_paragraph("  · (제시된 값만으로도 다른 경우가 모두 배제된다.)")

        # row reasons
        for lab in ["X", "Y", "Z"]:
            true_bt = pdata["witness"]["row_map"][lab]
            doc.add_paragraph(f"- {lab}={true_bt}형")
            lines = pdata["reasons"].get(f"row:{lab}", [])
            if lines:
                for s in lines[:2]:
                    doc.add_paragraph(f"  · {s}")
            else:
                doc.add_paragraph("  · (제시된 값만으로도 다른 경우가 모두 배제된다.)")

        doc.add_paragraph("")
        doc.add_paragraph("① 완성표(?, 없음)")
        add_table(doc, pdata["full"])
        doc.add_paragraph("② 제시표(물음표 포함)")
        add_table(doc, pdata["masked"])
        doc.add_paragraph("")

    doc.save(filename)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {filename}")


if __name__ == "__main__":
    make_docx()
