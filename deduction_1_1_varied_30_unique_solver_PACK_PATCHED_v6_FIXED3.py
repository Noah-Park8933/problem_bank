# deduction_1_1_varied_30_unique_solver.py
# ------------------------------------------------------------
# ✅ Deduction 1-1 자동 출제기 (형식 고정, 자료만 다양화)
# - (A/a), (B/b), (D/d) 중 2개는 4번(연관), 1개는 18번(독립)
# - 부모 유전자형 랜덤 생성(다양화)
# - 조건:
#   (1) 자손 표현형이 P1과 같을 확률 = P2와 같을 확률
#   (2) "자손이 가질 수 있는 유전자형에는 ____가 있다" (배열/연관 강제 단서)
# - ✅ 솔버: AB/AD/BD 3케이스 + 결합/반발(필요시) 모두 탐색
#          조건을 만족하는 해가 "유일"할 때만 문제로 채택
# - DOCX 출력(2단, 2문제/페이지) + 정답/해설
#
# 실행:
#   pip install python-docx
#   python deduction_1_1_varied_30_unique_solver.py
# ------------------------------------------------------------

import os, time, random
from dataclasses import dataclass
from typing import Dict, Tuple, List, Optional, Any

from docx import Document
from docx.shared import Pt

# ------------------------------------------------------------
# 표시용 유전자형 포맷터
# - dict가 그대로 출력되는 문제(예: {'A': 'Aa', ...}) 방지
# ------------------------------------------------------------
def fmt_geno(x) -> str:
    """유전자형 표시를 'AaBbdd'처럼 붙여서 출력.
    - dict이면 A,B,D,E... 정해진 순서대로 value만 이어붙임
    - 문자열이면 그대로 반환
    """
    if isinstance(x, dict):
        order = ["A","B","D","E","F","G","H","Q","R","T"]
        s = ""
        for k in order:
            if k in x:
                s += str(x[k])
        # 혹시 예상 밖 키가 있으면 알파벳 순으로 뒤에 붙임
        extra = sorted([k for k in x.keys() if k not in order])
        for k in extra:
            s += str(x[k])
        return s
    return str(x)

from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



# =========================
# CONFIG
# =========================
N_PROBLEMS = 30
MAX_TRIES_PER_PROBLEM = 4000
LOG_EVERY = 200

OUT_DIR = "output"
FONT_NAME = "바탕"
FONT_SIZE = 9

LOCI = ["A", "B", "D"]  # (A/a), (B/b), (D/d)

# 우열 고정
DOM = {"A": "A", "B": "B", "D": "D"}

# 완전연관(재조합 0)으로 고정: Deduction(배열 강제) 안정성 최강
R = 0.0

# 자손 "가능 유전자형" 후보들을 다양화(배열 강제용)
FORCED_POOL = [
    "aaBbDD", "aaBbDd", "aaBbdd",
    "AabbDD", "AabbDd", "Aabbdd",
    "AaBbDD", "AaBbdd", "AABbDD",
    "aaBBdd", "AAbbDD", "AaBBdd",
]

# =========================
# Utilities
# =========================
def norm_pair(x: str, y: str) -> str:
    # e.g., 'a','A' -> 'Aa'
    if x.isupper() and y.islower():
        return x + y
    if x.islower() and y.isupper():
        return y + x
    return x + y

def genotype_random(locus: str, allow_homo: bool = True) -> str:
    # returns 'AA','Aa','aa' style for given locus
    U = locus
    L = locus.lower()
    choices = [U+U, U+L, L+L] if allow_homo else [U+L]
    return random.choice(choices)

def pheno_from_genotype(gA: str, gB: str, gD: str) -> str:
    # phenotype label just as booleans: A- / bb / D-
    A_dom = (gA != "aa")
    B_dom = (gB != "bb")
    D_dom = (gD != "dd")
    return f"{'A-' if A_dom else 'aa'} {'B-' if B_dom else 'bb'} {'D-' if D_dom else 'dd'}"

def match_P1_P2_pheno(child: Dict[str,str], P1: Dict[str,str], P2: Dict[str,str]) -> Tuple[bool,bool]:
    # P1 phenotype defined by P1's genotype
    # P2 phenotype defined by P2's genotype
    ph_child = pheno_from_genotype(child["A"], child["B"], child["D"])
    ph_P1 = pheno_from_genotype(P1["A"], P1["B"], P1["D"])
    ph_P2 = pheno_from_genotype(P2["A"], P2["B"], P2["D"])
    return (ph_child == ph_P1, ph_child == ph_P2)

# =========================
# Linkage model (r=0)
# phase: "coupling" => AB/ab, "repulsion" => Ab/aB
# For a linked pair XY:
#  - if parent is hetero at both loci: phase matters
#  - if one/both loci homozygous: phase effectively irrelevant, gametes collapse
# =========================
def linked_gametes_for_parent(pairXY: Tuple[str,str], geno: Dict[str,str], phase: str) -> Dict[Tuple[str,str], float]:
    X, Y = pairXY
    gX = geno[X]
    gY = geno[Y]

    # allele options at each locus
    def alleles(g: str) -> List[str]:
        if g[0] == g[1]:
            return [g[0]]
        return [g[0], g[1]]  # e.g., ['A','a']

    Ax = alleles(gX)
    Ay = alleles(gY)

    # if not hetero at both, just independent product BUT we treat it as collapsed linked distribution
    # under r=0, still only possible combos are those compatible.
    if len(Ax) == 1 and len(Ay) == 1:
        return {(Ax[0], Ay[0]): 1.0}
    if len(Ax) == 1 and len(Ay) == 2:
        return {(Ax[0], Ay[0]): 0.5, (Ax[0], Ay[1]): 0.5}
    if len(Ax) == 2 and len(Ay) == 1:
        return {(Ax[0], Ay[0]): 0.5, (Ax[1], Ay[0]): 0.5}

    # both hetero (2x2) -> phase controls
    # coupling: XY and xy only (each 1/2)
    # repulsion: Xy and xY only (each 1/2)
    Ux, lx = X, X.lower()
    Uy, ly = Y, Y.lower()

    if phase == "coupling":
        return {(Ux, Uy): 0.5, (lx, ly): 0.5}
    else:
        return {(Ux, ly): 0.5, (lx, Uy): 0.5}

def indep_gametes_for_parent(locus: str, geno: Dict[str,str]) -> Dict[str,float]:
    g = geno[locus]
    if g[0] == g[1]:
        return {g[0]: 1.0}
    return {g[0]: 0.5, g[1]: 0.5}

def child_genotype_from_gametes(gam1: Dict[str,str], gam2: Dict[str,str]) -> Dict[str,str]:
    # gam dict locus->allele letter
    out = {}
    for L in LOCI:
        out[L] = norm_pair(gam1[L], gam2[L])
    return out

# =========================
# Compute probabilities for a given hypothesis (which pair linked + phases)
# =========================
def compute_probs(P1: Dict[str,str], P2: Dict[str,str],
                  linked_pair: Tuple[str,str],
                  phase1: str, phase2: str) -> Tuple[float,float, bool]:
    # linked_pair on chr4, remaining locus independent on chr18
    X, Y = linked_pair
    Z = [l for l in LOCI if l not in linked_pair][0]

    link1 = linked_gametes_for_parent(linked_pair, P1, phase1)
    link2 = linked_gametes_for_parent(linked_pair, P2, phase2)
    ind1 = indep_gametes_for_parent(Z, P1)
    ind2 = indep_gametes_for_parent(Z, P2)

    ph_P1 = pheno_from_genotype(P1["A"], P1["B"], P1["D"])
    ph_P2 = pheno_from_genotype(P2["A"], P2["B"], P2["D"])

    p_same_P1 = 0.0
    p_same_P2 = 0.0

    for (x1,y1), pr1 in link1.items():
        for (x2,y2), pr2 in link2.items():
            for z1, prz1 in ind1.items():
                for z2, prz2 in ind2.items():
                    pr = pr1*pr2*prz1*prz2
                    gam1 = {X:x1, Y:y1, Z:z1}
                    gam2 = {X:x2, Y:y2, Z:z2}
                    child = child_genotype_from_gametes(gam1, gam2)
                    ph_child = pheno_from_genotype(child["A"], child["B"], child["D"])
                    if ph_child == ph_P1:
                        p_same_P1 += pr
                    if ph_child == ph_P2:
                        p_same_P2 += pr

    # feasibility always true if distributions non-empty
    feasible = (sum(link1.values()) > 0 and sum(link2.values()) > 0)
    return p_same_P1, p_same_P2, feasible

def forced_genotype_possible(P1: Dict[str,str], P2: Dict[str,str],
                             linked_pair: Tuple[str,str],
                             phase1: str, phase2: str,
                             forced: str) -> bool:
    # forced like "aaBbDD" -> parse into target genotypes
    tA, tB, tD = forced[0:2], forced[2:4], forced[4:6]
    target = {"A": tA, "B": tB, "D": tD}

    X,Y = linked_pair
    Z = [l for l in LOCI if l not in linked_pair][0]

    link1 = linked_gametes_for_parent(linked_pair, P1, phase1)
    link2 = linked_gametes_for_parent(linked_pair, P2, phase2)
    ind1 = indep_gametes_for_parent(Z, P1)
    ind2 = indep_gametes_for_parent(Z, P2)

    for (x1,y1), pr1 in link1.items():
        for (x2,y2), pr2 in link2.items():
            if pr1*pr2 == 0: 
                continue
            for z1, prz1 in ind1.items():
                for z2, prz2 in ind2.items():
                    if prz1*prz2 == 0:
                        continue
                    gam1 = {X:x1, Y:y1, Z:z1}
                    gam2 = {X:x2, Y:y2, Z:z2}
                    child = child_genotype_from_gametes(gam1, gam2)
                    if child == target:
                        return True
    return False

# =========================
# Solver: enumerate all hypotheses and check conditions
# =========================
@dataclass
class Solution:
    linked_pair: Tuple[str,str]
    indep_locus: str
    phase1: str
    phase2: str
    p_same_P1: float
    p_same_P2: float

def all_hypotheses() -> List[Tuple[Tuple[str,str], str, str, str]]:
    pairs = [("A","B"), ("A","D"), ("B","D")]
    phases = ["coupling","repulsion"]
    out = []
    for pair in pairs:
        Z = [l for l in LOCI if l not in pair][0]
        for ph1 in phases:
            for ph2 in phases:
                out.append((pair, Z, ph1, ph2))
    return out

def solve_unique(P1: Dict[str,str], P2: Dict[str,str], forced: str, eps: float = 1e-9) -> Tuple[Optional[Solution], List[Solution]]:
    sols: List[Solution] = []
    for pair, Z, ph1, ph2 in all_hypotheses():
        p1, p2, feas = compute_probs(P1, P2, pair, ph1, ph2)
        if not feas:
            continue
        if abs(p1 - p2) > eps:
            continue
        if not forced_genotype_possible(P1, P2, pair, ph1, ph2, forced):
            continue
        sols.append(Solution(pair, Z, ph1, ph2, p1, p2))

    # "유일정답" 기준: linked_pair가 유일하게 결정되도록
    # (배열까지 유일하게 하고 싶으면 조건 강화하면 됨)
    if not sols:
        return None, sols

    unique_pairs = set(tuple(sorted(s.linked_pair)) for s in sols)
    if len(unique_pairs) != 1:
        return None, sols

    # pick one canonical solution
    # preference: phase info까지 유일하면 그대로, 아니면 첫 해를 대표로
    return sols[0], sols

# =========================
# Problem generation (varied parents + forced genotype)
# =========================
@dataclass
class Problem:
    P1: Dict[str,str]
    P2: Dict[str,str]
    forced: str
    sol: Solution
    all_solutions: List[Solution]

def random_parents() -> Tuple[Dict[str,str], Dict[str,str]]:
    # 다양화: P1은 대체로 2~3개 헤테로, P2는 1~3개 헤테로 랜덤
    # 너무 단순(전부 동형)이나 너무 불능(조건 안 맞음) 케이스는 솔버가 필터링
    while True:
        P1 = {L: genotype_random(L, allow_homo=True) for L in LOCI}
        P2 = {L: genotype_random(L, allow_homo=True) for L in LOCI}

        # 최소한 "P1 표현형"과 "P2 표현형"이 서로 달라야 문제 맛이 있음
        if pheno_from_genotype(P1["A"], P1["B"], P1["D"]) == pheno_from_genotype(P2["A"], P2["B"], P2["D"]):
            continue

        # 너무 쉬운 케이스 제한: 둘 다 거의 전부 동형이면 스킵
        def hetero_count(G): 
            return sum(1 for L in LOCI if G[L][0] != G[L][1])
        if hetero_count(P1) + hetero_count(P2) < 3:
            continue

        return P1, P2

def make_one_problem(pnum: int) -> Problem:
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        P1, P2 = random_parents()
        forced = random.choice(FORCED_POOL)

        sol, all_sols = solve_unique(P1, P2, forced)
        if sol is not None:
            return Problem(P1, P2, forced, sol, all_sols)

        if attempt % LOG_EVERY == 0:
            print(f"[P{pnum:02d}] attempt={attempt} still searching...")

    raise RuntimeError(f"[P{pnum:02d}] 유일정답 문제 생성 실패: MAX_TRIES_PER_PROBLEM 증가 또는 조건 완화 필요")

# =========================
# Pretty explanation (짧고 결정적 단서 중심)
# =========================
def phase_kor(ph: str) -> str:
    return "결합형" if ph == "coupling" else "반발형"

def sol_text(sol: Solution) -> str:
    a,b = sol.linked_pair
    return f"4번 염색체 연관: ({a}/{a.lower()})–({b}/{b.lower()}) , 18번 염색체: ({sol.indep_locus}/{sol.indep_locus.lower()})"

def build_explanation(pr: Problem) -> str:
    # 핵심 단서: forced genotype이 왜 연관/배열을 제한하는지 "틀"을 자동 작성
    a,b = pr.sol.linked_pair
    Z = pr.sol.indep_locus
    f = pr.forced

    # forced genotype에서 a,b가 어떤 조합을 요구하는지 간단히 언급
    # (완전 자동 자연어까지는 길어지니, 평가원식 "결정 단서" 느낌으로만)
    exp = []
    exp.append(f"자손에 {f}가 가능하므로, 4번 염색체에 있는 두 유전자좌에서 해당 대립유전자 조합이 배우자로 생성되어야 한다.")
    exp.append(f"따라서 연관쌍을 AB/AD/BD 중에서 시험하여 {a}–{b}만 조건(두 확률 같음 + {f} 가능)을 동시에 만족한다.")
    exp.append(f"(필요 시 배열) P1은 {phase_kor(pr.sol.phase1)}, P2는 {phase_kor(pr.sol.phase2)}로 해석된다.")
    exp.append(f"결론적으로 {sol_text(pr.sol)}.")
    exp.append(f"또한 P(자손 표현형=P1) = P(자손 표현형=P2) = {pr.sol.p_same_P1:.6f}.")
    return " ".join(exp)

# =========================
# DOCX
# =========================
def set_style(doc: Document):
    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = Pt(FONT_SIZE)

def add_problem_to_cell(cell, idx: int, pr: Problem):
    p = cell.add_paragraph(f"[문제 {idx}]")
    if p.runs:
        p.runs[0].bold = True

    cell.add_paragraph("형질 (가)는 A/a에 의해 결정되고 형질 (나)는 B/b에 의해, 형질 (다)는 D/d에 의해 결정된다.")
    cell.add_paragraph("유전자의 우열관계는 A>a이고 B>b이며 D>d이다.")
    cell.add_paragraph("(A/a), (B/b), (D/d) 중 2개는 4번 염색체에 있고 나머지 1개는 18번 염색체에 있다.")
    cell.add_paragraph(f"유전자형이 {pr.P1['A']}{pr.P1['B']}{pr.P1['D']}인 부모 P1과 {pr.P2['A']}{pr.P2['B']}{pr.P2['D']}인 부모 P2 사이에서 태어나는 자손의")
    cell.add_paragraph("표현형이 P1과 같을 확률과 자손의 표현형이 P2와 같을 확률은 같고")
    cell.add_paragraph(f"자손이 가질 수 있는 유전자형에는 {pr.forced}가 있다.")
    cell.add_paragraph("P1과 P2에서 유전자 사이의 관계를 구하고 자손의 표현형이 P1과 같을 확률과 자손의 표현형이 P2와 같을 확률을 구하시오.")

def add_two_column(doc: Document, problems: List[Problem]):
    idx = 0
    pnum = 1
    while idx < len(problems):
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        add_problem_to_cell(left, pnum, problems[idx])
        idx += 1
        pnum += 1

        if idx < len(problems):
            add_problem_to_cell(right, pnum, problems[idx])
            idx += 1
            pnum += 1

        if idx < len(problems):
            doc.add_page_break()

def add_answer_section(doc: Document, problems: List[Problem]):
    doc.add_page_break()
    h = doc.add_paragraph("[정답·해설]")
    if h.runs:
        h.runs[0].bold = True

    for i, pr in enumerate(problems, start=1):
        doc.add_paragraph(f"{i}번 정답")
        doc.add_paragraph(f" - {sol_text(pr.sol)}")
        doc.add_paragraph(f" - P(자손 표현형=P1) = P(자손 표현형=P2) = {pr.sol.p_same_P1:.6f}")
        doc.add_paragraph(f" - 해설: {build_explanation(pr)}")
        doc.add_paragraph("")

def make_docx():
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(out_dir, exist_ok=True)
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="DED11V", out_dir=out_dir, id_prefix="DED11V_")

    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, f"Deduction_1_1_varied_{int(time.time())}.docx")

    problems: List[Problem] = []
    for pnum in range(1, N_PROBLEMS + 1):
        pr = make_one_problem(pnum)
        problems.append(pr)
        # PACK record
        # PACK record (표시용 텍스트/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "형질 (가)는 A/a에 의해 결정되고 형질 (나)는 B/b에 의해, 형질 (다)는 D/d에 의해 결정된다.",
            "유전자의 우열관계는 A>a이고 B>b이며 D>d이다.",
            "(A/a), (B/b), (D/d) 중 2개는 4번 염색체에 있고 나머지 1개는 18번 염색체에 있다.",
            f"유전자형이 {fmt_geno(pr.P1)}인 부모 P1과 {fmt_geno(pr.P2)}인 부모 P2 사이에서 태어나는 자손의 표현형이 P1과 같을 확률과",
            f"자손의 표현형이 P2와 같을 확률은 같고 자손이 가질 수 있는 유전자형에는 {pr.forced}가 있다.",
        ])
        ask_line_md = "P1과 P2에서 유전자 사이의 관계를 구하고, 자손의 표현형이 P1과 같을 확률과 P2와 같을 확률을 구하시오."
        X, Y = pr.sol.linked_pair
        answer_md = f"4번 염색체 연관={X}-{Y}, 18번 염색체={pr.sol.indep_locus}; P(P1형)={pr.sol.p_same_P1:.3f}, P(P2형)={pr.sol.p_same_P2:.3f}"
        explanation_md = f"연관쌍={X}/{Y}, 상염색체={pr.sol.indep_locus}; 위상(P1)={pr.sol.phase1}, (P2)={pr.sol.phase2}"

        pack_writer.new_problem(qnum=pnum, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": "",
            "full_table_md": "",
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "data": {
                "P1": pr.P1,
                "P2": pr.P2,
                "forced": pr.forced,
                "solution": pr.sol.__dict__,
            },
        }))
        print(f"[진행] {pnum}/{N_PROBLEMS} 생성 완료")

    doc = Document()
    set_style(doc)

    add_two_column(doc, problems)
    add_answer_section(doc, problems)

    doc.save(out_path)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print("✅ 저장 완료:", out_path)

if __name__ == "__main__":
    make_docx()