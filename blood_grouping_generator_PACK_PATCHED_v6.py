# blood_grouping_generator.py
# ------------------------------------------------------------
# Blood Grouping (ABO) - 임의의 순서(랜덤 행/열 라벨 매핑)만 출제
# - 4명 혈액형: O, A, B, AB (모두 다름)
# - 표: 행=가,나,다,라 (적혈구), 열=I,II,III,IV (혈장)
# - 칸값: + / - / ?
# - 힌트: 특정 열 혈장에 α/β 정보 1개로 시도 → 유일정답 아니면 2개까지 자동 추가
# - 솔버(브루트포스 576)로 유일정답만 통과
# - DOCX 출력: 문제(2단, 2문제/페이지) + 정답 + 해설(완성표 포함)
# ------------------------------------------------------------
import os
import time
import random
import itertools
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm
from problem_pack import ProblemPack

def _to_jsonable(x):
    # Make any object JSON serializable (safe for PACK).
    if x is None or isinstance(x, (str, int, float, bool)):
        return x
    if isinstance(x, dict):
        return {str(k): _to_jsonable(v) for k, v in x.items()}
    if isinstance(x, (list, tuple, set)):
        return [_to_jsonable(v) for v in x]
    # dataclass / custom objects
    if hasattr(x, "__dict__"):
        try:
            return _to_jsonable(vars(x))
        except Exception:
            pass
    return str(x)



def set_page_margins(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# =========================
# CONFIG
# =========================
N_PROBLEMS = 30

# ? 개수(난이도): 기본 8개, 유일정답 안 나오면 자동으로 줄여감
TARGET_MASK = 8
MIN_MASK = 4

MAX_TRIES_PER_PROBLEM = 500

ROWS = ["가", "나", "다", "라"]               # RBC rows
COLS = ["I", "II", "III", "IV"]             # Plasma cols
TYPES = ["O", "A", "B", "AB"]

# 정답 분포(원하면 바꿔도 됨): 여기선 별도 "정답번호" 없음(혈액형 찾기형)
# 대신 유일정답만 통과

# =========================
# ABO rules (truth engine)
# =========================
RBC_ANTIGEN = {
    "O": set(),
    "A": {"A"},
    "B": {"B"},
    "AB": {"A", "B"},
}
PLASMA_AB = {  # antibodies
    "O": {"α", "β"},
    "A": {"β"},
    "B": {"α"},
    "AB": set(),
}

def reacts(rbc_type: str, plasma_type: str) -> str:
    ant = RBC_ANTIGEN[rbc_type]
    ab = PLASMA_AB[plasma_type]
    if ("α" in ab and "A" in ant) or ("β" in ab and "B" in ant):
        return "+"
    return "-"

# =========================
# Hint system
# =========================
HINT_KINDS = ["α", "β", "αβ", "없음"]  # plasma contains α / β / both / none

def hint_ok(plasma_type: str, kind: str) -> bool:
    ab = PLASMA_AB[plasma_type]
    if kind == "α":
        return "α" in ab
    if kind == "β":
        return "β" in ab
    if kind == "αβ":
        return ("α" in ab and "β" in ab)  # O형 혈장
    if kind == "없음":
        return (len(ab) == 0)            # AB형 혈장
    raise ValueError("bad hint kind")

def hint_sentence(col: str, kind: str) -> str:
    if kind == "α":
        return f"{col}의 혈장에는 응집소 α가 있다."
    if kind == "β":
        return f"{col}의 혈장에는 응집소 β가 있다."
    if kind == "αβ":
        return f"{col}의 혈장에는 응집소 α와 β가 모두 있다."
    if kind == "없음":
        return f"{col}의 혈장에는 응집소 α와 β가 없다."
    return ""

def pick_hint(existing_cols: set) -> Tuple[str, str]:
    # 가능하면 다른 열을 뽑기
    cols = COLS[:]
    random.shuffle(cols)
    col = None
    for c in cols:
        if c not in existing_cols:
            col = c
            break
    if col is None:
        col = random.choice(COLS)
    kind = random.choice(HINT_KINDS)
    return (col, kind)

# =========================
# Solver (unique check)
# =========================
def count_solutions(grid: Dict[str, Dict[str, str]], hints: List[Tuple[str, str]]):
    """
    grid[r][c] in {"+","-","?"}
    hints: list of (col, kind)
    returns (n_solutions, witness) witness=(r_map,c_map) for one solution
    """
    sol = 0
    witness = None

    for perm_r in itertools.permutations(TYPES, 4):
        r_map = dict(zip(ROWS, perm_r))
        for perm_c in itertools.permutations(TYPES, 4):
            c_map = dict(zip(COLS, perm_c))

            # hints
            ok = True
            for col, kind in hints:
                if not hint_ok(c_map[col], kind):
                    ok = False
                    break
            if not ok:
                continue

            # grid constraints
            for r in ROWS:
                for c in COLS:
                    v = grid[r][c]
                    if v == "?":
                        continue
                    if reacts(r_map[r], c_map[c]) != v:
                        ok = False
                        break
                if not ok:
                    break
            if not ok:
                continue

            sol += 1
            witness = (r_map, c_map)
            if sol >= 2:
                return sol, witness

    return sol, witness

# =========================
# World build (full table)
# =========================
def build_full_grid() -> Tuple[Dict[str, Dict[str, str]], Dict]:
    # assign true types to 4 people (distinct)
    people_types = TYPES[:]  # O,A,B,AB
    random.shuffle(people_types)

    # Random mapping: rows and cols each correspond to a permutation of the 4 people
    row_person = dict(zip(ROWS, people_types))  # row label -> RBC type
    # For columns we must permute the SAME four people types, but independently
    col_types = TYPES[:]
    random.shuffle(col_types)
    col_person = dict(zip(COLS, col_types))  # col label -> Plasma type

    full = {r: {} for r in ROWS}
    for r in ROWS:
        for c in COLS:
            full[r][c] = reacts(row_person[r], col_person[c])

    meta = {
        "row_type": row_person,  # RBC type per row label
        "col_type": col_person,  # Plasma type per col label
    }
    return full, meta

# =========================
# Masking
# =========================
def mask_grid(full: Dict[str, Dict[str, str]], n_mask: int) -> Dict[str, Dict[str, str]]:
    masked = {r: dict(full[r]) for r in ROWS}
    cells = [(r, c) for r in ROWS for c in COLS]
    random.shuffle(cells)

    for i in range(min(n_mask, 16)):
        r, c = cells[i]
        masked[r][c] = "?"
    return masked

# =========================
# Generate one problem (hint 1 -> 2)
# =========================
def generate_one_problem(problem_index: int):
    for attempt in range(1, MAX_TRIES_PER_PROBLEM + 1):
        full, meta = build_full_grid()

        # start masking; if not unique, reduce masks gradually
        n_mask = TARGET_MASK
        while n_mask >= MIN_MASK:
            masked = mask_grid(full, n_mask)

            # 1 hint
            hints = []
            used_cols = set()
            h1 = pick_hint(used_cols)
            hints = [h1]
            used_cols.add(h1[0])

            sol1, wit1 = count_solutions(masked, hints)
            if sol1 == 1:
                return full, masked, meta, hints, wit1, n_mask

            # 2 hints
            found = False
            for _ in range(25):
                h2 = pick_hint(used_cols)
                hints2 = [h1, h2]
                sol2, wit2 = count_solutions(masked, hints2)
                if sol2 == 1:
                    return full, masked, meta, hints2, wit2, n_mask

            # not unique even with 2 hints -> reduce masking
            n_mask -= 1

    raise RuntimeError(f"[P{problem_index:02d}] 생성 실패")

# =========================
# DOCX rendering
# =========================
def add_abo_table(container, grid: Dict[str, Dict[str, str]]):
    # rows: header + 4, cols: header + 4
    t = container.add_table(rows=1 + len(ROWS), cols=1 + len(COLS))
    t.style = "Table Grid"

    # header
    t.cell(0, 0).text = ""
    for j, c in enumerate(COLS):
        t.cell(0, j + 1).text = f"{c}의 혈장"

    # body
    for i, r in enumerate(ROWS):
        t.cell(i + 1, 0).text = f"{r}의 적혈구"
        for j, c in enumerate(COLS):
            t.cell(i + 1, j + 1).text = grid[r][c]


def grid_to_md(grid: Dict[str, Dict[str, str]], rows: List[str], cols: List[str], row_hdr_suffix: str = "", col_hdr_suffix: str = "") -> str:
    """dict-of-dict(+,-,?)를 Markdown 표로 변환."""
    header = [""] + [f"{c}{col_hdr_suffix}" for c in cols]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for r in rows:
        row = [f"{r}{row_hdr_suffix}"] + [str(grid[r][c]) for c in cols]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

def make_docx():
    # output 폴더 생성
    out_dir = os.path.join(os.path.dirname(__file__), "output")
    # --- PACK JSON setup (auto) ---
    pack_writer = ProblemPack(module_code="BLOODG", out_dir=out_dir, id_prefix="BLOODG_")

    os.makedirs(out_dir, exist_ok=True)
    # 저장 파일명 (시간 붙여서 중복 방지)
    filename = os.path.join(out_dir, f"BloodGrouping_{int(time.time())}.docx")
    doc = Document()
    set_page_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "바탕"
    style.font.size = Pt(9)

    problems = []
    for i in range(1, N_PROBLEMS + 1):
        full, masked, meta, hints, witness, mask_used = generate_one_problem(i)
        problems.append((full, masked, meta, hints, witness, mask_used))
        # PACK record
        # PACK record (표시용 텍스트/표/해설 포함: 웹 문제은행 호환)
        problem_text_md = "\n".join([
            "아래의 표는 ABO식 혈액형이 모두 다른 4명의 혈액의 응집반응을 나타낸 것이다.",
            "I, II, III, IV는 가, 나, 다, 라를 순서없이 나타낸 것이다.",
        ] + [f"※ {hint_sentence(col, kind)}" for (col, kind) in hints])
        ask_line_md = "표에서 빈칸을 채우고 가, 나, 다, 라와 I, II, III, IV의 ABO식 혈액형을 찾으시오."

        table_md = grid_to_md(masked, ROWS, COLS, row_hdr_suffix="의 적혈구", col_hdr_suffix="의 혈장")
        full_table_md = grid_to_md(full, ROWS, COLS, row_hdr_suffix="의 적혈구", col_hdr_suffix="의 혈장")

        r_map, c_map = witness
        explanation_md = "\n".join([
            f"- 힌트: " + "; ".join([hint_sentence(col, kind) for col, kind in hints]),
            f"- 적혈구(가/나/다/라) 매핑: {r_map}",
            f"- 혈장(I/II/III/IV) 매핑: {c_map}",
        ])
        answer_md = f"적혈구={r_map} / 혈장={c_map}"

        pack_writer.new_problem(qnum=i, payload=_to_jsonable({
            "problem_text_md": problem_text_md,
            "ask_line_md": ask_line_md,
            "table_md": table_md,
            "full_table_md": full_table_md,
            "answer_md": answer_md,
            "explanation_md": explanation_md,
            "mask_used": mask_used,
            "hints": hints,
            "data": {
                "full": full,
                "masked": masked,
                "meta": meta,
                "witness": witness,
            },
        }))

    # problems (2 columns, 2 per page)
    idx = 0
    pnum = 1
    while idx < N_PROBLEMS:
        tbl = doc.add_table(rows=1, cols=2)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]

        def put(cell, pnum, pdata):
            full, masked, meta, hints, witness, mask_used = pdata
            h = cell.add_paragraph(f"[문제 {pnum}]")
            h.runs[0].bold = True
            cell.add_paragraph("아래의 표는 ABO식 혈액형이 모두 다른 4명의 혈액의 응집반응을 나타낸 것이다.")
            cell.add_paragraph("I, II, III, IV는 가, 나, 다, 라를 순서없이 나타낸 것이다.")
            for (col, kind) in hints:
                cell.add_paragraph("※ " + hint_sentence(col, kind))
            cell.add_paragraph("")
            add_abo_table(cell, masked)
            cell.add_paragraph("")
            cell.add_paragraph("표에서 빈칸을 채우고 가, 나, 다, 라와 I, II, III, IV의 ABO식 혈액형을 찾으시오.")

        put(left, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            put(right, pnum, problems[idx]); idx += 1; pnum += 1
        if idx < N_PROBLEMS:
            doc.add_page_break()

    # Answers
    doc.add_page_break()
    h = doc.add_paragraph("[정답]")
    h.runs[0].bold = True

    for i, pdata in enumerate(problems, start=1):
        full, masked, meta, hints, witness, mask_used = pdata
        r_map, c_map = witness  # label->type
        # Convert col plasma types to "I=??" etc
        hint_txt = "; ".join([hint_sentence(col, kind) for col, kind in hints])
        doc.add_paragraph(
            f"{i}번 | ?={mask_used} | 힌트: {hint_txt} | "
            f"적혈구(가/나/다/라)={r_map} | 혈장(I/II/III/IV)={c_map}"
        )

    # Explanation: full table
    doc.add_page_break()
    h2 = doc.add_paragraph("[해설(완성표)]")
    h2.runs[0].bold = True

    for i, pdata in enumerate(problems, start=1):
        full, masked, meta, hints, witness, mask_used = pdata
        title = doc.add_paragraph(f"{i}번")
        title.runs[0].bold = True
        for (col, kind) in hints:
            doc.add_paragraph("조건: " + hint_sentence(col, kind))
        doc.add_paragraph("① 완성표(?, 없음)")
        add_abo_table(doc, full)
        doc.add_paragraph("② 제시표(물음표 포함)")
        add_abo_table(doc, masked)
        doc.add_paragraph("")

    doc.save(filename)
    # --- PACK JSON save (auto) ---
    pack_path = pack_writer.save_json()
    print(f"✅ PACK JSON 저장: {pack_path}")
    print(f"✅ 저장 완료: {filename}")

if __name__ == "__main__":
    make_docx()